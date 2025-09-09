"""
🧰 Helper functions for the automated label generation tool.

Includes:
- fill_full_page(): populates a DOCX template with ID values for printing.
- Data processing utilities with DRY pattern support.

Example:
    from .utils import fill_full_page
"""

from docx import Document
from docx.shared import Pt
from copy import deepcopy
from typing import List, Tuple, Union, Optional, Dict, Any
from pathlib import Path
import pandas as pd

# Import common utilities using the cu pattern
from ...common import cu

SETS_PER_PAGE = 8  # Number of ID sets per label sheet


def fill_full_page(template_doc: Document, id_pairs: List[Tuple[str, str, str]]) -> Document:
    """
    Fills a Word template with up to 8 sets of IDs (RID, MID, Visit).

    Args:
        template_doc (Document): A loaded Word document with placeholders.
        id_pairs (list of tuples): List of (RID, MID, Visit) values.

    Returns:
        Document: A modified document page with all placeholders replaced.
    """
    page = deepcopy(template_doc)

    def replace_placeholders(paragraphs, rid: str, mid: str, visit: str, idx: int) -> None:
        target_rid = f'{{RID {idx}}}'
        target_mid = f'{{MID {idx}}}'
        target_v = f'{{V {idx}}}'

        for para in paragraphs:
            full_text = ''.join(run.text for run in para.runs)
            if target_rid in full_text or target_mid in full_text or target_v in full_text:
                full_text = (
                    full_text.replace(target_rid, str(rid))
                             .replace(target_mid, str(mid))
                             .replace(target_v, str(visit))
                )
                # Clear all runs
                for run in para.runs:
                    run.text = ''
                # Write updated text to the first run
                if para.runs:
                    para.runs[0].text = full_text
                    font = para.runs[0].font
                    font.name = 'Calibri'
                    font.size = Pt(11)

    for idx in range(1, SETS_PER_PAGE + 1):
        rid, mid, visit = ('', '', '')
        if idx <= len(id_pairs):
            rid, mid, visit = id_pairs[idx - 1]

        replace_placeholders(page.paragraphs, rid, mid, visit, idx)

        for table in page.tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_placeholders(cell.paragraphs, rid, mid, visit, idx)

    return page


def apply_labeling_rules(data: pd.DataFrame, rules: Optional[Dict[str, Any]] = None, domain: Optional[str] = None) -> pd.DataFrame:
    """
    Apply labeling rules to the data.
    
    Args:
        data: DataFrame to label
        rules: Labeling rules to apply
        domain: Optional domain to filter labeling
        
    Returns:
        DataFrame with labels applied
    """
    # Create a copy to avoid modifying original data
    labeled_data = data.copy()
    
    # Add a basic label column if it doesn't exist
    if 'label' not in labeled_data.columns:
        labeled_data['label'] = 'unlabeled'
    
    # Apply domain filtering if specified
    if domain and 'domain' in labeled_data.columns:
        domain_mask = labeled_data['domain'] == domain
        labeled_data.loc[~domain_mask, 'label'] = 'out_of_domain'
    
    # Apply custom rules if provided
    if rules:
        for rule_name, rule_config in rules.items():
            # This is a placeholder for rule application logic
            # In a real implementation, you would parse and apply the rules
            pass
    
    return labeled_data


def save_labeled_data(data: pd.DataFrame, output_path: Union[str, Path], format: str = 'excel') -> None:
    """
    Save labeled data to file.
    
    Args:
        data: DataFrame to save
        output_path: Path where to save the data
        format: Output format ('excel' or 'csv')
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    if format.lower() == 'excel':
        data.to_excel(output_path, index=False)
    else:
        data.to_csv(output_path, index=False)


def process_data(input_path: Union[str, Path], **kwargs: Any) -> pd.DataFrame:
    """Process input data file."""
    return cu.load_data(input_path, **kwargs)
