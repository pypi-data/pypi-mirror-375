"""
pip install python-docx
"""
import os
import re
import uuid
from collections import defaultdict
from dataclasses import dataclass
from enum import Enum
from urllib.parse import urljoin
from typing import Union

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.table import _Cell
from docx.text.paragraph import Paragraph


class SectionType(str, Enum):
    HEADING = 'heading'
    PARAGRAPH = 'paragraph'
    TABLE = 'table'


@dataclass
class Section:
    type: SectionType
    content: Union[str, None] = None
    level: Union[str, None] = None
    contains_img: bool = False


class DocxParser:

    def __init__(self, docx_path: str, image_handler=None, preserve_style=True):
        self.docx_path = docx_path
        self.doc = Document(self.docx_path)
        if image_handler is None:
            image_handler = DocxParser.save_img_to_local
        self.image_handler = image_handler
        self.preserve_style = preserve_style

        self.para_number_set = defaultdict(list)
        self.numbering_part = {}

    def parse(self) -> list[Section]:
        sections = []
        self.para_number_set.clear()
        self.numbering_part = self.get_numbering_part(self.doc)
        for child in self.doc.element.body.iterchildren():
            if child.tag.endswith('p'):
                paragraph = Paragraph(child, self.doc)
                section = self.handle_paragraph(paragraph)
                if section:
                    sections.append(section)
            elif child.tag.endswith('tbl'):
                table = Table(child, self.doc)
                section = self.handle_table(table)
                sections.append(section)

        return sections

    def to_markdown(self) -> str:
        """
        Read the docx and parse it to markdown.
        :param docx_path:
        :param image_handler: The handler who responsible for storing images.
         The signature is image_handler(file_bytes, file_ext)
        :return: markdown
        """
        sections = self.parse()

        texts = []
        for section in sections:
            text = section.content
            if section.type != SectionType.HEADING:
                text += '\n'
            texts.append(text)
        return "\n".join(texts)  # md中，两个空格+换行才能表示换行

    def handle_paragraph(self, paragraph: Paragraph) -> Union[Section, None]:
        texts = []
        level = self.get_paragraph_heading_level(paragraph)
        type = SectionType.PARAGRAPH
        if level > 0:
            texts.append('#' * level + ' ')
            type = SectionType.HEADING
        numbering = self.get_paragraph_numbering(paragraph)
        if numbering:
            texts.append(numbering)

        # 空标题跳过，markdown解析器无法正确识别
        if not paragraph.runs and level > 0:
            return None

        buffer = []
        current_style = None

        def flush_text_buffer():
            if not buffer:
                return
            buffer_text = ''.join(buffer)
            if self.preserve_style and buffer_text.strip():
                if current_style == "bold":
                    prefix, core, suffix = self.split_text_by_whitespace(buffer_text)
                    buffer_text = f"{prefix}**{core}**{suffix}"
                elif current_style == "italic":
                    prefix, core, suffix = self.split_text_by_whitespace(buffer_text)
                    buffer_text = f"{prefix}*{core}*{suffix}"
            texts.append(buffer_text)
            buffer.clear()

        for run in paragraph.runs:
            if run.text:
                text = run.text
                style = "bold" if run.bold else "italic" if run.italic else "normal"
                if style != current_style:
                    flush_text_buffer()
                    current_style = style
                buffer.append(text)

            if run.element.drawing_lst:
                for drawing in run.element.drawing_lst:
                    image_data, ext = self.extract_image_bytes_and_ext(drawing, run.part)
                    if image_data:
                        flush_text_buffer()
                        img_path = self.image_handler(image_data, ext)
                        image_name = os.path.splitext(os.path.basename(img_path))[0]
                        texts.append(f'\n![{image_name}]({img_path})\n')

        flush_text_buffer()
        text = ''.join(texts)
        if not text.strip():  # 去掉空白行
            return None
        if re.fullmatch(r'\s*-{3,}\s*', text):  # 分隔线前面加个换行才会有效果
            text = '\n' + text
        return Section(type=type, content=text)

    @staticmethod
    def split_text_by_whitespace(text: str) -> tuple[str, str, str]:
        match = re.match(r'^(\s*)(.*?)(\s*)$', text, re.DOTALL)
        if match:
            return match.group(1), match.group(2), match.group(3)
        return '', text, ''

    @staticmethod
    def get_paragraph_heading_level(paragraph: Paragraph):
        if paragraph.style is not None:
            if paragraph.style.name.startswith('Heading'):
                # 使用样式名中的 "Heading X"
                return int(paragraph.style.name.split(' ')[-1])
            else:
                # 查找样式中的大纲级别
                style_element = paragraph.style._element
                pPr = style_element.find(qn('w:pPr'))
                if pPr is not None:
                    outlineLvl = pPr.find(qn('w:outlineLvl'))
                    if outlineLvl is not None:
                        outline_level = int(outlineLvl.get(qn('w:val')))
                        return outline_level + 1

        # 查找段落中的大纲级别
        pPr = paragraph._element.find(qn('w:pPr'))
        if pPr is not None:
            outlineLvl = pPr.find(qn('w:outlineLvl'))
            if outlineLvl is not None:
                outline_level = int(outlineLvl.get(qn('w:val')))
                return outline_level + 1

        return 0

    def get_paragraph_numbering(self, paragraph: Paragraph):
        pPr = paragraph._p.find(qn('w:pPr'))
        if pPr is None:
            return ''
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None:
            # 如果存在编号，取具体编号值（这里只是示例，真实编号格式需要结合 numbering.xml 判断）
            ilvl = numPr.find(qn('w:ilvl'))
            ilvl_val = ilvl.val if ilvl is not None else 0
            numId = numPr.find(qn('w:numId'))
            numId_val = numId.val if numId is not None else 0
            if numId_val > 0:
                para_number_arr = self.para_number_set[numId_val]
                while ilvl_val >= len(para_number_arr):
                    para_number_arr.append(0)
                para_num = para_number_arr[ilvl_val] + 1
                para_number_arr[ilvl_val] = para_num
                for i in range(ilvl_val + 1, len(para_number_arr)):
                    para_number_arr[i] = 0

                number_str = DocxNumberFormatter.generate_numbering(
                    para_number_arr[:ilvl_val + 1], self.numbering_part.get(str(numId_val), {})
                )
                if number_str:
                    return '    ' * ilvl_val + number_str + ' '
        return ''

    @staticmethod
    def extract_image_bytes_and_ext(drawing, part):
        """从 drawing XML 中提取图片数据及扩展名"""
        # 虑组合图片的场景
        blip = drawing.findall('.//' + qn('a:blip'))
        if blip:
            rid = blip[0].get(qn('r:embed'))
            image_part  = part.related_parts[rid]
            content_type = image_part .content_type  # 例如：image/png、image/jpeg
            ext = content_type.split('/')[-1]  # 提取扩展名部分，如 png、jpeg
            # 一些扩展名需要调整为标准格式（如 jpeg -> jpg）
            ext = {'jpeg': 'jpg'}.get(ext, ext)
            return image_part.blob, ext
        return None, None

    def handle_table(self, table: Table) -> Section:
        html = ['<table>']

        def get_cell_vMerge(tr_list, row_idx, col_idx):
            """获取指定单元格的纵向合并标志：continue代表是上一行单元格的合并部分"""
            if row_idx < len(tr_list):
                row = tr_list[row_idx]
                current_col_idx = 0
                for cell in row.tc_lst:
                    current_col_idx += cell.grid_span
                    if current_col_idx == col_idx:
                        return cell.vMerge
            return None

        for row_idx, row in enumerate(table._tbl.tr_lst):
            html.append('<tr>')
            col_idx = 0
            for tc in row.tc_lst:
                rowspan = 1
                colspan = tc.grid_span
                col_idx += colspan
                # 处理 vMerge（纵向合并 rowspan）
                if tc.vMerge == 'continue':
                    continue
                elif tc.vMerge == 'restart':
                    idx = row_idx + 1
                    while idx < len(table._tbl.tr_lst):
                        vMerge = get_cell_vMerge(table._tbl.tr_lst, idx, col_idx)
                        idx += 1
                        if vMerge == 'continue':
                            rowspan += 1
                        else:
                            break

                # 构造属性
                attr = []
                if rowspan > 1:
                    attr.append(f'rowspan="{rowspan}"')
                if colspan > 1:
                    attr.append(f'colspan="{colspan}"')

                # 构造临时 cell 获取内容
                cell = _Cell(tc, table)
                texts = []
                for paragraph in cell.paragraphs:
                    section = self.handle_paragraph(paragraph)
                    if section:
                        texts.append(section.content)
                content = '\n'.join(texts).strip()

                html.append(f'<td{(" " + " ".join(attr)) if attr else ""}>{content}</td>')
                col_idx += colspan
            html.append('</tr>')
        html.append('</table>')
        return Section(type=SectionType.TABLE, content=''.join(html))

    @staticmethod
    def get_numbering_part(doc):
        try:
            numbering_part_element = doc.part.numbering_part.element
        except NotImplementedError:
            return {}
        w_val = qn('w:val')
        # 1. Locate all w:num elements. Each w:num element contains a numId attribute that associates it with a paragraph.
        w_num = qn('w:num')
        num_elements = numbering_part_element.findall(w_num)

        # 2. Locate all abstractNumId elements. Typically, there is one abstractNumId element.
        absNumId_to_numId = defaultdict(list)
        w_abstractNumId = qn('w:abstractNumId')
        w_numId = qn('w:numId')
        for num_element in num_elements:
            abstractNumId = num_element.findall(w_abstractNumId)
            if len(abstractNumId) == 0:
                continue
            abstractNumId = abstractNumId[0]
            abstractNumId = abstractNumId.get(w_val)
            numId = num_element.get(w_numId)
            if abstractNumId is not None and numId is not None:
                absNumId_to_numId[abstractNumId].append(numId)

        # 3. Locate all abstractNum elements.
        w_abstractNum = qn('w:abstractNum')
        abstractNum_elements = numbering_part_element.findall(w_abstractNum)

        # 4. Within each abstractNum element, examine the abstractNumId, lvl, lvlText, and numFmt elements.
        # Under normal circumstances, there would be only one lvlText and one numFmt element for each level.
        w_lvl = qn('w:lvl')
        w_ilvl = qn('w:ilvl')
        w_lvlText = qn('w:lvlText')
        w_numFmt = qn('w:numFmt')

        numbering_part = {}
        for abstractNum_element in abstractNum_elements:
            abstractNumId = abstractNum_element.get(w_abstractNumId)
            if abstractNumId is None:
                continue
            bucket = {}
            lvl_elements = abstractNum_element.findall(w_lvl)
            numFmt = 'decimal'
            for lvl_element in lvl_elements:
                ilvl = lvl_element.get(w_ilvl)
                if ilvl is None:
                    continue

                lvlText_elements = lvl_element.findall(w_lvlText)
                numFmt_elements = lvl_element.findall(w_numFmt)
                if len(lvlText_elements) == 0 or len(numFmt_elements) == 0:
                    continue
                numFmt = numFmt_elements[0].get(w_val)
                text = lvlText_elements[0].get(w_val)
                bucket.update({ilvl: [text, numFmt]})

            if abstractNumId in absNumId_to_numId.keys():
                for numId in absNumId_to_numId[abstractNumId]:
                    numbering_part.update({
                        numId: bucket
                    })
        return numbering_part

    @staticmethod
    def save_img_to_local(image_bytes, ext, prefix='images'):
        os.makedirs(prefix, exist_ok=True)
        file_name = f"{uuid.uuid4().hex}.{ext}"
        img_path = urljoin(prefix + '/', file_name)
        with open(img_path, 'wb') as f:
            f.write(image_bytes)
        return img_path

    @classmethod
    def _run_contains_chart(cls, run) -> bool:
        """
        判断一个 Word run 是否包含图表（chart）
        """
        if run.element.drawing_lst:
            for drawing in run.element.drawing_lst:
                graphic_data_list = drawing.findall('.//' + qn('a:graphicData'))
                for graphic_data in graphic_data_list:
                    # 通过 tag 判断是否为 chart
                    for child in graphic_data:
                        if child.tag.endswith('chart') or child.tag.endswith('diagram') \
                                or child.tag.endswith('oleobject') or child.tag.endswith('grpSp'):
                            return True
        return False

    def contains_chart(self):
        for para in self.doc.paragraphs:
            for run in para.runs:
                if self._run_contains_chart(run):
                    return True
        return False


class DocxNumberFormatter:

    @staticmethod
    def int_to_lower_letter(n):
        """1 -> a, 2 -> b, ... 27 -> aa"""
        result = ""
        while n > 0:
            n -= 1
            result = chr(97 + n % 26) + result
            n //= 26
        return result

    @staticmethod
    def int_to_lower_roman(n):
        """1 -> i, 2 -> ii, etc."""
        val_map = [
            (1000, 'm'), (900, 'cm'), (500, 'd'), (400, 'cd'),
            (100, 'c'), (90, 'xc'), (50, 'l'), (40, 'xl'),
            (10, 'x'), (9, 'ix'), (5, 'v'), (4, 'iv'), (1, 'i')
        ]
        result = ""
        for val, sym in val_map:
            while n >= val:
                result += sym
                n -= val
        return result

    @staticmethod
    def get_styled_value(val: int, style: str) -> str:
        if style == 'decimal':
            return str(val)
        elif style == 'lowerLetter':
            return DocxNumberFormatter.int_to_lower_letter(val)
        elif style == 'upperLetter':
            return DocxNumberFormatter.int_to_lower_letter(val).upper()
        elif style == 'lowerRoman':
            return DocxNumberFormatter.int_to_lower_roman(val)
        elif style == 'upperRoman':
            return DocxNumberFormatter.int_to_lower_roman(val).upper()
        return str(val)

    @staticmethod
    def generate_numbering(level_values, numbering_map):
        """
        level_values: [1, 2, 3] 表示各层级序号值（从1开始）
        numbering_map: dict，层级样式定义
        """
        level = str(len(level_values) - 1)
        if level not in numbering_map:
            return ''
        fmt, style = numbering_map[level]

        # 遍历格式中的 %1 ~ %9，并替换为带格式的值
        for i, val in enumerate(level_values):
            placeholder = f'%{i + 1}'
            if placeholder in fmt:
                fmt = fmt.replace(placeholder, DocxNumberFormatter.get_styled_value(val, style))

        return fmt


if __name__ == "__main__":
    docx_path = r'C:\Users\zhousf-a\Desktop\图片敏感信息识别技术方案.docx'
    parser = DocxParser(docx_path)
    sections = parser.parse()
    for section in sections:
        print(f"Type: {section.type}, Content: {section.content}, Level: {section.level}, Contains Image: {section.contains_img}")

    markdown = parser.to_markdown()
    print(markdown)

    if parser.contains_chart():
        print("The document contains charts.")
    else:
        print("No charts found in the document.")
    # 如果需要将markdown保存到文件
    with open('output.md', 'w', encoding='utf-8') as f:
        f.write(markdown)
    print("Markdown content saved to output.md")
