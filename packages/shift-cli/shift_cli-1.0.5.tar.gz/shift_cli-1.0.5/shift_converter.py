#!/usr/bin/env python3
"""
Shift - Universal Document Converter
Convert between various document formats: PDF, Word, HTML, Markdown, Text, etc.

Usage:
    shift document.docx --to pdf
    shift report.md --to html --css style.css
    shift file.pdf --to text --output extracted.txt
    shift folder/ --batch --from docx --to pdf
"""

import argparse
import sys
from pathlib import Path
from typing import List, Optional, Dict, Any
import tempfile
import shutil
import subprocess

try:
    import pypdf  # type: ignore
    from fpdf import FPDF  # type: ignore
    import markdown  # type: ignore
    from bs4 import BeautifulSoup  # type: ignore
    import docx  # type: ignore
    from docx2python import docx2python  # type: ignore
    import pdfkit  # type: ignore
    import html2text  # type: ignore
except ImportError as e:
    missing_package = str(e).split("'")[1] if "'" in str(e) else str(e)
    print(f"Missing required package: {missing_package}")
    print("Install with: pip install pypdf fpdf2 markdown beautifulsoup4 python-docx docx2python pdfkit html2text")
    print("Note: Some conversions may require additional tools like wkhtmltopdf, pandoc, or LibreOffice")
    sys.exit(1)


def shift_converter_main(input_path=None, output_path=None):
    print(f"Converting {input_path} to {output_path}")
    # Your logic here

class DocumentConverter:
    def __init__(self):
        """Initialize the document converter with supported formats and converters."""
        
        # Define supported formats and their converters
        self.converters = {
            # PDF conversions
            ('pdf', 'text'): self.pdf_to_text,
            ('pdf', 'html'): self.pdf_to_html,
            ('pdf', 'md'): self.pdf_to_markdown,
            
            # Word document conversions
            ('docx', 'pdf'): self.docx_to_pdf,
            ('docx', 'html'): self.docx_to_html,
            ('docx', 'text'): self.docx_to_text,
            ('docx', 'md'): self.docx_to_markdown,
            
            # Markdown conversions
            ('md', 'html'): self.markdown_to_html,
            ('md', 'pdf'): self.markdown_to_pdf,
            ('md', 'docx'): self.markdown_to_docx,
            
            # HTML conversions
            ('html', 'pdf'): self.html_to_pdf,
            ('html', 'text'): self.html_to_text,
            ('html', 'md'): self.html_to_markdown,
            
            # Text conversions
            ('text', 'pdf'): self.text_to_pdf,
            ('text', 'html'): self.text_to_html,
            ('text', 'md'): self.text_to_markdown,
        }
        
        # Format aliases
        self.format_aliases = {
            'txt': 'text',
            'markdown': 'md',
            'htm': 'html',
            'doc': 'docx',  # Will try to convert via LibreOffice if available
        }
        
        # Check for external tools
        self.external_tools = {
            'pandoc': shutil.which('pandoc'),
            'wkhtmltopdf': shutil.which('wkhtmltopdf'),
            'libreoffice': shutil.which('libreoffice'),
        }
    
    def normalize_format(self, format_str: str) -> str:
        """Normalize format string and handle aliases."""
        format_str = format_str.lower().lstrip('.')
        return self.format_aliases.get(format_str, format_str)
    
    def get_supported_conversions(self) -> List[tuple]:
        """Get list of all supported conversion pairs."""
        conversions = list(self.converters.keys())
        
        # Add pandoc conversions if available
        if self.external_tools['pandoc']:
            pandoc_formats = ['docx', 'pdf', 'html', 'md', 'text', 'rtf', 'odt', 'epub']
            for from_fmt in pandoc_formats:
                for to_fmt in pandoc_formats:
                    if from_fmt != to_fmt and (from_fmt, to_fmt) not in conversions:
                        conversions.append((from_fmt, to_fmt))
        
        return sorted(conversions)
    
    # PDF conversion methods
    def pdf_to_text(self, input_path: Path, output_path: Path, **kwargs) -> bool:
        """Convert PDF to plain text."""
        try:
            reader = pypdf.PdfReader(input_path)
            text_content = []
            
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text.strip():
                    text_content.append(f"--- Page {i+1} ---\n{text}\n")
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(text_content))
            
            return True
        except Exception as e:
            print(f"Error converting PDF to text: {e}")
            return False
    
    def pdf_to_html(self, input_path: Path, output_path: Path, **kwargs) -> bool:
        """Convert PDF to HTML."""
        try:
            # First convert to text
            text_content = []
            reader = pypdf.PdfReader(input_path)
            
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text.strip():
                    text_content.append(f"<h2>Page {i+1}</h2>\n<pre>{text}</pre>\n")
            
            html_content = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <title>{input_path.stem}</title>
                <meta charset="utf-8">
                <style>
                    body {{ font-family: Arial, sans-serif; margin: 40px; }}
                    pre {{ white-space: pre-wrap; background: #f5f5f5; padding: 15px; }}
                    h2 {{ color: #333; border-bottom: 1px solid #ddd; }}
                </style>
            </head>
            <body>
                <h1>{input_path.stem}</h1>
                {''.join(text_content)}
            </body>
            </html>
            """
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            return True
        except Exception as e:
            print(f"Error converting PDF to HTML: {e}")
            return False
    
    def pdf_to_markdown(self, input_path: Path, output_path: Path, **kwargs) -> bool:
        """Convert PDF to Markdown."""
        try:
            reader = pypdf.PdfReader(input_path)
            md_content = [f"# {input_path.stem}\n"]
            
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text.strip():
                    md_content.append(f"## Page {i+1}\n\n{text}\n")
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(md_content))
            
            return True
        except Exception as e:
            print(f"Error converting PDF to Markdown: {e}")
            return False
    
    # Word document conversion methods
    def docx_to_text(self, input_path: Path, output_path: Path, **kwargs) -> bool:
        """Convert Word document to plain text."""
        try:
            doc = docx.Document(str(input_path))
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(text_content))
            
            return True
        except Exception as e:
            print(f"Error converting DOCX to text: {e}")
            return False
    
    def docx_to_html(self, input_path: Path, output_path: Path, **kwargs) -> bool:
        """Convert Word document to HTML."""
        try:
            # Use pandoc if available (better formatting)
            if self.external_tools['pandoc']:
                return self.convert_with_pandoc(input_path, output_path, 'docx', 'html')
            
            # Fallback to basic conversion
            doc = docx.Document(str(input_path))
            html_content = [f"<h1>{input_path.stem}</h1>"]
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    html_content.append(f"<p>{paragraph.text}</p>")
            
            full_html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <title>{input_path.stem}</title>
                <meta charset="utf-8">
            </head>
            <body>
                {''.join(html_content)}
            </body>
            </html>
            """
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(full_html)
            
            return True
        except Exception as e:
            print(f"Error converting DOCX to HTML: {e}")
            return False
    
    def docx_to_pdf(self, input_path: Path, output_path: Path, **kwargs) -> bool:
        """Convert Word document to PDF."""
        try:
            # Try LibreOffice first (best quality)
            if self.external_tools['libreoffice']:
                cmd = [
                    'libreoffice',
                    '--headless',
                    '--convert-to', 'pdf',
                    '--outdir', str(output_path.parent),
                    str(input_path)
                ]
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
                
                # LibreOffice creates file with same name but .pdf extension
                expected_output = output_path.parent / f"{input_path.stem}.pdf"
                if expected_output.exists():
                    if expected_output != output_path:
                        shutil.move(expected_output, output_path)
                    return True
            
            # Fallback: convert to HTML then PDF
            temp_html = output_path.with_suffix('.html')
            if self.docx_to_html(input_path, temp_html):
                success = self.html_to_pdf(temp_html, output_path)
                temp_html.unlink(missing_ok=True)
                return success
            
            return False
        except Exception as e:
            print(f"Error converting DOCX to PDF: {e}")
            return False
    
    def docx_to_markdown(self, input_path: Path, output_path: Path, **kwargs) -> bool:
        """Convert Word document to Markdown."""
        try:
            if self.external_tools['pandoc']:
                return self.convert_with_pandoc(input_path, output_path, 'docx', 'md')
            
            # Basic fallback
            doc = docx.Document(str(input_path))
            md_content = [f"# {input_path.stem}\n"]
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # Simple formatting detection
                    if (
                        paragraph.style is not None
                        and hasattr(paragraph.style, 'name')
                        and paragraph.style.name is not None
                        and paragraph.style.name.startswith('Heading')
                    ):
                        level = '##' if '1' in paragraph.style.name else '###'
                        md_content.append(f"{level} {paragraph.text}\n")
                    else:
                        md_content.append(f"{paragraph.text}\n")
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(md_content))
            
            return True
        except Exception as e:
            print(f"Error converting DOCX to Markdown: {e}")
            return False
    
    # Markdown conversion methods
    def markdown_to_html(self, input_path: Path, output_path: Path, **kwargs) -> bool:
        """Convert Markdown to HTML."""
        try:
            with open(input_path, 'r', encoding='utf-8') as f:
                md_content = f.read()
            
            # Convert markdown to HTML
            html_body = markdown.markdown(md_content, extensions=['tables', 'fenced_code', 'toc'])
            
            # Add CSS if provided
            css_content = ""
            if 'css' in kwargs and kwargs['css']:
                css_path = Path(kwargs['css'])
                
                # If relative path, try both current working directory and script directory
                if not css_path.is_absolute():
                    # First try current working directory (where user ran the command)
                    cwd_css_path = Path.cwd() / css_path
                    if cwd_css_path.exists():
                        css_path = cwd_css_path
                    # Then try script directory (for bundled CSS files)
                    elif css_path.exists():
                        pass  # Use relative to script directory
                
                if css_path.exists():
                    with open(css_path, 'r', encoding='utf-8') as f:
                        css_content = f"<style>\n{f.read()}\n</style>"
                else:
                    css_content = f'<link rel="stylesheet" href="{kwargs["css"]}">'
            
            full_html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <title>{input_path.stem}</title>
                <meta charset="utf-8">
                {css_content}
            </head>
            <body>
                {html_body}
            </body>
            </html>
            """
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(full_html)
            
            return True
        except Exception as e:
            print(f"Error converting Markdown to HTML: {e}")
            return False
    
    def markdown_to_pdf(self, input_path: Path, output_path: Path, **kwargs) -> bool:
        """Convert Markdown to PDF."""
        try:
            # Convert to HTML first
            temp_html = output_path.with_suffix('.html')
            if self.markdown_to_html(input_path, temp_html, **kwargs):
                success = self.html_to_pdf(temp_html, output_path)
                temp_html.unlink(missing_ok=True)
                return success
            return False
        except Exception as e:
            print(f"Error converting Markdown to PDF: {e}")
            return False
    
    def markdown_to_docx(self, input_path: Path, output_path: Path, **kwargs) -> bool:
        """Convert Markdown to Word document."""
        try:
            if self.external_tools['pandoc']:
                return self.convert_with_pandoc(input_path, output_path, 'md', 'docx')
            
            # Basic fallback - create simple DOCX
            doc = docx.Document()
            
            with open(input_path, 'r', encoding='utf-8') as f:
                lines = f.readlines()
            
            for line in lines:
                line = line.strip()
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line:
                    doc.add_paragraph(line)
            
            doc.save(str(output_path))
            return True
        except Exception as e:
            print(f"Error converting Markdown to DOCX: {e}")
            return False
    
    # HTML conversion methods
    def html_to_pdf(self, input_path: Path, output_path: Path, **kwargs) -> bool:
        """Convert HTML to PDF."""
        try:
            if self.external_tools['wkhtmltopdf']:
                options = {
                    'page-size': 'A4',
                    'margin-top': '0.75in',
                    'margin-right': '0.75in',
                    'margin-bottom': '0.75in',
                    'margin-left': '0.75in',
                    'encoding': "UTF-8",
                    'no-outline': None,
                    'enable-local-file-access': None,
                    'disable-smart-shrinking': '',
                    'print-media-type': '',
                    'load-error-handling': 'ignore',
                    'load-media-error-handling': 'ignore'
                }
                
                # Try to convert with robust error handling
                try:
                    pdfkit.from_file(str(input_path), str(output_path), options=options)
                    return True
                except Exception as e:
                    # If the conversion fails, try with minimal options
                    print(f"Initial conversion failed, trying with minimal options: {e}")
                    minimal_options = {
                        'page-size': 'A4',
                        'enable-local-file-access': None,
                        'load-error-handling': 'ignore',
                        'disable-external-links': '',
                        'disable-internal-links': ''
                    }
                    try:
                        pdfkit.from_file(str(input_path), str(output_path), options=minimal_options)
                        return True
                    except Exception as e2:
                        print(f"Minimal conversion also failed: {e2}")
                        return False
            else:
                print("wkhtmltopdf not found. Install with: sudo apt-get install wkhtmltopdf")
                return False
        except Exception as e:
            print(f"Error converting HTML to PDF: {e}")
            return False
    
    def html_to_text(self, input_path: Path, output_path: Path, **kwargs) -> bool:
        """Convert HTML to plain text."""
        try:
            with open(input_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            # Convert HTML to text
            h = html2text.HTML2Text()
            h.ignore_links = kwargs.get('ignore_links', False)
            text = h.handle(html_content)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)
            
            return True
        except Exception as e:
            print(f"Error converting HTML to text: {e}")
            return False
    
    def html_to_markdown(self, input_path: Path, output_path: Path, **kwargs) -> bool:
        """Convert HTML to Markdown."""
        try:
            with open(input_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            # Convert HTML to Markdown
            h = html2text.HTML2Text()
            h.body_width = 0  # Don't wrap lines
            markdown_content = h.handle(html_content)
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
            
            return True
        except Exception as e:
            print(f"Error converting HTML to Markdown: {e}")
            return False
    
    # Text conversion methods
    def text_to_pdf(self, input_path: Path, output_path: Path, **kwargs) -> bool:
        """Convert plain text to PDF."""
        try:
            with open(input_path, 'r', encoding='utf-8') as f:
                text_content = f.read()
            
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font('Arial', size=12)
            
            # Split text into lines and add to PDF
            lines = text_content.split('\n')
            for line in lines:
                if line.strip():
                    # Handle long lines
                    if len(line) > 80:
                        words = line.split(' ')
                        current_line = ""
                        for word in words:
                            if len(current_line + word) < 80:
                                current_line += word + " "
                            else:
                                pdf.cell(0, 10, current_line.strip(), ln=True)
                                current_line = word + " "
                        if current_line:
                            pdf.cell(0, 10, current_line.strip(), ln=True)
                    else:
                        pdf.cell(0, 10, line, ln=True)
                else:
                    pdf.cell(0, 10, "", ln=True)  # Empty line
            
            pdf.output(str(output_path))
            return True
        except Exception as e:
            print(f"Error converting text to PDF: {e}")
            return False
    
    def text_to_html(self, input_path: Path, output_path: Path, **kwargs) -> bool:
        """Convert plain text to HTML with proper paragraph formatting."""
        try:
            with open(input_path, 'r', encoding='utf-8') as f:
                text_content = f.read()
            
            # Get CSS file if provided
            css_file = kwargs.get('css')
            css_content = ""
            if css_file and Path(css_file).exists():
                with open(css_file, 'r', encoding='utf-8') as f:
                    css_content = f"<style>\n{f.read()}\n</style>"
            else:
                # Default styling for better readability
                css_content = """
                <style>
                    body { 
                        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; 
                        line-height: 1.6; 
                        max-width: 800px; 
                        margin: 40px auto; 
                        padding: 20px;
                        color: #333;
                        background: #fafafa;
                    }
                    h1, h2, h3 { color: #2c3e50; }
                    p { margin: 1em 0; }
                    .company-entry { 
                        margin: 1.5em 0; 
                        padding: 1em; 
                        background: white; 
                        border-radius: 5px;
                        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
                    }
                    .company-name { 
                        font-weight: bold; 
                        color: #2c3e50; 
                        font-size: 1.1em;
                    }
                    .contact-info { 
                        color: #666; 
                        font-size: 0.9em; 
                    }
                </style>"""
            
            # Smart paragraph detection and formatting
            html_content = self._format_text_content(text_content)
            
            full_html = f"""<!DOCTYPE html>
<html>
<head>
    <title>{input_path.stem}</title>
    <meta charset="utf-8">{css_content}
</head>
<body>
{html_content}
</body>
</html>"""
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(full_html)
            
            return True
        except Exception as e:
            print(f"Error converting text to HTML: {e}")
            return False
    
    def _format_text_content(self, text: str) -> str:
        """Format plain text content into proper HTML with paragraphs and structure."""
        lines = text.strip().split('\n')
        html_parts = []
        current_paragraph = []
        in_company_entry = False
        company_lines = []
        
        for line in lines:
            line = line.strip()
            
            # Skip empty lines
            if not line:
                if in_company_entry:
                    # Empty line might end a company entry
                    if len(company_lines) >= 2:  # Minimum lines for a company entry
                        html_parts.append(self._format_company_entry(company_lines))
                        company_lines = []
                        in_company_entry = False
                    continue
                elif current_paragraph:
                    # End current paragraph
                    para_text = ' '.join(current_paragraph)
                    html_parts.append(self._format_paragraph(para_text))
                    current_paragraph = []
                continue
            
            # Check if this looks like a company name/header
            if self._looks_like_company_header(line):
                # End any current paragraph
                if current_paragraph:
                    para_text = ' '.join(current_paragraph)
                    html_parts.append(self._format_paragraph(para_text))
                    current_paragraph = []
                
                # Start or continue company entry
                company_lines.append(line)
                in_company_entry = True
            elif in_company_entry:
                # Continue collecting company-related lines
                company_lines.append(line)
                
                # Check if this line looks like it ends company info (like a phone number)
                if self._looks_like_contact_info(line):
                    # This might be the end of company entry, but continue collecting
                    pass
            else:
                # Regular content - add to current paragraph
                current_paragraph.append(line)
        
        # Handle any remaining content
        if company_lines and in_company_entry:
            html_parts.append(self._format_company_entry(company_lines))
        elif current_paragraph:
            para_text = ' '.join(current_paragraph)
            html_parts.append(self._format_paragraph(para_text))
        
        return '\n'.join(html_parts)
    
    def _format_company_entry(self, company_lines: List[str]) -> str:
        """Format multiple lines as a single company entry."""
        if not company_lines:
            return ""
        
        # First line is typically the company name
        company_name = company_lines[0]
        
        # Group remaining lines into contact info
        contact_lines = company_lines[1:] if len(company_lines) > 1 else []
        
        html_parts = ['<div class="company-entry">']
        html_parts.append(f'<div class="company-name">{company_name}</div>')
        
        # Add contact information as individual paragraphs within the company entry
        for contact_line in contact_lines:
            if self._looks_like_contact_info(contact_line):
                html_parts.append(f'<div class="contact-info">{contact_line}</div>')
            else:
                html_parts.append(f'<p>{contact_line}</p>')
        
        html_parts.append('</div>')
        return '\n'.join(html_parts)
    
    def _looks_like_contact_info(self, line: str) -> bool:
        """Check if a line looks like contact information (phone, address, etc.)."""
        # Phone patterns
        if any(pattern in line for pattern in ['Phone:', 'Tel:', 'Call:', '(']):
            return True
        
        # Address patterns (city, state, zip)
        import re
        if re.search(r'[A-Za-z\s]+,\s*[A-Z]{2}\s*\d{5}', line):
            return True
        
        # Email patterns
        if '@' in line and '.' in line:
            return True
            
        return False
    
    def _looks_like_company_header(self, line: str) -> bool:
        """Check if a line looks like a company name/header."""
        # Look for company keywords
        company_words = ['Inc', 'LLC', 'Corp', 'Company', 'Co.', 'Ltd', 'Industries', 'Manufacturing', 'Services', 'Group', 'Associates', 'Solutions', 'Systems', 'Technologies']
        has_company_word = any(word in line for word in company_words)
        
        # Look for patterns that suggest it's a business name (not contact info)
        is_not_contact = not any(pattern in line.lower() for pattern in ['phone:', 'tel:', 'email:', 'fax:', 'www.', 'http'])
        
        # If it has company words and isn't contact info, likely a company name
        if has_company_word and is_not_contact:
            return True
            
        # Alternative: look for patterns like "Business Name" followed by location info
        # This is less reliable but catches some cases
        if len(line) > 20 and not line.startswith(('Phone', 'Tel', 'Email', 'Address')):
            # Check if it might be a company name with embedded contact info
            import re
            # Look for business name patterns (Title Case words)
            title_case_pattern = re.search(r'^[A-Z][a-z]+(\s+[A-Z][a-z]*)*', line)
            if title_case_pattern and len(title_case_pattern.group()) > 10:
                return True
        
        return False
    
    def _split_company_info(self, line: str) -> dict:
        """Split a company header line into name and contact info."""
        # Try to find where the location starts (usually after company name)
        import re
        
        # Look for city, state pattern
        location_match = re.search(r'([A-Za-z\s]+,\s*[A-Z]{2}[\d\-\s]*)', line)
        if location_match:
            location_start = location_match.start()
            name = line[:location_start].strip()
            location = line[location_start:].strip()
            return {"name": name, "location": location}
        else:
            return {"name": line, "location": ""}
    
    def _format_paragraph(self, text: str) -> str:
        """Format a paragraph of text as HTML."""
        if not text.strip():
            return ""
        return f'<p>{text}</p>'
    
    def text_to_markdown(self, input_path: Path, output_path: Path, **kwargs) -> bool:
        """Convert plain text to Markdown."""
        try:
            with open(input_path, 'r', encoding='utf-8') as f:
                text_content = f.read()
            
            # Basic conversion - wrap in code block to preserve formatting
            markdown_content = f"# {input_path.stem}\n\n```\n{text_content}\n```\n"
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
            
            return True
        except Exception as e:
            print(f"Error converting text to Markdown: {e}")
            return False
    
    # External tool methods
    def convert_with_pandoc(self, input_path: Path, output_path: Path, 
                           from_format: str, to_format: str) -> bool:
        """Convert using Pandoc (if available)."""
        try:
            cmd = [
                'pandoc',
                str(input_path),
                '-f', from_format,
                '-t', to_format,
                '-o', str(output_path)
            ]
            
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
            
            if result.returncode == 0:
                return True
            else:
                print(f"Pandoc error: {result.stderr}")
                return False
                
        except Exception as e:
            print(f"Error using Pandoc: {e}")
            return False
    
    # Main conversion method
    def convert(self, input_path: Path, output_path: Path, 
                from_format: Optional[str] = None, to_format: Optional[str] = None, **kwargs) -> bool:
        """
        Convert a document from one format to another.
        
        Args:
            input_path: Path to input file
            output_path: Path to output file
            from_format: Source format (auto-detected if None)
            to_format: Target format (auto-detected if None)
            **kwargs: Additional options for specific converters
            
        Returns:
            True if conversion successful, False otherwise
        """
        # Auto-detect formats if not provided
        if from_format is None:
            from_format = self.normalize_format(input_path.suffix)
        else:
            from_format = self.normalize_format(from_format)
        
        if to_format is None:
            to_format = self.normalize_format(output_path.suffix)
        else:
            to_format = self.normalize_format(to_format)
        
        # Check if input file exists
        if not input_path.exists():
            print(f"Error: Input file '{input_path}' not found.")
            return False
        
        # Check if conversion is supported
        conversion_key = (from_format, to_format)
        
        if conversion_key in self.converters:
            print(f"Converting {from_format.upper()} to {to_format.upper()}...")
            return self.converters[conversion_key](input_path, output_path, **kwargs)
        
        # Try Pandoc for unsupported conversions
        elif self.external_tools['pandoc']:
            print(f"Converting {from_format.upper()} to {to_format.upper()} using Pandoc...")
            return self.convert_with_pandoc(input_path, output_path, from_format, to_format)
        
        else:
            print(f"Conversion from {from_format.upper()} to {to_format.upper()} not supported.")
            print("Supported conversions:")
            for from_fmt, to_fmt in self.get_supported_conversions():
                print(f"  {from_fmt} → {to_fmt}")
            return False
    
    def batch_convert(self, input_folder: Path, output_folder: Path,
                     from_format: str, to_format: str, **kwargs) -> List[Path]:
        """
        Convert all files of a specific format in a folder.
        
        Args:
            input_folder: Folder containing files to convert
            output_folder: Folder to save converted files
            from_format: Source format
            to_format: Target format
            **kwargs: Additional options for converters
            
        Returns:
            List of successfully converted files
        """
        if not input_folder.exists() or not input_folder.is_dir():
            print(f"Error: Input folder '{input_folder}' not found.")
            return []
        
        # Create output folder
        output_folder.mkdir(parents=True, exist_ok=True)
        
        # Find files to convert
        from_format = self.normalize_format(from_format)
        pattern = f"*.{from_format}"
        files_to_convert = list(input_folder.glob(pattern))
        
        if not files_to_convert:
            print(f"No {from_format.upper()} files found in '{input_folder}'")
            return []
        
        print(f"Found {len(files_to_convert)} {from_format.upper()} files to convert")
        
        converted_files = []
        for file_path in files_to_convert:
            # Create output filename
            output_file = output_folder / f"{file_path.stem}.{to_format}"
            
            print(f"Converting: {file_path.name} → {output_file.name}")
            
            # Convert file
            if self.convert(file_path, output_file, from_format, to_format, **kwargs):
                converted_files.append(output_file)
                print(f"✓ Converted: {output_file.name}")
            else:
                print(f"✗ Failed: {file_path.name}")
        
        return converted_files


def main():
    import sys
    parser = argparse.ArgumentParser(
        description="Convert between document formats (PDF, Word, HTML, Markdown, Text)",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  shift document.docx --to pdf
  shift report.md --to html --css style.css
  shift file.pdf --to text --output extracted.txt
  shift folder/ --batch --from docx --to pdf --output converted/
  shift presentation.html --to pdf
  
Supported formats: pdf, docx, html, md, text
Additional formats available with pandoc: rtf, odt, epub, and more
        """
    )
    
    parser.add_argument('input', help='Input file or folder (with --batch)')
    parser.add_argument('--to', '-t', required=True, help='Target format')
    parser.add_argument('--from', '-f', help='Source format (auto-detected if not specified)')
    parser.add_argument('--output', '-o', help='Output file or folder')
    parser.add_argument('--batch', '-b', action='store_true',
                       help='Convert all files of source format in input folder')
    parser.add_argument('--css', help='CSS file for HTML output')
    parser.add_argument('--ignore-links', action='store_true',
                       help='Ignore links when converting HTML to text')
    
    args = parser.parse_args()
    
    # Initialize converter
    converter = DocumentConverter()
    
    input_path = Path(args.input).resolve()
    to_format = args.to
    from_format = getattr(args, 'from', None)  # 'from' is a keyword
    
    # Batch processing
    if args.batch:
        if not from_format:
            print("Error: --from format required for batch processing")
            return
        
        if not args.output:
            output_folder = input_path.parent / f"{input_path.name}_converted"
        else:
            output_folder = Path(args.output).resolve()
        
        kwargs = {}
        if args.css:
            kwargs['css'] = args.css
        if args.ignore_links:
            kwargs['ignore_links'] = True
        
        converted_files = converter.batch_convert(
            input_path, output_folder, from_format, to_format, **kwargs
        )
        
        if converted_files:
            print(f"\n✓ Converted {len(converted_files)} files")
            print(f"Output folder: {output_folder}")
        else:
            print("No files were converted successfully")
        return
    
    # Single file processing
    if not input_path.exists():
        print(f"Error: Input file '{input_path}' not found.")
        return
    
    # Set default output path
    if not args.output:
        output_path = input_path.parent / f"{input_path.stem}.{to_format}"
    else:
        output_path = Path(args.output).resolve()
    
    # Prepare conversion options
    kwargs = {}
    if args.css:
        kwargs['css'] = args.css
    if args.ignore_links:
        kwargs['ignore_links'] = True
    
    # Convert file
    success = converter.convert(input_path, output_path, from_format, to_format, **kwargs)
    
    if success:
        print(f"✓ Converted to: {output_path}")
    else:
        print("✗ Conversion failed")
        sys.exit(1)


if __name__ == "__main__":
    main()
