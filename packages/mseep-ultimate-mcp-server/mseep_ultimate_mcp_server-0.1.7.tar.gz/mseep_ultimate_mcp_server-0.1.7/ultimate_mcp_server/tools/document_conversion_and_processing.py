# ultimate_mcp_server/tools/document_conversion_and_processing.py
"""Standalone Document Processing Toolkit functions for MCP Server.

A comprehensive, fault-tolerant toolkit for document processing, providing:
(Functionality remains the same as the original class docstring)
"""

###############################################################################
# Imports                                                                     #
###############################################################################
# Standard library imports
import asyncio
import base64
import csv
import functools
import hashlib
import html
import io
import json
import math
import os
import re
import tempfile
import textwrap
import time
from contextlib import contextmanager
from io import StringIO
from pathlib import Path
from typing import (
    TYPE_CHECKING,
    Any,
    Awaitable,
    Callable,
    Dict,
    List,
    Optional,
    Pattern,
    Sequence,
    Set,
    Tuple,
    Union,
)

# Third-party imports
import html2text
from bs4 import BeautifulSoup, Tag
from rapidfuzz import fuzz

# Local application imports
from ultimate_mcp_server.constants import Provider
from ultimate_mcp_server.exceptions import ProviderError, ToolError, ToolInputError
from ultimate_mcp_server.tools.base import (
    with_error_handling,
    with_retry,
    with_tool_metrics,
)
from ultimate_mcp_server.tools.completion import generate_completion
from ultimate_mcp_server.utils import get_logger

# Type checking imports
if TYPE_CHECKING:
    import numpy as np
    import pandas as pd
    import tiktoken
    from docling.datamodel.pipeline_options import AcceleratorDevice as _AcceleratorDeviceType
    from docling_core.types.doc import DoclingDocument as _DoclingDocumentType
    from docling_core.types.doc import ImageRefMode as _ImageRefModeType
    from PIL import Image as PILImage
    from tiktoken import Encoding

# ───────────────────── Optional Dependency Check & Initialization ───────────────────
_DOCLING_AVAILABLE = False
try:
    from docling.datamodel.base_models import InputFormat
    from docling.datamodel.pipeline_options import (
        AcceleratorDevice,
        AcceleratorOptions,
        PdfPipelineOptions,
    )
    from docling.document_converter import DocumentConverter, PdfFormatOption
    from docling_core.types.doc import DoclingDocument, ImageRefMode

    _DOCLING_AVAILABLE = True
    _DoclingDocumentType = DoclingDocument
    _ImageRefModeType = ImageRefMode
    _AcceleratorDeviceType = AcceleratorDevice

except ImportError:
    _DoclingDocumentType = Any
    _ImageRefModeType = Any
    _AcceleratorDeviceType = Any
    InputFormat = None
    AcceleratorDevice = None
    AcceleratorOptions = None
    PdfPipelineOptions = None
    DocumentConverter = None
    PdfFormatOption = None
    pass

_PANDAS_AVAILABLE = False
try:
    import pandas as pd

    _PANDAS_AVAILABLE = True
except ModuleNotFoundError:
    pd = None

_TIKTOKEN_AVAILABLE = False
try:
    import tiktoken

    _TIKTOKEN_AVAILABLE = True
except ModuleNotFoundError:
    tiktoken = None

_PYPDF2_AVAILABLE = False
try:
    import PyPDF2

    _PYPDF2_AVAILABLE = True
except ImportError:
    PyPDF2 = None

_DOCX_AVAILABLE = False
try:
    import docx

    _DOCX_AVAILABLE = True
except ImportError:
    docx = None

_NUMPY_AVAILABLE = False
try:
    import numpy as np

    _NUMPY_AVAILABLE = True
except ImportError:
    np = None

_PIL_AVAILABLE = False
try:
    from PIL import Image, ImageEnhance, ImageFilter

    _PIL_AVAILABLE = True
except ImportError:
    Image, ImageEnhance, ImageFilter = None, None, None

_CV2_AVAILABLE = False
try:
    import cv2

    _CV2_AVAILABLE = True
except ImportError:
    cv2 = None

_PYTESSERACT_AVAILABLE = False
try:
    import pytesseract

    _PYTESSERACT_AVAILABLE = True
except ImportError:
    pytesseract = None

_PDF2IMAGE_AVAILABLE = False
try:
    from pdf2image import convert_from_bytes, convert_from_path

    _PDF2IMAGE_AVAILABLE = True
except ImportError:
    convert_from_bytes, convert_from_path = None, None

_PDFPLUMBER_AVAILABLE = False
try:
    import pdfplumber

    _PDFPLUMBER_AVAILABLE = True
except ImportError:
    pdfplumber = None

_PYMUPDF_AVAILABLE = False
try:
    import pymupdf  # PyMuPDF

    _PYMUPDF_AVAILABLE = True
except ImportError:
    pymupdf = None

_TRAFILATURA_AVAILABLE = False
try:
    import trafilatura

    _TRAFILATURA_AVAILABLE = True
except ImportError:
    trafilatura = None

_READABILITY_AVAILABLE = False
try:
    import readability

    _READABILITY_AVAILABLE = True
except ImportError:
    readability = None

try:
    from markdownify import markdownify as _markdownify_fallback
except ModuleNotFoundError:
    _markdownify_fallback = None

# ───────────────────── Module Level Logger ─────────────────────────
logger = get_logger("ultimate_mcp_server.tools.document_processing")

# ───────────────────── Module Level Config & State ──────────────────
DEFAULT_EXTRACTION_STRATEGY = "hybrid_direct_ocr"
_VALID_FORMATS = {"markdown", "text", "html", "json", "doctags"}
_OCR_COMPATIBLE_FORMATS = {"text", "markdown"}
_VALID_EXTRACTION_STRATEGIES = {
    "docling",
    "direct_text",
    "ocr",
    "hybrid_direct_ocr",
}

# Acceleration Device Mapping (Docling)
if _DOCLING_AVAILABLE and AcceleratorDevice:
    _ACCEL_MAP = {
        "auto": AcceleratorDevice.AUTO,
        "cpu": AcceleratorDevice.CPU,
        "cuda": AcceleratorDevice.CUDA,
        "mps": AcceleratorDevice.MPS,
    }
else:
    _ACCEL_MAP = {"auto": "auto", "cpu": "cpu", "cuda": "cuda", "mps": "mps"}

# HTML Detection Patterns
_RE_FLAGS = re.MULTILINE | re.IGNORECASE
_HTML_PATTERNS: Sequence[Pattern] = [
    re.compile(p, _RE_FLAGS)
    for p in (
        r"<\s*[a-zA-Z]+[^>]*>",
        r"<\s*/\s*[a-zA-Z]+\s*>",
        r"&[a-zA-Z]+;",
        r"&#[0-9]+;",
        r"<!\s*DOCTYPE",
        r"<!\s*--",
    )
]

# Content Type Patterns (Used by detect_content_type)
_CONTENT_PATTERNS: Dict[str, List[Tuple[Pattern, float]]] = {
    "html": [
        (re.compile(r"<html", re.I), 5.0),
        (re.compile(r"<head", re.I), 4.0),
        (re.compile(r"<body", re.I), 4.0),
        (re.compile(r"</(div|p|span|a|li)>", re.I), 1.0),
        (re.compile(r"<[a-z][a-z0-9]*\s+[^>]*>", re.I), 0.8),
        (re.compile(r"<!DOCTYPE", re.I), 5.0),
        (re.compile(r"&\w+;"), 0.5),
    ],
    "markdown": [
        (re.compile(r"^#{1,6}\s+", re.M), 4.0),
        (re.compile(r"^\s*[-*+]\s+", re.M), 2.0),
        (re.compile(r"^\s*\d+\.\s+", re.M), 2.0),
        (re.compile(r"`[^`]+`"), 1.5),
        (re.compile(r"^```", re.M), 5.0),
        (re.compile(r"\*{1,2}[^*\s]+?\*{1,2}"), 1.0),
        (re.compile(r"!\[.*?\]\(.*?\)", re.M), 3.0),
        (re.compile(r"\[.*?\]\(.*?\)", re.M), 2.5),
        (re.compile(r"^>.*", re.M), 2.0),
        (re.compile(r"^-{3,}$", re.M), 3.0),
    ],
    "code": [
        (re.compile(r"def\s+\w+\(.*\):"), 3.0),
        (re.compile(r"class\s+\w+"), 3.0),
        (re.compile(r"import\s+|from\s+"), 3.0),
        (
            re.compile(r"((function\s+\w+\(|const|let|var)\s*.*?=>|\b(document|console|window)\.)"),
            3.0,
        ),
        (re.compile(r"public\s+|private\s+|static\s+"), 2.5),
        (re.compile(r"#include"), 3.0),
        (re.compile(r"<\?php"), 4.0),
        (re.compile(r"console\.log"), 2.0),
        (re.compile(r";\s*$"), 1.0),
        (re.compile(r"\b(var|let|const|int|float|string|bool)\b"), 1.5),
        (re.compile(r"//.*$"), 1.0),
        (re.compile(r"/\*.*?\*/", re.S), 1.5),
    ],
}
_LANG_PATTERNS: List[Tuple[Pattern, str]] = [
    (re.compile(r"(def\s+\w+\(.*?\):|import\s+|from\s+\S+\s+import)"), "python"),
    (
        re.compile(r"((function\s+\w+\(|const|let|var)\s*.*?=>|\b(document|console|window)\.)"),
        "javascript",
    ),
    (re.compile(r"<(\w+)(.*?)>.*?</\1>", re.S), "html"),
    (re.compile(r"<\?php"), "php"),
    (re.compile(r"(public|private|protected)\s+(static\s+)?(void|int|String)"), "java"),
    (re.compile(r"#include\s+<"), "c/c++"),
    (re.compile(r"using\s+System;"), "c#"),
    (re.compile(r"(SELECT|INSERT|UPDATE|DELETE)\s+.*FROM", re.I), "sql"),
    (re.compile(r":\s+\w+\s*\{"), "css"),
    (re.compile(r"^[^:]+:\s* # YAML key-value", re.M | re.X), "yaml"),
    (re.compile(r"\$\w+"), "shell/bash"),
]

# Markdown processing regex
_BULLET_RX = re.compile(r"^[•‣▪◦‧﹒∙·] ?", re.MULTILINE)

# Lazy Loading State
_tiktoken_enc_instance: Union["Encoding", bool, None] = None

# OCR Caching (Simple in-memory - can be extended)
_OCR_CACHE: Dict[str, Any] = {}

# Domain Rules and Compiled Regex (Loaded Lazily)
_DOMAIN_RULES_CACHE: Optional[Dict] = None
_ACTIVE_DOMAIN: Optional[str] = None
_BOUND_RX: Optional[re.Pattern] = None
_CUSTOM_SECT_RX: Optional[List[Tuple[re.Pattern, str]]] = None
_METRIC_RX: Optional[List[Tuple[str, re.Pattern]]] = None
_RISK_RX: Optional[Dict[str, re.Pattern]] = None
_DOC_LABELS: Optional[List[str]] = None
_CLASS_PROMPT_PREFIX: Optional[str] = None

###############################################################################
# Utility & Helper Functions (Private Module Level)                           #
###############################################################################


def _log_dependency_warnings():
    """Logs warnings for missing optional dependencies on first use."""
    if not _DOCLING_AVAILABLE:
        logger.warning(
            "Docling library not available. Advanced PDF/Office conversion features disabled."
        )
    if not _PYPDF2_AVAILABLE:
        logger.warning("PyPDF2 not available. Basic PDF fallback conversion disabled.")
    if not _DOCX_AVAILABLE:
        logger.warning("python-docx not available. Basic DOCX fallback conversion disabled.")
    if not _PANDAS_AVAILABLE:
        logger.warning("Pandas not available. Pandas output format for tables disabled.")
    if not _TIKTOKEN_AVAILABLE:
        logger.warning(
            "Tiktoken not available. Token-based chunking will fallback to character chunking."
        )
    ocr_deps = {
        "Pillow": _PIL_AVAILABLE,
        "numpy": _NUMPY_AVAILABLE,
        "opencv-python": _CV2_AVAILABLE,
        "pytesseract": _PYTESSERACT_AVAILABLE,
        "pdf2image": _PDF2IMAGE_AVAILABLE,
    }
    missing_ocr = [name for name, avail in ocr_deps.items() if not avail]
    if missing_ocr:
        logger.warning(
            f"Missing OCR dependencies: {', '.join(missing_ocr)}. OCR functionality limited/disabled."
        )
    if not _PDFPLUMBER_AVAILABLE and not _PYMUPDF_AVAILABLE:
        logger.warning(
            "Missing direct PDF text extraction libraries (pdfplumber/pymupdf). Direct text extraction disabled."
        )
    elif not _PDFPLUMBER_AVAILABLE:
        logger.warning(
            "pdfplumber not available. Will rely solely on PyMuPDF for direct text extraction."
        )
    elif not _PYMUPDF_AVAILABLE:
        logger.warning(
            "PyMuPDF not available. Will rely solely on pdfplumber for direct text extraction."
        )
    if not _TRAFILATURA_AVAILABLE:
        logger.warning("Trafilatura not installed. Trafilatura HTML extraction disabled.")
    if not _READABILITY_AVAILABLE:
        logger.warning("Readability-lxml not installed. Readability HTML extraction disabled.")
    if not _markdownify_fallback:
        logger.warning("Markdownify not installed. HTML to Markdown fallback disabled.")


# Call once on import to log status
_log_dependency_warnings()


def _load_and_compile_domain_rules():
    """Loads domain rules from config and compiles regex patterns."""
    global _DOMAIN_RULES_CACHE, _ACTIVE_DOMAIN, _BOUND_RX, _CUSTOM_SECT_RX
    global _METRIC_RX, _RISK_RX, _DOC_LABELS, _CLASS_PROMPT_PREFIX

    if _DOMAIN_RULES_CACHE is not None:  # Already loaded
        return

    logger.debug("Lazily loading and compiling domain rules...")
    default_rules = {
        "generic": {
            "classification": {
                "labels": ["Report", "Contract", "Presentation", "Memo", "Email", "Manual"],
                "prompt_prefix": "Classify the document into exactly one of: ",
            },
            "sections": {
                "boundary_regex": r"^\s*(chapter\s+\d+|section\s+\d+|[A-Z][A-Za-z\s]{3,80})$",
                "custom": [],
            },
            "metrics": {
                "metric_1": {"aliases": ["metric one", "m1"]},
                "metric_2": {"aliases": ["metric two", "m2"]},
            },
            "risks": {"Risk_A": r"risk a", "Risk_B": r"risk b"},
        },
        "finance": {
            "classification": {
                "labels": [
                    "10-K",
                    "Credit Agreement",
                    "Investor Deck",
                    "Press Release",
                    "Board Minutes",
                    "NDA",
                    "LPA",
                    "CIM",
                ],
                "prompt_prefix": "Identify the document type (finance domain): ",
            },
            "sections": {
                "boundary_regex": r"^\s*(item\s+\d+[a-z]?\.|[A-Z][A-Za-z\s]{3,80})$",
                "custom": [
                    {"regex": r"item\s+1a?\.? .*business", "label": "Business"},
                    {"regex": r"item\s+1a\.? .*risk factors", "label": "Risk Factors"},
                    {"regex": r"item\s+7\.? .*management'?s discussion", "label": "MD&A"},
                    {"regex": r"covena[nv]ts", "label": "Covenants"},
                ],
            },
            "metrics": {
                "revenue": {
                    "aliases": ["revenue", "net sales", "total sales", "sales revenue", "turnover"]
                },
                "ebitda": {
                    "aliases": ["ebitda", "adj. ebitda", "operating profit", "operating income"]
                },
                "gross_profit": {"aliases": ["gross profit"]},
                "net_income": {"aliases": ["net income", "net profit", "earnings"]},
                "capex": {"aliases": ["capital expenditures", "capex"]},
                "debt": {"aliases": ["total debt", "net debt", "long-term debt"]},
            },
            "risks": {
                "Change_of_Control": r"change\s+of\s+control",
                "ESG_Risk": r"(child\s+labor|environmental\s+violation|scope\s+3)",
                "PII": r"(\bSSN\b|social security number|passport no)",
            },
        },
        "legal": {
            "classification": {
                "labels": ["Contract", "NDA", "Lease", "Policy", "License", "Settlement"],
                "prompt_prefix": "Classify the legal document into exactly one of: ",
            },
            "sections": {
                "boundary_regex": r"^\s*(article\s+\d+|section\s+\d+|[A-Z][A-Za-z\s]{3,80})$",
                "custom": [
                    {"regex": r"definitions", "label": "Definitions"},
                    {"regex": r"termination", "label": "Termination"},
                    {"regex": r"confidentiality", "label": "Confidentiality"},
                ],
            },
            "metrics": {},
            "risks": {
                "Indemnity": r"indemnif(y|ication)",
                "Liquidated_Damages": r"liquidated damages",
                "Governing_Law_NY": r"governing law.*new york",
                "Governing_Law_DE": r"governing law.*delaware",
            },
        },
        "medical": {
            "classification": {
                "labels": [
                    "Clinical Study",
                    "Patient Report",
                    "Lab Results",
                    "Prescription",
                    "Care Plan",
                ],
                "prompt_prefix": "Classify the medical document: ",
            },
            "sections": {
                "boundary_regex": r"^\s*(section\s+\d+|[A-Z][A-Za-z\s]{3,80})$",
                "custom": [
                    {"regex": r"diagnosis", "label": "Diagnosis"},
                    {"regex": r"treatment", "label": "Treatment"},
                    {"regex": r"medications", "label": "Medications"},
                    {"regex": r"allergies", "label": "Allergies"},
                ],
            },
            "metrics": {
                "blood_pressure": {"aliases": ["blood pressure", "bp"]},
                "heart_rate": {"aliases": ["heart rate", "hr"]},
                "temperature": {"aliases": ["temperature", "temp"]},
                "bmi": {"aliases": ["bmi", "body mass index"]},
            },
            "risks": {
                "Allergy": r"allergic reaction",
                "Contraindication": r"contraindicat(ed|ion)",
                "Adverse_Event": r"adverse event",
            },
        },
    }
    _DOMAIN_RULES_CACHE = default_rules

    _ACTIVE_DOMAIN = os.getenv("DOC_DOMAIN", "generic")
    # Config loading if needed:
    # from ultimate_mcp_server.config import get_config
    # try:
    #     cfg = get_config()
    #     _ACTIVE_DOMAIN = cfg.document_processing.domain if cfg and hasattr(cfg, 'document_processing') else "generic"
    # except Exception as e:
    #     logger.warning(f"Failed to load document processing domain from config: {e}. Defaulting to 'generic'.")
    #     _ACTIVE_DOMAIN = "generic"

    if _ACTIVE_DOMAIN not in _DOMAIN_RULES_CACHE:
        logger.warning(f"Unknown DOC_DOMAIN '{_ACTIVE_DOMAIN}', defaulting to 'generic'.")
        _ACTIVE_DOMAIN = "generic"

    instruction_json = _DOMAIN_RULES_CACHE[_ACTIVE_DOMAIN]

    try:
        _BOUND_RX = re.compile(instruction_json["sections"].get("boundary_regex", r"$^"), re.M)
    except re.error as e:
        logger.error(f"Invalid boundary regex for domain {_ACTIVE_DOMAIN}: {e}")
        _BOUND_RX = re.compile(r"$^")

    _CUSTOM_SECT_RX = []
    for d in instruction_json["sections"].get("custom", []):
        try:
            _CUSTOM_SECT_RX.append((re.compile(d["regex"], re.I), d["label"]))
        except re.error as e:
            logger.error(
                f"Invalid custom section regex '{d['regex']}' for domain {_ACTIVE_DOMAIN}: {e}"
            )

    _METRIC_RX = []
    for key, cfg in instruction_json.get("metrics", {}).items():
        aliases = cfg.get("aliases", [])
        if aliases:
            try:
                sorted_aliases = sorted(aliases, key=len, reverse=True)
                joined = "|".join(re.escape(a) for a in sorted_aliases)
                if joined:
                    pattern = re.compile(
                        rf"""(?i)\b({joined})\b[\s:–-]*([$€£]?\s?-?\d[\d,.]*)""",
                        re.VERBOSE | re.MULTILINE,
                    )
                    _METRIC_RX.append((key, pattern))
            except re.error as e:
                logger.error(
                    f"Invalid metric regex for alias group '{key}' in domain {_ACTIVE_DOMAIN}: {e}"
                )

    _RISK_RX = {}
    for t, pat_str in instruction_json.get("risks", {}).items():
        try:
            _RISK_RX[t] = re.compile(pat_str, re.I)
        except re.error as e:
            logger.error(
                f"Invalid risk regex for '{t}' in domain {_ACTIVE_DOMAIN}: '{pat_str}'. Error: {e}"
            )

    _DOC_LABELS = instruction_json["classification"].get("labels", [])
    _CLASS_PROMPT_PREFIX = instruction_json["classification"].get("prompt_prefix", "")
    logger.info(f"Domain rules loaded and compiled for domain: '{_ACTIVE_DOMAIN}'")


def _get_active_domain_rules():
    """Ensures domain rules are loaded and returns them."""
    if _DOMAIN_RULES_CACHE is None:
        _load_and_compile_domain_rules()
    return {
        "active_domain": _ACTIVE_DOMAIN,
        "bound_rx": _BOUND_RX,
        "custom_sect_rx": _CUSTOM_SECT_RX,
        "metric_rx": _METRIC_RX,
        "risk_rx": _RISK_RX,
        "doc_labels": _DOC_LABELS,
        "class_prompt_prefix": _CLASS_PROMPT_PREFIX,
    }


def _get_tiktoken_encoder() -> Optional["Encoding"]:
    """Lazy load and return the tiktoken encoder instance."""
    global _tiktoken_enc_instance
    if _tiktoken_enc_instance is not None:
        return (
            _tiktoken_enc_instance
            if isinstance(_tiktoken_enc_instance, tiktoken.Encoding)
            else None
        )
    if not _TIKTOKEN_AVAILABLE:
        _tiktoken_enc_instance = False
        return None
    try:
        encoding_name = os.getenv("TIKTOKEN_ENCODING", "cl100k_base")
        logger.info(f"Lazy-loading tiktoken encoding: {encoding_name}")
        _tiktoken_enc_instance = tiktoken.get_encoding(encoding_name)  # type: ignore
        logger.info("Successfully lazy-loaded tiktoken encoder.")
        return _tiktoken_enc_instance  # type: ignore
    except Exception as e:
        logger.error(f"Failed to lazy-load tiktoken: {e}", exc_info=True)
        _tiktoken_enc_instance = False
        return None


async def _standalone_llm_call(
    *,
    prompt: str,
    provider: str = Provider.OPENAI.value,
    model: str | None = None,
    temperature: float = 0.3,
    max_tokens: int | None = None,
    extra: Dict[str, Any] | None = None,
) -> str:
    """Standalone wrapper to make LLM calls using the completion tool."""
    if not callable(generate_completion):
        logger.error("LLM generation function 'generate_completion' is not available.")
        raise ToolError("LLM_UNAVAILABLE", details={"reason": "generate_completion not available"})

    chosen_provider = provider
    try:
        additional_params = extra or {}
        response_dict = await generate_completion(
            prompt=prompt,
            provider=chosen_provider,
            model=model,
            temperature=temperature,
            max_tokens=max_tokens,
            additional_params=additional_params,
        )
        if isinstance(response_dict, dict):
            if response_dict.get("isError", False) or not response_dict.get("success", True):
                err_detail = response_dict.get("error", {})
                err_msg = err_detail.get("message", "Unknown LLM Error")
                err_code = err_detail.get("type", "LLM_CALL_FAILED")
                logger.error(
                    f"LLM call failed [{err_code}]: {err_msg}. Raw Response: {response_dict}"
                )
                raise ToolError(
                    err_code,
                    details={
                        "provider": chosen_provider,
                        "error": err_msg,
                        "raw_response": str(response_dict),
                    },
                )
            llm_content = response_dict.get("text") or response_dict.get("content")
            if llm_content is None:
                logger.error(f"LLM response missing 'text'/'content': {response_dict}")
                raise ToolError(
                    "LLM_INVALID_RESPONSE",
                    details={"reason": "Missing content", "response_received": str(response_dict)},
                )
            if isinstance(llm_content, str):
                return llm_content.strip()
            else:
                logger.warning(f"LLM content not string: {type(llm_content)}. Converting.")
                return str(llm_content).strip()
        else:
            logger.error(f"LLM response unexpected format: {response_dict}")
            raise ToolError(
                "LLM_INVALID_RESPONSE", details={"response_received": str(response_dict)}
            )
    except ProviderError as pe:
        logger.error(f"LLM provider error ({chosen_provider}): {pe}", exc_info=True)
        raise ToolError(
            "LLM_PROVIDER_ERROR",
            details={"provider": chosen_provider, "error_code": pe.error_code, "error": str(pe)},
        ) from pe
    except ToolError as te:
        raise te
    except Exception as e:
        logger.error(f"LLM call failed ({chosen_provider}): {e}", exc_info=True)
        raise ToolError(
            "LLM_CALL_FAILED", details={"provider": chosen_provider, "error": str(e)}
        ) from e


@contextmanager
def _span(label: str):
    """Context manager for timing operations (module level)."""
    st = time.perf_counter()
    logger.debug(f"Starting span: {label}")
    try:
        yield
    finally:
        elapsed = time.perf_counter() - st
        logger.debug(f"Finished span: {label} ({elapsed:.3f}s)")


def _get_docling_converter(device, threads: int):
    """Create a Docling DocumentConverter."""
    if not _DOCLING_AVAILABLE:
        raise ToolError("DEPENDENCY_MISSING", details={"dependency": "docling"})
    if (
        not PdfPipelineOptions
        or not AcceleratorOptions
        or not InputFormat
        or not PdfFormatOption
        or not DocumentConverter
    ):
        raise ToolError(
            "INTERNAL_ERROR", details={"reason": "Docling partially imported but types missing"}
        )
    opts = PdfPipelineOptions()
    opts.do_ocr = False
    opts.generate_page_images = False
    opts.do_table_extraction = True # Explicitly enable table extraction in the pipeline options
    opts.accelerator_options = AcceleratorOptions(num_threads=threads, device=device)
    try:
        converter_options = {InputFormat.PDF: PdfFormatOption(pipeline_options=opts)}
        return DocumentConverter(format_options=converter_options)
    except Exception as e:
        logger.error(f"Failed to initialize Docling DocumentConverter: {e}", exc_info=True)
        raise ToolError(
            "INITIALIZATION_FAILED", details={"component": "DocumentConverter", "error": str(e)}
        ) from e


def _get_input_path_or_temp(
    document_path: Optional[str], document_data: Optional[bytes]
) -> Tuple[Path, bool]:
    """Gets a valid Path object for input. Saves data to temp file if needed."""
    is_temp = False
    if document_path:
        path = Path(document_path)
        if not path.is_file():
            raise ToolInputError(
                f"Input file not found: {document_path}", param_name="document_path"
            )
        return path, is_temp
    elif document_data:
        try:
            suffix = ".bin"
            if document_data.startswith(b"%PDF"):
                suffix = ".pdf"
            elif len(document_data) > 10 and document_data[6:10] in (b"JFIF", b"Exif"):
                suffix = ".jpg"
            elif document_data.startswith(b"\x89PNG\r\n\x1a\n"):
                suffix = ".png"
            elif document_data.startswith((b"II*\x00", b"MM\x00*")):
                suffix = ".tiff"
            elif document_data.startswith(b"PK\x03\x04"):
                suffix = ".zip"
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
                tmp_file.write(document_data)
                path = Path(tmp_file.name)
            is_temp = True
            logger.debug(f"Saved input data to temporary file: {path}")
            return path, is_temp
        except Exception as e:
            raise ToolError(
                "TEMP_FILE_ERROR",
                details={"error": f"Failed to save input data to temporary file: {e}"},
            ) from e
    else:
        raise ToolInputError("Either 'document_path' or 'document_data' must be provided.")


@contextmanager
def _handle_temp_file(path: Path, is_temp: bool):
    """Context manager to clean up temporary file."""
    try:
        yield path
    finally:
        if is_temp and path.exists():
            try:
                path.unlink()
                logger.debug(f"Cleaned up temporary file: {path}")
            except OSError as e:
                logger.warning(f"Failed to delete temporary file {path}: {e}")


def _tmp_path(src: str, fmt: str) -> Path:
    """Generate a temporary file path for output."""
    src_path = Path(src.split("?")[0])
    stem = src_path.stem or "document"
    ext = "md" if fmt == "markdown" else fmt
    timestamp = int(time.time() * 1000)
    temp_dir = Path(tempfile.gettempdir())
    temp_dir.mkdir(parents=True, exist_ok=True)
    return temp_dir / f"{stem}_{timestamp}.{ext}"


def _get_docling_metadata(doc: Any) -> dict[str, Any]:
    """Extract metadata from a Docling document."""
    if not _DOCLING_AVAILABLE or not doc:
        return {"error": "Docling not available or document object missing"}
    num_pages = 0
    try:
        num_pages = doc.num_pages() if callable(getattr(doc, "num_pages", None)) else 0
        has_tables = False
        has_figures = False
        has_sections = False
        if hasattr(doc, "pages") and isinstance(doc.pages, list):
            for page in doc.pages:
                if hasattr(page, "content") and page.content:
                    if (
                        callable(getattr(page.content, "has_tables", None))
                        and page.content.has_tables()
                    ):
                        has_tables = True
                    if (
                        callable(getattr(page.content, "has_figures", None))
                        and page.content.has_figures()
                    ):
                        has_figures = True
                if has_tables and has_figures:
                    break
        if hasattr(doc, "texts") and isinstance(doc.texts, list):
            for item in doc.texts:
                if hasattr(item, "__class__") and item.__class__.__name__ == "SectionHeaderItem":
                    has_sections = True
                    break
                elif hasattr(item, "label") and getattr(item, "label", None) == "section_header":
                    has_sections = True
                    break
        return {
            "num_pages": num_pages,
            "has_tables": has_tables,
            "has_figures": has_figures,
            "has_sections": has_sections,
        }
    except Exception as e:
        logger.warning(f"Docling metadata collection failed: {e}", exc_info=True)
        return {
            "num_pages": num_pages,
            "has_tables": False,
            "has_figures": False,
            "has_sections": False,
            "metadata_error": str(e),
        }


def _get_basic_metadata(text_content: str, num_pages: int = 0) -> dict[str, Any]:
    """Generate basic metadata for non-Docling content."""
    has_tables = "| --- |" in text_content or "\t" in text_content
    has_figures = "![" in text_content
    has_sections = bool(re.search(r"^#{1,6}\s+", text_content, re.M))
    return {
        "num_pages": num_pages,
        "has_tables": has_tables,
        "has_figures": has_figures,
        "has_sections": has_sections,
    }


def _json(obj: Any) -> str:
    """Utility to serialize objects to JSON."""
    return json.dumps(obj, ensure_ascii=False, separators=(",", ":"))


def _hash(txt: str) -> str:
    """Generate SHA-1 hash of text."""
    return hashlib.sha1(txt.encode("utf-8", "ignore")).hexdigest()


# --- HTML Helpers ---
def _is_html_fragment(text: str) -> bool:
    """Check if text contains likely HTML markup using precompiled patterns."""
    check_len = min(len(text), 5000)
    sample = text[:check_len]
    return any(p.search(sample) for p in _HTML_PATTERNS)


def _best_soup(html_txt: str) -> Tuple[BeautifulSoup, str]:
    """Try progressively more forgiving parsers; fall back to empty soup."""
    parsers = ("html.parser", "lxml", "html5lib")
    last_exception = None
    for p_name in parsers:
        try:
            # Use ModuleNotFoundError for library availability checks
            return BeautifulSoup(html_txt, p_name), p_name
        except ModuleNotFoundError:
            logger.debug(f"HTML parser '{p_name}' not installed, skipping.")
            continue
        except Exception as e_parse:
            last_exception = e_parse
            logger.debug(f"HTML parsing with '{p_name}' failed: {e_parse}")
            continue

    if last_exception:
        logger.warning(
            f"All standard HTML parsers failed ({last_exception}), attempting fragment parsing."
        )
    wrapped_html = (
        f"<!DOCTYPE html><html><head><title>Fragment</title></head><body>{html_txt}</body></html>"
    )
    try:
        return BeautifulSoup(wrapped_html, "html.parser"), "html.parser-fragment"
    except Exception as e_frag:
        logger.error(
            f"Fragment parsing also failed: {e_frag}. Returning empty soup.", exc_info=True
        )
        return BeautifulSoup("", "html.parser"), "failed"


def _clean_html(html_txt: str) -> Tuple[str, str]:
    """Remove dangerous/pointless elements & attempt structural repair."""
    soup, parser_used = _best_soup(html_txt)
    if parser_used == "failed":
        logger.warning("HTML cleaning skipped due to parsing failure.")
        return html_txt, parser_used

    tags_to_remove = [
        "script",
        "style",
        "svg",
        "iframe",
        "canvas",
        "noscript",
        "meta",
        "link",
        "form",
        "input",
        "button",
        "select",
        "textarea",
        "nav",
        "aside",
        "header",
        "footer",
        "video",
        "audio",
    ]
    for el in soup(tags_to_remove):
        el.decompose()

    unsafe_attrs = ["style", "onclick", "onload", "onerror", "onmouseover", "onmouseout", "target"]
    for tag in soup.find_all(True):
        current_attrs = list(tag.attrs.keys())
        for attr in current_attrs:
            attr_val_str = str(tag.get(attr, "")).lower()
            is_unsafe = (
                attr in unsafe_attrs
                or attr.startswith("on")
                or attr.startswith("data-")
                or (attr == "src" and ("javascript:" in attr_val_str or "data:" in attr_val_str))
                or (attr == "href" and attr_val_str.startswith("javascript:"))
            )
            if is_unsafe and attr in tag.attrs:
                del tag[attr]
    try:
        text = str(soup)
        text = html.unescape(text)
        text = re.sub(r"[ \t\r\f\v]+", " ", text)
        text = re.sub(r"\n\s*\n", "\n\n", text)
        text = text.strip()
    except Exception as e:
        logger.error(f"Error during HTML text processing (unescape/regex): {e}", exc_info=True)
        try:
            return str(soup), parser_used
        except Exception as stringify_error:
            logger.error(f"Could not stringify soup after error: {stringify_error}")
            return html_txt, parser_used

    return text, parser_used


# --- Markdown Helpers ---
def _sanitize(md: str) -> str:
    """Basic Markdown sanitization."""
    if not md:
        return ""
    md = md.replace("\u00a0", " ")
    md = _BULLET_RX.sub("- ", md)
    md = re.sub(r"\n{3,}", "\n\n", md)
    md = re.sub(r"[ \t]+$", "", md, flags=re.MULTILINE)
    md = re.sub(r"^[ \t]+", "", md, flags=re.MULTILINE)
    md = re.sub(r"(^|\n)(#{1,6})([^#\s])", r"\1\2 \3", md)
    md = re.sub(r"```\s*\n", "```\n", md)
    md = re.sub(r"\n\s*```", "\n```", md)
    md = re.sub(r"^[*+]\s", "- ", md, flags=re.MULTILINE)
    md = re.sub(r"^\d+\.\s", lambda m: f"{m.group(0).strip()} ", md, flags=re.MULTILINE)
    return md.strip()


def _improve(md: str) -> str:
    """Apply structural improvements to Markdown text."""
    if not md:
        return ""
    # Ensure blank lines around major block elements
    md = re.sub(r"(?<=\S)\n(#{1,6}\s)", r"\n\n\1", md)
    md = re.sub(r"(^#{1,6}\s.*\S)\n(?!\n|#|```|>|\s*[-*+]|\s*\d+\.)", r"\1\n\n", md, flags=re.M)
    md = re.sub(r"(?<=\S)\n(```)", r"\n\n\1", md)
    md = re.sub(r"(```)\n(?!\n)", r"\1\n\n", md)
    md = re.sub(r"(?<=\S)\n(> )", r"\n\n\1", md)
    md = re.sub(r"(\n> .*\S)\n(?!\n|>\s)", r"\1\n\n", md, flags=re.M)
    md = re.sub(r"(?<=\S)\n(\s*([-*+]|\d+\.)\s)", r"\n\n\1", md)
    md = re.sub(
        r"(\n(\s*[-*+]\s+|\s*\d+\.\s+).*\S)\n(?!\n|\s*([-*+]|\d+\.)\s)", r"\1\n\n", md, flags=re.M
    )
    md = re.sub(r"(?<=\S)\n(-{3,}|\*{3,}|_{3,})$", r"\n\n\1", md, flags=re.M)
    md = re.sub(r"(^-{3,}|\*{3,}|_{3,})\n(?!\n)", r"\1\n\n", md, flags=re.M)
    md = re.sub(r"\n{3,}", "\n\n", md)
    return md.strip()


def _convert_html_table_to_markdown(table_tag: Tag) -> str:
    """Converts a single BeautifulSoup table Tag to a Markdown string."""
    md_rows = []
    num_cols = 0
    header_row_tag = table_tag.find("thead")
    header_cells_tags = []
    if header_row_tag:
        header_cells_tags = header_row_tag.find_all(["th", "td"], recursive=False)
        if not header_cells_tags:
            header_row_tr = header_row_tag.find("tr")
            if header_row_tr:
                header_cells_tags = header_row_tr.find_all(["th", "td"])
    if not header_cells_tags:
        first_row = table_tag.find("tr")
        if first_row:
            temp_cells = first_row.find_all(["th", "td"])
            is_header = any(c.name == "th" for c in temp_cells) or (
                len(temp_cells) > 0
                and not any(re.match(r"^\s*[\d.,-]+\s*$", c.get_text()) for c in temp_cells)
            )
            if is_header:
                header_cells_tags = temp_cells

    if header_cells_tags:
        num_cols = len(header_cells_tags)
        hdr = [
            " ".join(c.get_text(" ", strip=True).replace("|", "\\|").split())
            for c in header_cells_tags
        ]
        md_rows.append("| " + " | ".join(hdr) + " |")
        md_rows.append("| " + " | ".join(["---"] * num_cols) + " |")
    else:
        body_rows_tags = (
            table_tag.find("tbody").find_all("tr")
            if table_tag.find("tbody")
            else table_tag.find_all("tr")
        )
        if not body_rows_tags:
            return ""
        for r in body_rows_tags:
            num_cols = max(num_cols, len(r.find_all(["th", "td"])))
        if num_cols == 0:
            return ""
        logger.debug(f"Table has no clear header, assuming {num_cols} columns.")
        md_rows.append("| " + " | ".join([f"Col {i + 1}" for i in range(num_cols)]) + " |")
        md_rows.append("| " + " | ".join(["---"] * num_cols) + " |")

    body_rows_tags = []
    tbody = table_tag.find("tbody")
    if tbody:
        body_rows_tags = tbody.find_all("tr")
    else:
        all_trs = table_tag.find_all("tr")
        start_index = (
            1
            if header_cells_tags
            and all_trs
            and header_cells_tags[0].find_parent("tr") == all_trs[0]
            else 0
        )
        body_rows_tags = all_trs[start_index:]

    for r in body_rows_tags:
        cells = r.find_all(["td", "th"])
        cell_texts = [
            " ".join(c.get_text(" ", strip=True).replace("|", "\\|").split()) for c in cells
        ]
        cell_texts.extend([""] * (num_cols - len(cells)))
        cell_texts = cell_texts[:num_cols]
        md_rows.append("| " + " | ".join(cell_texts) + " |")

    return "\n".join(md_rows)


def _convert_html_tables_to_markdown(html_txt: str) -> str:
    """Finds HTML tables and replaces them with Markdown format within the HTML string."""
    soup, parser_used = _best_soup(html_txt)
    if parser_used == "failed":
        logger.warning("Skipping HTML table conversion due to parsing failure.")
        return html_txt
    tables = soup.find_all("table")
    if not tables:
        return html_txt
    logger.debug(f"Found {len(tables)} HTML tables to convert to Markdown.")
    for table_tag in tables:
        try:
            md_table_str = _convert_html_table_to_markdown(table_tag)
            if md_table_str:
                placeholder = soup.new_string(f"\n\n{md_table_str}\n\n")
                table_tag.replace_with(placeholder)
            else:
                table_tag.decompose()
        except Exception as e:
            logger.error(f"Failed to convert a table to Markdown: {e}", exc_info=True)
    return str(soup)


def _html_to_md_core(html_txt: str, links: bool, imgs: bool, tbls: bool, width: int) -> str:
    """Convert HTML to Markdown using primary and fallback libraries."""
    try:
        h = html2text.HTML2Text()
        h.ignore_links = not links
        h.ignore_images = not imgs
        processed_html = html_txt
        if tbls:
            processed_html = _convert_html_tables_to_markdown(html_txt)
            h.ignore_tables = True
        else:
            h.ignore_tables = True

        h.body_width = width if width > 0 else 0
        h.unicode_snob = True
        h.escape_snob = True
        h.skip_internal_links = True
        h.single_line_break = True

        md_text = h.handle(processed_html)
        logger.debug("html2text conversion successful.")
        return md_text.strip()
    except Exception as e_html2text:
        logger.warning(f"html2text failed ({e_html2text}); attempting fallback with markdownify")
        if _markdownify_fallback and callable(_markdownify_fallback):
            try:
                md_opts = {
                    "strip": [
                        "script",
                        "style",
                        "meta",
                        "link",
                        "head",
                        "iframe",
                        "form",
                        "button",
                        "input",
                        "select",
                        "textarea",
                        "nav",
                        "aside",
                        "header",
                        "footer",
                        "svg",
                        "canvas",
                        "video",
                        "audio",
                    ],
                    "convert": [
                        "a",
                        "p",
                        "img",
                        "br",
                        "hr",
                        "h1",
                        "h2",
                        "h3",
                        "h4",
                        "h5",
                        "h6",
                        "li",
                        "ul",
                        "ol",
                        "blockquote",
                        "code",
                        "pre",
                        "strong",
                        "em",
                        "b",
                        "i",
                        "table",
                        "tr",
                        "td",
                        "th",
                    ],
                    "heading_style": "ATX",
                    "bullets": "-",
                    "strong_em_symbol": "*",
                    "autolinks": False,
                }
                if not links:
                    md_opts["convert"] = [tag for tag in md_opts["convert"] if tag != "a"]
                if not imgs:
                    md_opts["convert"] = [tag for tag in md_opts["convert"] if tag != "img"]
                if not tbls:
                    md_opts["convert"] = [
                        tag for tag in md_opts["convert"] if tag not in ["table", "tr", "td", "th"]
                    ]
                md_text = _markdownify_fallback(html_txt, **md_opts)
                logger.debug("Markdownify fallback conversion successful.")
                return md_text.strip()
            except Exception as e_markdownify:
                logger.error(f"Markdownify fallback also failed: {e_markdownify}", exc_info=True)
                raise ToolError(
                    "MARKDOWN_CONVERSION_FAILED",
                    details={
                        "reason": "Both failed",
                        "html2text_error": str(e_html2text),
                        "markdownify_error": str(e_markdownify),
                    },
                ) from e_markdownify
        else:
            logger.error("html2text failed and markdownify fallback is not available.")
            raise ToolError(
                "MARKDOWN_CONVERSION_FAILED",
                details={"reason": "html2text failed, no fallback", "error": str(e_html2text)},
            ) from e_html2text


###############################################################################
# Core OCR & PDF Helper Functions (Standalone)                                #
###############################################################################


def _ocr_check_dep(dep_name: str, is_available: bool, feature: str):
    """Checks if a required dependency is available, raising ToolError if not."""
    if not is_available:
        logger.error(f"Missing required dependency '{dep_name}' for feature '{feature}'.")
        raise ToolError("DEPENDENCY_MISSING", details={"dependency": dep_name, "feature": feature})


def _ocr_extract_text_from_pdf_direct(
    file_path: Path, start_page: int = 0, max_pages: int = 0
) -> Tuple[List[str], bool]:
    """
    Extracts text directly from PDF using PyMuPDF or PDFPlumber (sync function).

    Args:
        file_path: Path to the PDF file.
        start_page: 0-based starting page index.
        max_pages: Maximum number of pages to extract (0 for all from start_page).

    Returns:
        Tuple containing:
        - List of strings, one per extracted page (or error marker).
        - Boolean indicating if meaningful text was found on at least one page.
    """
    texts: List[str] = []
    has_text = False
    min_chars = 50  # Threshold for considering a page to have meaningful text

    # --- Try PyMuPDF First ---
    if _PYMUPDF_AVAILABLE and pymupdf:
        logger.debug(f"Attempting direct text extraction with PyMuPDF for {file_path}")
        try:
            with pymupdf.open(file_path) as doc:  # type: ignore
                total_pages = len(doc)
                # Calculate 0-based end page index (exclusive)
                end_page = (
                    total_pages if max_pages <= 0 else min(start_page + max_pages, total_pages)
                )
                # Ensure start_page is valid
                start_page = min(start_page, total_pages)
                end_page = max(start_page, end_page)  # Ensure end is not before start

                for i in range(start_page, end_page):
                    try:
                        page = doc.load_page(i)  # Use load_page for clarity
                        page_text = page.get_text("text") or ""  # Specify text format
                        texts.append(page_text)
                        if len(page_text.strip()) >= min_chars:
                            has_text = True
                    except Exception as e_page:
                        logger.warning(
                            f"PyMuPDF: Error extracting text from page {i + 1}: {e_page}"
                        )
                        texts.append(f"[Page {i + 1} Extraction Error: PyMuPDF]")
                logger.debug(
                    f"PyMuPDF extracted {len(texts)} pages. Found meaningful text: {has_text}"
                )
                return texts, has_text
        except Exception as e_pymupdf:
            logger.warning(
                f"PyMuPDF direct text extraction failed: {e_pymupdf}. Trying PDFPlumber..."
            )
            # Fall through to PDFPlumber if PyMuPDF failed

    # --- Try PDFPlumber as Fallback ---
    if _PDFPLUMBER_AVAILABLE and pdfplumber:
        logger.debug(f"Attempting direct text extraction with PDFPlumber for {file_path}")
        try:
            # pdfplumber might require explicit closing
            pdf = pdfplumber.open(file_path)  # type: ignore
            try:
                total_pages = len(pdf.pages)
                end_page = (
                    total_pages if max_pages <= 0 else min(start_page + max_pages, total_pages)
                )
                start_page = min(start_page, total_pages)
                end_page = max(start_page, end_page)

                for i in range(start_page, end_page):
                    try:
                        page = pdf.pages[i]
                        # Use slightly more tolerant settings
                        page_text = (
                            page.extract_text(x_tolerance=2, y_tolerance=2, keep_blank_chars=True)
                            or ""
                        )
                        texts.append(page_text)
                        if len(page_text.strip()) >= min_chars:
                            has_text = True
                    except Exception as e_page:
                        logger.warning(
                            f"PDFPlumber: Error extracting text from page {i + 1}: {e_page}"
                        )
                        texts.append(f"[Page {i + 1} Extraction Error: PDFPlumber]")
                logger.debug(
                    f"PDFPlumber extracted {len(texts)} pages. Found meaningful text: {has_text}."
                )
                return texts, has_text
            finally:
                pdf.close()  # Ensure file handle is closed
        except Exception as e_plumber:
            logger.error(f"PDFPlumber direct text extraction failed: {e_plumber}", exc_info=True)
            # If PyMuPDF also failed (or wasn't available), raise the final error
            if (
                not _PYMUPDF_AVAILABLE
            ):  # Only raise if it was the only option tried or PyMuPDF failed before
                raise ToolError(
                    "DIRECT_EXTRACTION_FAILED",
                    details={"reason": "PDFPlumber failed", "error": str(e_plumber)},
                ) from e_plumber
            else:  # PyMuPDF failed first, now PDFPlumber failed
                raise ToolError(
                    "DIRECT_EXTRACTION_FAILED",
                    details={
                        "reason": "Both PyMuPDF and PDFPlumber failed",
                        "error": str(e_plumber),
                    },
                ) from e_plumber

    # --- If neither library worked ---
    logger.error(
        "No functional direct PDF text extraction library (PyMuPDF or PDFPlumber) available or both failed."
    )
    raise ToolError("DIRECT_EXTRACTION_FAILED", details={"reason": "No available/working library"})


def _ocr_convert_pdf_to_images(
    file_path: Path, start_page: int = 0, max_pages: int = 0, dpi: int = 300
) -> List["PILImage.Image"]:
    """Converts PDF path pages to PIL Images using 0-based indexing internally (sync function)."""
    _ocr_check_dep("pdf2image", _PDF2IMAGE_AVAILABLE, "PDF->Image Conversion")
    _ocr_check_dep("Pillow", _PIL_AVAILABLE, "PDF->Image Conversion")
    if convert_from_path is None:
        raise ToolError("INTERNAL_ERROR", details={"reason": "pdf2image.convert_from_path is None"})

    try:
        # pdf2image uses 1-based indexing for first_page/last_page args
        first_page_1based = start_page + 1
        last_page_1based = None if max_pages <= 0 else first_page_1based + max_pages - 1
        logger.debug(
            f"Converting PDF {file_path} (pages {first_page_1based}-{last_page_1based or 'end'}, dpi={dpi})"
        )

        with _span(f"pdf2image_path_p{first_page_1based}-{last_page_1based or 'end'}"):
            # pdf2image handles its own temporary files internally if output_folder=None
            # Using a TemporaryDirectory might be slightly less efficient but ensures cleanup
            with tempfile.TemporaryDirectory() as temp_dir:
                images = convert_from_path(  # type: ignore
                    file_path,
                    dpi=dpi,
                    first_page=first_page_1based,
                    last_page=last_page_1based,
                    output_folder=temp_dir,  # Recommended for multi-threading stability
                    fmt="png",
                    thread_count=max(1, os.cpu_count() // 2 if os.cpu_count() else 1),
                    use_pdftocairo=True,  # Often more reliable than pdftoppm
                )
        logger.info(f"Converted {len(images)} pages from PDF path.")
        return images  # type: ignore
    except Exception as e:  # Catch specific pdf2image errors if library defines them
        logger.error(f"PDF path to image conversion failed: {e}", exc_info=True)
        raise ToolError(
            "PDF_CONVERSION_FAILED", details={"reason": "pdf2image path failed", "error": str(e)}
        ) from e


def _ocr_preprocess_image(
    image: "PILImage.Image", preprocessing_options: Optional[Dict[str, Any]] = None
) -> "PILImage.Image":
    """Preprocesses an image for better OCR results (sync function)."""
    if not _PIL_AVAILABLE:
        logger.warning("Pillow (PIL) not available. Skipping preprocessing.")
        return image
    if not ImageEnhance or not ImageFilter:  # Check specifically for submodules
        logger.warning("PIL ImageEnhance or ImageFilter missing. Some enhancements skipped.")

    can_use_cv2 = _CV2_AVAILABLE and _NUMPY_AVAILABLE and cv2 is not None and np is not None
    if (
        not can_use_cv2
        and preprocessing_options
        and any(k in preprocessing_options for k in ["denoise", "threshold", "deskew"])
    ):
        logger.warning("OpenCV/NumPy missing. Advanced preprocessing disabled.")

    prep_opts = {
        "denoise": True,
        "threshold": "otsu",
        "deskew": True,
        "enhance_contrast": True,
        "enhance_brightness": False,
        "enhance_sharpness": False,
        "apply_filters": [],
        "resize_factor": 1.0,
        **(preprocessing_options or {}),
    }
    logger.debug(f"Applying preprocessing with options: {prep_opts}")

    img_pil = image.copy()
    # Apply PIL enhancements first
    if ImageEnhance:
        if prep_opts.get("enhance_brightness"):
            img_pil = ImageEnhance.Brightness(img_pil).enhance(1.3)
        if prep_opts.get("enhance_contrast") and not can_use_cv2:
            img_pil = ImageEnhance.Contrast(img_pil).enhance(1.4)
        if prep_opts.get("enhance_sharpness"):
            img_pil = ImageEnhance.Sharpness(img_pil).enhance(1.5)
    if ImageFilter:
        filters = prep_opts.get("apply_filters", [])
        for filter_name in filters:
            try:
                if filter_name == "unsharp_mask":
                    img_pil = img_pil.filter(ImageFilter.UnsharpMask(radius=2, percent=150))
                elif filter_name == "detail":
                    img_pil = img_pil.filter(ImageFilter.DETAIL)
                elif filter_name == "edge_enhance":
                    img_pil = img_pil.filter(ImageFilter.EDGE_ENHANCE)
                elif filter_name == "smooth":
                    img_pil = img_pil.filter(ImageFilter.SMOOTH)
                else:
                    logger.warning(f"Unknown PIL filter: {filter_name}")
            except Exception as e:
                logger.warning(f"PIL filter '{filter_name}' failed: {e}")

    if not can_use_cv2:
        return img_pil  # Return PIL-enhanced if CV2 unavailable

    # OpenCV Processing
    try:
        img_cv = np.array(img_pil)
        if len(img_cv.shape) == 3 and img_cv.shape[2] == 3:
            gray = cv2.cvtColor(img_cv, cv2.COLOR_RGB2GRAY)
        elif len(img_cv.shape) == 3 and img_cv.shape[2] == 4:
            gray = cv2.cvtColor(img_cv, cv2.COLOR_RGBA2GRAY)
        else:
            gray = img_cv

        original_height, original_width = gray.shape[:2]
        deskewed_gray = gray.copy()  # Operate on this copy

        # Deskewing (best on grayscale before thresholding might change shapes)
        if prep_opts.get("deskew", True):
            try:
                # Use inverted image for finding text blocks if background is light
                mean_intensity = np.mean(gray)
                invert_for_deskew = mean_intensity > 128
                deskew_input = cv2.bitwise_not(gray) if invert_for_deskew else gray

                # Use a less aggressive threshold for finding angle
                _, angle_thresh = cv2.threshold(
                    deskew_input, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU
                )
                coords = cv2.findNonZero(angle_thresh)
                if coords is not None and len(coords) > 10:
                    angle = cv2.minAreaRect(coords)[-1]
                    if angle < -45:
                        angle = -(90 + angle)
                    else:
                        angle = -angle

                    if abs(angle) > 0.1:
                        (h, w) = gray.shape[:2]
                        center = (w // 2, h // 2)
                        M = cv2.getRotationMatrix2D(center, angle, 1.0)
                        # Rotate original grayscale image
                        deskewed_gray = cv2.warpAffine(
                            gray, M, (w, h), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE
                        )
                        logger.debug(f"Deskewed image by {angle:.2f} degrees.")
            except Exception as e_deskew:
                logger.warning(f"Deskewing failed: {e_deskew}. Using original orientation.")
                deskewed_gray = gray  # Reset to original gray if deskew fails

        processed_img = deskewed_gray  # Start processing from (potentially) deskewed gray

        # Adaptive scaling calculation (applied later)
        resize_factor = prep_opts.get("resize_factor", 1.0)
        if resize_factor == 1.0:
            longest_edge = max(original_width, original_height)
            target_low, target_high = 1500, 3500
            if 0 < longest_edge < target_low:
                resize_factor = math.ceil(target_low / longest_edge * 10) / 10
            elif longest_edge > target_high:
                resize_factor = math.floor(target_high / longest_edge * 10) / 10
            resize_factor = max(0.5, min(3.0, resize_factor))

        # Contrast enhancement on grayscale
        if prep_opts.get("enhance_contrast", True):
            clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
            processed_img = clahe.apply(processed_img)

        # Denoising grayscale (before thresholding)
        if prep_opts.get("denoise", True):
            # Adjust h based on image size? Might be overkill.
            # h_param = math.ceil(10 * math.log10(max(10, min(original_width, original_height))))
            processed_img = cv2.fastNlMeansDenoising(processed_img, None, 10, 7, 21)

        # Thresholding
        threshold_method = prep_opts.get("threshold", "otsu")
        if threshold_method == "otsu":
            # No need for blur if denoised already
            _, processed_img = cv2.threshold(
                processed_img, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU
            )
        elif threshold_method == "adaptive":
            block_size = max(11, math.floor(min(processed_img.shape[:2]) / 20) * 2 + 1)
            processed_img = cv2.adaptiveThreshold(
                processed_img, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, block_size, 5
            )
        # If no threshold, check background and invert if needed for Tesseract
        elif np.mean(processed_img) < 128:
            processed_img = cv2.bitwise_not(processed_img)

        # Resizing (applied last)
        if resize_factor != 1.0:
            current_h, current_w = processed_img.shape[:2]
            new_w = math.ceil(current_w * resize_factor)
            new_h = math.ceil(current_h * resize_factor)
            processed_img = cv2.resize(processed_img, (new_w, new_h), interpolation=cv2.INTER_CUBIC)
            logger.debug(f"Resized image by factor {resize_factor:.2f} to {new_w}x{new_h}")

        final_pil_image = Image.fromarray(processed_img)
        return final_pil_image
    except Exception as e_cv:
        logger.error(f"OpenCV preprocessing failed: {e_cv}", exc_info=True)
        return img_pil  # Fallback to PIL-processed image


def _ocr_run_tesseract(
    image: "PILImage.Image", ocr_language: str = "eng", ocr_config: str = ""
) -> str:
    """Extracts text from an image using Tesseract OCR (sync function)."""
    _ocr_check_dep("pytesseract", _PYTESSERACT_AVAILABLE, "OCR Text Extraction")
    _ocr_check_dep("Pillow", _PIL_AVAILABLE, "OCR Text Extraction")
    if pytesseract is None:
        raise ToolError("INTERNAL_ERROR", details={"reason": "pytesseract is None"})
    try:
        # Combine language and custom config flags
        custom_config = f"-l {ocr_language} {ocr_config}".strip()
        logger.debug(f"Running Tesseract with config: '{custom_config}'")
        with _span(f"pytesseract_ocr_{ocr_language}"):
            # Use a timeout? Tesseract can sometimes hang. Requires subprocess handling.
            # For simplicity, no timeout implemented here.
            text = pytesseract.image_to_string(
                image, config=custom_config, timeout=60
            )  # Add 60s timeout
        logger.debug(f"Tesseract extracted {len(text)} characters.")
        return text or ""  # Ensure string return
    except pytesseract.TesseractNotFoundError as e:
        logger.error("Tesseract executable not found or not in PATH.")
        raise ToolError("DEPENDENCY_MISSING", details={"dependency": "Tesseract OCR Engine"}) from e
    except RuntimeError as e_runtime:  # Catch Tesseract runtime errors (like timeout)
        logger.error(f"Tesseract runtime error: {e_runtime}", exc_info=True)
        raise ToolError(
            "OCR_FAILED", details={"engine": "Tesseract", "error": f"Runtime error: {e_runtime}"}
        ) from e_runtime
    except Exception as e:
        logger.error(f"Tesseract OCR extraction failed: {e}", exc_info=True)
        raise ToolError("OCR_FAILED", details={"engine": "Tesseract", "error": str(e)}) from e


def _ocr_is_text_mostly_noise(text: str, noise_threshold: float = 0.4) -> bool:
    """
    Determine if extracted text is mostly noise based on character distribution.
    Considers alphanumeric, whitespace, and common punctuation as 'valid'.

    Args:
        text: The text string to analyze.
        noise_threshold: The ratio (0.0 to 1.0) of non-valid characters above which
                         the text is considered noisy. Default is 0.4 (40%).

    Returns:
        True if the text is considered mostly noise, False otherwise.
    """
    if not text or not isinstance(text, str):
        return False  # Empty or invalid input is not noise

    text_length = len(text)
    if text_length < 20:  # Don't evaluate very short strings
        return False

    # Define a set of characters generally expected in non-noisy text
    # (alphanumeric, whitespace, common punctuation/symbols)
    # Adding more symbols that might appear legitimately in documents
    valid_char_pattern = re.compile(r"[a-zA-Z0-9\s.,;:!?\"'()\[\]{}%/$£€¥₽₹#@&*+=<>~|_^-]")

    valid_chars_count = len(valid_char_pattern.findall(text))

    # Calculate the ratio of characters *not* matching the valid pattern
    noise_ratio = 1.0 - (valid_chars_count / text_length)

    is_noise = noise_ratio > noise_threshold
    if is_noise:
        # Log only a snippet to avoid flooding logs with potentially large noisy text
        snippet = text.replace("\n", " ")[:100]  # Replace newlines for cleaner log output
        logger.debug(
            f"Text flagged as noisy (Ratio: {noise_ratio:.2f} > {noise_threshold}): '{snippet}...'"
        )

    return is_noise


def _ocr_is_likely_header_or_footer(text: str, line_length_threshold: int = 80) -> bool:
    """
    Determine if a single line of text is likely a header or footer based on common patterns.

    Args:
        text: The line of text to evaluate.
        line_length_threshold: Lines longer than this are less likely to be headers/footers. Default 80.

    Returns:
        True if the line matches common header/footer patterns, False otherwise.
    """
    text = text.strip()
    if not text or len(text) > line_length_threshold:
        return False

    # --- Pattern Checks ---
    # 1. Page Number patterns (robust check)
    #    - "Page X", "P. X", "X / Y", "- X -", etc.
    #    - Allows for variations in spacing and separators
    if re.search(r"(?i)\b(page|p[ag]{1,2}\.?|seite|s\.?)\s*\d+", text):
        return True
    if re.match(r"^\s*[-–—]?\s*\d+\s*[/of\s]+\s*\d+\s*[-–—]?\s*$", text):
        return True  # e.g., "1 / 10", "1 of 10"
    if re.match(r"^\s*[-–—]?\s*\d+\s*[-–—]?\s*$", text):
        return True  # Just a number, possibly bracketed

    # 2. Date patterns
    #    - "Month Day, Year", "DD/MM/YYYY", "YYYY-MM-DD", etc.
    if re.search(
        r"\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{1,2},?\s+\d{4}", text, re.I
    ):
        return True
    if re.search(r"\b\d{1,2}[./-]\d{1,2}[./-]\d{2,4}\b", text):
        return True
    if re.search(r"\b\d{4}[./-]\d{1,2}[./-]\d{1,2}\b", text):
        return True  # ISO-like

    # 3. Common keywords (case-insensitive start of line)
    if re.match(
        r"^(confidential|internal use only|draft|proprietary|for discussion purposes)", text, re.I
    ):
        return True
    if re.match(r"^(copyright|\(c\)|©)\s*\d*", text, re.I):
        return True

    # 4. Repeated characters (often used as separators)
    #    - Check if the line consists mostly of one or two non-alphanumeric characters
    non_alnum_chars = re.sub(r"[a-zA-Z0-9\s]", "", text)
    if len(non_alnum_chars) > 5 and len(set(non_alnum_chars)) <= 2:
        return True

    # 5. Company Names / Document Titles (Heuristic - might be too broad)
    #    - Check if it's short, title-cased, and doesn't end in punctuation?
    # if len(text.split()) < 7 and text == text.title() and not text.endswith(('.', '?', '!')):
    #     # Further check: Is this text repeated elsewhere? (Needs broader context)
    #     pass # This heuristic is often unreliable without more context.

    # 6. All Caps Short Lines (Potential titles/headers)
    if text.isupper() and len(text.split()) < 7 and len(text) > 3:
        return True

    return False  # Default: Not a header/footer


def _ocr_remove_headers_and_footers(text: str, max_lines_check: int = 5) -> str:
    """
    Removes likely headers and footers from the top/bottom of the text block.

    Args:
        text: The block of text (potentially multiple pages concatenated).
        max_lines_check: How many lines from the top and bottom to examine. Default 5.

    Returns:
        The text with potential header/footer lines removed.
    """
    if not text or not isinstance(text, str):
        return ""

    lines = text.splitlines()
    num_lines = len(lines)

    # Don't process if text is too short to reliably identify headers/footers
    if num_lines < max_lines_check * 2:
        return text

    lines_to_remove_indices: Set[int] = set()

    # Check top lines
    for i in range(max_lines_check):
        if i < num_lines:  # Ensure index is valid
            line_text = lines[i]
            # Also check if the line is very short (e.g., just whitespace remnants)
            if _ocr_is_likely_header_or_footer(line_text) or len(line_text.strip()) <= 2:
                lines_to_remove_indices.add(i)
            # Stop checking top lines if a probable content line is found early
            elif len(line_text) > 80 and i < max_lines_check // 2:  # Heuristic for content line
                break
        else:  # Should not happen given initial num_lines check, but safety
            break

    # Check bottom lines
    for i in range(max_lines_check):
        idx = num_lines - 1 - i
        # Ensure index is valid and not already marked for removal from top scan
        if idx >= 0 and idx not in lines_to_remove_indices:
            line_text = lines[idx]
            if _ocr_is_likely_header_or_footer(line_text) or len(line_text.strip()) <= 2:
                lines_to_remove_indices.add(idx)
            # Stop checking bottom lines if a probable content line is found early
            elif len(line_text) > 80 and i < max_lines_check // 2:
                break
        elif idx < 0:  # Reached top of file during bottom check
            break

    if not lines_to_remove_indices:
        return text  # No lines identified for removal

    logger.debug(f"Removing {len(lines_to_remove_indices)} potential header/footer lines.")

    # Build the result, skipping removed lines
    result_lines = [line for i, line in enumerate(lines) if i not in lines_to_remove_indices]

    # Remove leading/trailing blank lines potentially left after removal
    # This needs care: find first/last non-blank line indices
    first_content_line = -1
    last_content_line = -1
    for i, line in enumerate(result_lines):
        if line.strip():
            if first_content_line == -1:
                first_content_line = i
            last_content_line = i

    if first_content_line == -1:  # All lines were removed or blank
        return ""
    else:
        # Join only the content lines, preserving internal blank lines
        cleaned_text = "\n".join(result_lines[first_content_line : last_content_line + 1])
        return cleaned_text


async def _ocr_enhance_text_chunk(
    chunk: str, output_format: str = "markdown", remove_headers: bool = False
) -> str:
    """Enhances OCR text chunk using LLM (standalone internal helper)."""
    # --- Apply Basic Rule-based Cleaning First ---
    cleaned_text = chunk.strip()  # Work on a copy

    # Join words incorrectly split across lines (common OCR artifact)
    cleaned_text = re.sub(r"(\w+)-\s*\n\s*(\w+)", r"\1\2", cleaned_text)

    # Normalize multiple whitespace characters (including newlines within paragraphs)
    # This is aggressive and might merge intended line breaks within code/poetry
    # Consider a less aggressive approach if preserving specific line breaks is crucial.
    # cleaned_text = re.sub(r"\s+", " ", cleaned_text) # Too aggressive

    # Normalize space/tab characters to single space
    cleaned_text = re.sub(r"[ \t]+", " ", cleaned_text)
    # Collapse multiple blank lines (2+ newlines) into exactly two newlines
    cleaned_text = re.sub(r"\n\s*\n", "\n\n", cleaned_text)

    # Optional header/footer removal using rules *before* LLM
    if remove_headers:
        original_len = len(cleaned_text)
        cleaned_text = _ocr_remove_headers_and_footers(cleaned_text)
        if len(cleaned_text) < original_len:
            logger.debug("Applied rule-based header/footer removal pre-LLM.")

    # Check for noise after initial cleaning
    if _ocr_is_text_mostly_noise(cleaned_text):
        logger.warning(
            "Text chunk noisy after basic cleaning, LLM enhancement might be less effective."
        )
        # Decide whether to proceed or return early based on noise level?
        # For now, proceed with LLM enhancement.

    # --- LLM Prompt Generation ---
    format_instruction = ""
    if output_format == "markdown":
        format_instruction = """
2. Format as clean, readable markdown:
   - Use appropriate heading levels (#, ##, etc.). Infer structure where possible.
   - Format lists correctly (bulleted or numbered).
   - Apply emphasis (*italic*) and strong (**bold**) sparingly where appropriate.
   - Represent tabular data using markdown table syntax IF table structure is clearly identifiable.
   - Use code blocks (```) for code snippets or equations if detected."""
    else:  # output_format == "text"
        format_instruction = """
2. Format as clean, readable plain text:
   - Ensure clear paragraph separation (double newline).
   - Maintain list structures with standard markers (e.g., -, 1.).
   - Avoid markdown syntax like #, *, _, ```, etc."""

    header_footer_instruction = (
        "Remove any remaining headers, footers, and page numbers."
        if remove_headers
        else "Preserve all content including potential headers/footers."
    )
    prompt = f"""You are an expert text processor specialized in correcting OCR errors from scanned documents. Please process the following text according to these instructions:

1. Fix OCR-induced errors:
   - Correct character recognition errors (e.g., 'rn' vs 'm', 'O' vs '0', 'l' vs '1', 'S' vs '5').
   - Join words incorrectly split across lines (e.g., "hyphen-\nation").
   - Merge paragraphs that were artificially split by page breaks or scanning artifacts.
   - Split run-on paragraphs where a clear topic shift or structural break (like a list starting) occurs.
   - Use context to resolve ambiguities and reconstruct the original meaning accurately.
{format_instruction}
3. Clean up formatting:
   - Remove redundant spaces within lines.
   - Ensure consistent paragraph spacing (double newline between paragraphs).
   - {header_footer_instruction}

4. IMPORTANT: Preserve all meaningful content and the original structure as much as possible. Do not add information or summaries. Do not change the substance of the text. Focus solely on fixing OCR errors and applying the requested formatting based *only* on the input text provided.

Input Text:
```text
{cleaned_text}
```

Corrected Output ({output_format}):"""

    try:
        logger.debug(
            f"Sending chunk (len={len(cleaned_text)}) to LLM for enhancement (format={output_format}, rm_hdrs={remove_headers})."
        )
        # Use a capable model (adjust model name as needed)
        provider = Provider.OPENAI.value
        model = "gpt-4o-mini"

        # Estimate max tokens needed
        estimated_input_tokens = len(cleaned_text) // 3
        buffer_factor = 1.4 if output_format == "markdown" else 1.2  # Slightly more buffer
        llm_max_tokens = int(estimated_input_tokens * buffer_factor) + 500
        # Cap based on typical context window limits (e.g., ~16k tokens for GPT-4 Turbo input, allow ample output)
        llm_max_tokens = max(1000, min(llm_max_tokens, 8000))

        # Assume _standalone_llm_call is defined elsewhere
        enhanced_text = await _standalone_llm_call(
            prompt=prompt,
            provider=provider,
            model=model,
            temperature=0.1,  # Very low temperature for factual correction
            max_tokens=llm_max_tokens,
        )

        # --- Post-processing LLM Output ---
        # Remove potential preamble/apologies
        enhanced_text = re.sub(
            r"^(Okay, |Here is |Sure, |Here['’]s |Certainly, |Based on the text provided.*?\n)[:\n]?\s*",
            "",
            enhanced_text,
            flags=re.IGNORECASE | re.DOTALL,
        )
        # Remove potential markdown fences around the whole output
        enhanced_text = re.sub(
            r"^\s*```(?:\w+\n)?([\s\S]*?)\n?```\s*$", r"\1", enhanced_text
        ).strip()

        logger.debug(f"LLM enhancement returned text (len={len(enhanced_text)}).")
        return enhanced_text

    except ToolError as e:
        # Log the specific ToolError and fallback
        logger.error(
            f"LLM text enhancement failed with ToolError: {e.error_code} - {str(e)}. Returning pre-LLM cleaned text."
        )
        return cleaned_text
    except Exception as e:
        # Log unexpected errors and fallback
        logger.error(f"Unexpected error during LLM text enhancement: {e}", exc_info=True)
        return cleaned_text


def _ocr_validate_file_path(file_path: str, expected_extension: Optional[str] = None) -> Path:
    """Validates a file path exists and optionally has the expected extension."""
    if not file_path or not isinstance(file_path, str):
        raise ToolInputError("File path cannot be empty or non-string", param_name="file_path")

    try:
        # Expand user directory and normalize path separators
        path = Path(os.path.expanduser(os.path.normpath(file_path)))
    except Exception as e:
        raise ToolInputError(
            f"Invalid file path format: {file_path}. Error: {e}", param_name="file_path"
        ) from e

    if not path.exists():
        raise ToolInputError(f"File not found at path: {path}", param_name="file_path")
    if not path.is_file():
        raise ToolInputError(f"Path exists but is not a file: {path}", param_name="file_path")
    # Check extension case-insensitively
    if expected_extension and not path.suffix.lower() == expected_extension.lower():
        raise ToolInputError(
            f"File does not have the expected extension ({expected_extension}): {path}",
            param_name="file_path",
        )
    # Optional: Check read permissions?
    # if not os.access(path, os.R_OK):
    #     raise ToolInputError(f"Cannot read file (permission denied): {path}", param_name="file_path")

    return path


def _ocr_detect_tables(image: "PILImage.Image") -> List[Tuple[int, int, int, int]]:
    """Detects potential tables in an image using OpenCV (sync function)."""
    # Check dependencies first
    if not _CV2_AVAILABLE or not _NUMPY_AVAILABLE or not _PIL_AVAILABLE:
        logger.warning("Cannot detect tables: OpenCV, NumPy, or Pillow not available.")
        return []
    # Ensure library objects are valid
    if cv2 is None or np is None:
        logger.warning("Cannot detect tables: OpenCV or NumPy object is None.")
        return []

    try:
        img = np.array(image)
        # Convert to grayscale if necessary
        if len(img.shape) == 3:
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        elif len(img.shape) == 2:
            gray = img
        else:
            logger.warning(f"Unexpected image shape for table detection: {img.shape}")
            return []

        # --- Table Detection Logic (Example using line detection) ---
        # 1. Thresholding (Adaptive often works well for lines)
        thresh_inv = cv2.adaptiveThreshold(
            gray, 255, cv2.ADAPTIVE_THRESH_MEAN_C, cv2.THRESH_BINARY_INV, 15, 5
        )

        # 2. Detect Horizontal Lines
        horizontal_kernel = cv2.getStructuringElement(
            cv2.MORPH_RECT, (min(40, gray.shape[1] // 10), 1)
        )  # Kernel size relative to width
        detected_horizontal = cv2.morphologyEx(
            thresh_inv, cv2.MORPH_OPEN, horizontal_kernel, iterations=2
        )
        cnts_h, _ = cv2.findContours(
            detected_horizontal, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE
        )

        # 3. Detect Vertical Lines
        vertical_kernel = cv2.getStructuringElement(
            cv2.MORPH_RECT, (1, min(40, gray.shape[0] // 10))
        )  # Kernel size relative to height
        detected_vertical = cv2.morphologyEx(
            thresh_inv, cv2.MORPH_OPEN, vertical_kernel, iterations=2
        )
        cnts_v, _ = cv2.findContours(detected_vertical, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

        # 4. Combine contours or find bounding boxes of large contours containing lines
        # Strategy: Find large contours in the original inverted threshold image,
        # then check if those contours contain significant horiz/vert lines.
        contours, _ = cv2.findContours(thresh_inv, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

        table_regions = []
        img_area = img.shape[0] * img.shape[1]
        min_table_area = img_area * 0.01  # Lower threshold slightly (1%)
        min_dimension = 50  # Min width/height for a contour to be considered

        for contour in contours:
            x, y, w, h = cv2.boundingRect(contour)
            area = w * h
            aspect_ratio = w / max(1, h)

            # Basic filtering based on size and aspect ratio
            if (
                area > min_table_area
                and w > min_dimension
                and h > min_dimension
                and 0.1 < aspect_ratio < 10.0
            ):
                # Check for significant presence of detected lines within this bounding box
                roi_h = detected_horizontal[y : y + h, x : x + w]
                roi_v = detected_vertical[y : y + h, x : x + w]
                # Heuristic: Check if non-zero pixels (lines) exceed a small fraction of the ROI area or length
                min_line_pixels_h = w * 0.3  # Require horizontal lines covering ~30% width
                min_line_pixels_v = h * 0.3  # Require vertical lines covering ~30% height
                if (
                    cv2.countNonZero(roi_h) > min_line_pixels_h
                    and cv2.countNonZero(roi_v) > min_line_pixels_v
                ):
                    table_regions.append((x, y, w, h))
                # else:
                #    logger.debug(f"Contour rejected: area={area}, w={w}, h={h}, h_px={cv2.countNonZero(roi_h)}, v_px={cv2.countNonZero(roi_v)}")

        # Optional: Merge overlapping bounding boxes (omitted for simplicity)
        # merged_regions = merge_overlapping_boxes(table_regions) # Needs implementation

        logger.debug(f"Detected {len(table_regions)} potential table regions.")
        return table_regions

    except Exception as e:
        logger.error(f"OpenCV Table detection failed: {e}", exc_info=True)
        return []


def _ocr_process_toc(toc: List) -> List[Dict[str, Any]]:
    """Processes a PDF table of contents (from PyMuPDF) into a nested structure."""
    if not toc:
        return []
    result: List[Dict[str, Any]] = []
    # Stack stores tuples: (level, parent_list_to_append_to)
    stack: List[Tuple[int, List]] = [(-1, result)]
    for item in toc:
        # PyMuPDF TOC item format: [level, title, page, ?dest_dict]
        if not isinstance(item, (list, tuple)) or len(item) < 3:
            logger.warning(f"Skipping malformed TOC item: {item}")
            continue
        try:
            level = int(item[0])
            title = str(item[1])
            page = int(item[2])
        except (ValueError, TypeError, IndexError) as e:
            logger.warning(f"Error parsing TOC item '{item}': {e}")
            continue

        # Pop stack until parent level is found
        while stack[-1][0] >= level:
            stack.pop()
            if not stack:  # Should not happen with initial (-1, result)
                logger.error("TOC stack became empty unexpectedly.")
                return result  # Return what we have so far

        # Create new entry and add to parent's children list
        entry: Dict[str, Any] = {"title": title, "page": page, "children": []}
        stack[-1][1].append(entry)
        # Push current entry onto stack for potential children
        stack.append((level, entry["children"]))
    return result


def _ocr_split_text_into_chunks(
    text: str, max_chunk_size: int = 8000, overlap: int = 200
) -> List[str]:
    """Splits text into chunks, trying to respect paragraphs and sentences (sync function)."""
    if not text or not isinstance(text, str):
        return []

    max_chunk_size = max(1000, min(max_chunk_size, 15000))  # Sensible limits
    overlap = max(50, min(overlap, max_chunk_size // 4))
    # Ensure min_chunk_size is reasonable, at least larger than overlap
    min_chunk_size = max(overlap * 2, 100)

    chunks = []
    start_index = 0
    text_len = len(text)

    while start_index < text_len:
        end_index = min(start_index + max_chunk_size, text_len)

        # Handle the last chunk directly
        if end_index == text_len:
            chunk = text[start_index:end_index]
            # Only add if it has meaningful content (more than just whitespace)
            if chunk.strip():
                chunks.append(chunk)
            break  # End of text reached

        best_split_index = -1
        # Prefer double newline (paragraph break)
        split_point_para = text.rfind("\n\n", max(start_index, end_index - overlap * 2), end_index)
        if split_point_para != -1 and split_point_para > start_index:  # Ensure split is after start
            # Check if this split results in a reasonably sized chunk
            if (split_point_para + 2 - start_index) >= min_chunk_size:
                best_split_index = split_point_para + 2

        # If no good paragraph break, try sentence breaks
        if best_split_index == -1:
            sentence_break_pattern = r"[.?!]['\"]?(\s|\n|$)"  # Include end of string
            # Search within a reasonable lookback window
            search_region_start = max(start_index, end_index - overlap)
            search_region = text[search_region_start:end_index]
            matches = list(re.finditer(sentence_break_pattern, search_region))
            if matches:
                # Find the offset of the last match within the search region
                last_match_end_offset = matches[-1].end()
                # Calculate the split point relative to the original string
                split_point_sentence = search_region_start + last_match_end_offset
                # Check if this split is valid and creates a reasonably sized chunk
                if (
                    split_point_sentence > start_index
                    and (split_point_sentence - start_index) >= min_chunk_size
                ):
                    best_split_index = split_point_sentence

        # Fallback to single newline or space if still no good break
        if best_split_index == -1:
            split_point_newline = text.rfind("\n", max(start_index, end_index - overlap), end_index)
            split_point_space = text.rfind(" ", max(start_index, end_index - overlap), end_index)
            # Choose the latest valid break (newline or space)
            split_point_fallback = max(split_point_newline, split_point_space)
            if (
                split_point_fallback > start_index
                and (split_point_fallback + 1 - start_index) >= min_chunk_size
            ):
                best_split_index = split_point_fallback + 1

        # Force split at max_chunk_size boundary if no suitable break found,
        # or if the best found break is too early (making the chunk too small)
        if (
            best_split_index <= start_index
            or (best_split_index - start_index) < min_chunk_size // 2
        ):
            # Check if simply taking end_index results in a valid chunk start for next iteration
            potential_next_start = max(start_index + 1, end_index - overlap)
            if potential_next_start < text_len:  # Avoid forcing if it's the last chunk anyway
                best_split_index = end_index
            else:  # If forcing split here would make the loop end, try a slightly earlier hard split?
                # For simplicity, let's stick to end_index, the loop termination handles the last part.
                best_split_index = end_index

        # Extract the chunk
        chunk = text[start_index:best_split_index]
        if chunk.strip():  # Only add non-empty chunks
            chunks.append(chunk)

        # Calculate the start index for the next chunk
        next_start = max(start_index + 1, best_split_index - overlap)

        # Ensure substantial forward progress to avoid infinite loops on edge cases
        # Use max_chunk_size here instead of the undefined 'size'
        min_progress = min(max_chunk_size // 10, 50)  # Ensure we advance by at least a small amount
        next_start = max(next_start, start_index + min_progress)

        # Safety check: don't let next_start go beyond the text length
        start_index = min(next_start, text_len)

    # Filter out any potential empty strings added during edge cases
    final_chunks = [c for c in chunks if c]

    logger.debug(f"Split text ({text_len} chars) into {len(final_chunks)} chunks")
    return final_chunks


async def _ocr_assess_text_quality(original_text: str, enhanced_text: str) -> Dict[str, Any]:
    """Assesses the quality of OCR enhancement using LLM (Standalone)."""
    if not original_text and not enhanced_text:
        return {"score": 0, "explanation": "No text provided for assessment.", "examples": []}
    if not original_text:
        return {
            "score": 100,
            "explanation": "Original text was empty, enhanced text provided.",
            "examples": [],
        }
    if not enhanced_text:
        return {
            "score": 0,
            "explanation": "Enhanced text is empty, original text was not.",
            "examples": [],
        }

    max_sample = 4000
    original_sample = original_text[:max_sample] + (
        "..." if len(original_text) > max_sample else ""
    )
    enhanced_sample = enhanced_text[:max_sample] + (
        "..." if len(enhanced_text) > max_sample else ""
    )

    prompt = f"""Please assess the quality improvement from the 'Original OCR Text' to the 'Enhanced Text'. Focus on:
1. Correction of OCR errors (typos, spacing, broken words).
2. Improvement in formatting and readability (paragraphs, lists, structure).
3. Accuracy in preserving the original meaning and content.
4. Effectiveness of removing noise (like headers/footers if applicable).

Original OCR Text:
```
{original_sample}
```

Enhanced Text:
```
{enhanced_sample}
```

Provide your assessment ONLY in the following JSON format:
{{
  "score": <integer score 0-100, where 100 is perfect enhancement>,
  "explanation": "<brief explanation of the score, highlighting key improvements or remaining issues>",
  "examples": [
    "<example 1 of a specific correction or improvement>",
    "<example 2>",
    "<example 3 (optional)>"
  ]
}}
Do not add any text before or after the JSON object.
"""

    try:
        logger.debug("Requesting LLM quality assessment.")
        assessment_json_str = await _standalone_llm_call(
            prompt=prompt, max_tokens=500, temperature=0.2
        )
        try:
            json_match = re.search(r"```(?:json)?\s*([\s\S]+?)\s*```", assessment_json_str)
            json_str = json_match.group(1).strip() if json_match else assessment_json_str.strip()
            start_brace = json_str.find("{")
            end_brace = json_str.rfind("}")
            if start_brace != -1 and end_brace != -1 and start_brace < end_brace:
                json_str = json_str[start_brace : end_brace + 1]
            elif not json_str.startswith("{"):
                raise ValueError("Could not find JSON object boundaries.")

            assessment_data = json.loads(json_str)
            if (
                not isinstance(assessment_data, dict)
                or "score" not in assessment_data
                or "explanation" not in assessment_data
                or "examples" not in assessment_data
                or not isinstance(assessment_data["examples"], list)
            ):
                raise ValueError("Parsed JSON has incorrect structure.")
            try:
                assessment_data["score"] = (
                    int(assessment_data["score"]) if assessment_data["score"] is not None else None
                )
            except (ValueError, TypeError):
                assessment_data["score"] = None
            assessment_data["explanation"] = str(assessment_data["explanation"])
            assessment_data["examples"] = [str(ex) for ex in assessment_data["examples"]]
            logger.debug(
                f"Quality assessment received: Score {assessment_data.get('score', 'N/A')}"
            )
            return assessment_data
        except (json.JSONDecodeError, ValueError, TypeError) as e:
            logger.error(
                f"Failed to parse quality assessment JSON: {e}. Raw:\n{assessment_json_str}"
            )
            return {
                "score": None,
                "explanation": f"Parse failed: {e}",
                "examples": [],
                "raw_response": assessment_json_str,
            }
    except Exception as e:
        logger.error(f"Error during LLM quality assessment call: {e}", exc_info=True)
        return {"score": None, "explanation": f"LLM call failed: {e}", "examples": []}


# --- Fallback Conversion Helpers (module level) ---
async def _fallback_convert_pdf(file_path: Path) -> Dict[str, Any]:
    """Basic PDF conversion using PyPDF2."""
    _ocr_check_dep("PyPDF2", _PYPDF2_AVAILABLE, "Basic PDF Fallback Conversion")
    try:
        logger.info(f"Using PyPDF2 fallback for PDF: {file_path}")
        content = ""
        metadata: Dict[str, Any] = {"is_fallback": True}
        num_pages = 0
        if PyPDF2 is None:
            raise ImportError("PyPDF2 object is None despite _PYPDF2_AVAILABLE=True")
        with open(file_path, "rb") as f:
            try:
                reader = PyPDF2.PdfReader(f)
                num_pages = len(reader.pages)
                metadata["num_pages"] = num_pages
                pages = []
                for i in range(num_pages):
                    try:
                        page_text = reader.pages[i].extract_text() or ""
                        pages.append(page_text)
                    except Exception as page_err:
                        logger.warning(
                            f"PyPDF2 failed to extract text from page {i + 1}: {page_err}"
                        )
                        pages.append(f"[Page {i + 1} Extraction Error]")
                content = "\n\n".join(pages)
            except PyPDF2.errors.PdfReadError as pdf_err:
                logger.error(f"PyPDF2 could not read PDF {file_path}: {pdf_err}")
                raise ToolError(
                    "PDF_READ_ERROR", details={"library": "PyPDF2", "error": str(pdf_err)}
                ) from pdf_err
        metadata.update(_get_basic_metadata(content, num_pages))
        return {"content": content, "metadata": metadata}
    except Exception as e:
        logger.error(f"PyPDF2 fallback failed unexpectedly: {e}", exc_info=True)
        raise ToolError(
            "CONVERSION_FAILED",
            details={"file": str(file_path), "method": "PyPDF2 Fallback", "error": str(e)},
        ) from e


async def _fallback_convert_docx(file_path: Path) -> Dict[str, Any]:
    """Basic DOCX conversion using python-docx."""
    _ocr_check_dep("python-docx", _DOCX_AVAILABLE, "DOCX Fallback Conversion")
    try:
        logger.info(f"Using python-docx fallback for DOCX: {file_path}")
        if docx is None:
            raise ImportError("docx object is None despite _DOCX_AVAILABLE=True")
        doc = docx.Document(file_path)
        paragraphs = [para.text for para in doc.paragraphs if para.text]
        content = "\n\n".join(paragraphs)
        metadata: Dict[str, Any] = {
            "num_pages": 0,
            "has_tables": len(doc.tables) > 0,
            "has_figures": len(doc.inline_shapes) > 0,
            "has_sections": len(doc.sections) > 0,
            "is_fallback": True,
        }
        metadata.update(_get_basic_metadata(content))
        return {"content": content, "metadata": metadata}
    except Exception as e:
        logger.error(f"python-docx fallback failed: {e}", exc_info=True)
        raise ToolError(
            "CONVERSION_FAILED",
            details={"file": str(file_path), "method": "python-docx Fallback", "error": str(e)},
        ) from e


async def _fallback_convert_text(file_path: Path) -> Dict[str, Any]:
    """Simple text file reading."""
    try:
        logger.info(f"Reading text file directly: {file_path}")
        content = file_path.read_text(encoding="utf-8", errors="replace")
        line_count = content.count("\n") + 1
        page_estimate = max(1, int(line_count / 50))
        metadata = {"num_pages": page_estimate, "is_fallback": True}
        metadata.update(_get_basic_metadata(content, page_estimate))
        return {"content": content, "metadata": metadata}
    except Exception as e:
        logger.error(f"Text file reading failed: {e}", exc_info=True)
        raise ToolError(
            "CONVERSION_FAILED",
            details={"file": str(file_path), "method": "Direct Text Read", "error": str(e)},
        ) from e


###############################################################################
# Standalone Tool Functions (Exportable)                                      #
###############################################################################


# ------------------------ Document Conversion -----------------------------
@with_tool_metrics
@with_error_handling
async def convert_document(
    document_path: Optional[str] = None,
    document_data: Optional[bytes] = None,
    output_format: str = "markdown",
    extraction_strategy: str = DEFAULT_EXTRACTION_STRATEGY,
    enhance_with_llm: bool = True,
    ocr_options: Optional[Dict] = None,
    output_path: Optional[str] = None,
    save_to_file: bool = False,
    page_range: Optional[str] = None,
    section_filter: Optional[str] = None,
    accelerator_device: str = "auto",
    num_threads: int = 4,
) -> Dict[str, Any]:
    """
    Convert documents (PDF, Office formats, Images) to various formats (Standalone Function).
    (Args/Returns docs same as original class method)
    """
    t0 = time.time()
    strategy = extraction_strategy.lower()
    output_format = output_format.lower()
    ocr_options = ocr_options or {}

    # --- Input Validation ---
    if not document_path and not document_data:
        raise ToolInputError("Either 'document_path' or 'document_data' must be provided.")
    if document_path and document_data:
        raise ToolInputError("Provide either 'document_path' or 'document_data', not both.")
    if strategy not in _VALID_EXTRACTION_STRATEGIES:
        raise ToolInputError(
            f"Invalid extraction_strategy. Choose from: {', '.join(_VALID_EXTRACTION_STRATEGIES)}",
            param_name="extraction_strategy",
            provided_value=strategy,
        )
    if output_format not in _VALID_FORMATS:
        raise ToolInputError(
            f"Invalid output_format. Choose from: {', '.join(_VALID_FORMATS)}",
            param_name="output_format",
            provided_value=output_format,
        )

    # --- Dependency Checks based on strategy ---
    if strategy == "docling":
        _ocr_check_dep("docling", _DOCLING_AVAILABLE, "Docling extraction strategy")
    if strategy in ["direct_text", "hybrid_direct_ocr"]:
        if not (_PYMUPDF_AVAILABLE or _PDFPLUMBER_AVAILABLE):
            raise ToolError(
                "DEPENDENCY_MISSING",
                details={
                    "dependency": "PyMuPDF or PDFPlumber",
                    "feature": "Direct Text strategy",
                },
            )
    if strategy in ["ocr", "hybrid_direct_ocr"]:
        _ocr_check_dep("pdf2image", _PDF2IMAGE_AVAILABLE, "OCR strategy")
        _ocr_check_dep("pytesseract", _PYTESSERACT_AVAILABLE, "OCR strategy")
        _ocr_check_dep("Pillow", _PIL_AVAILABLE, "OCR strategy")
        if ocr_options.get("preprocessing") and not (_CV2_AVAILABLE and _NUMPY_AVAILABLE):
            logger.warning(
                "Preprocessing options provided but OpenCV/NumPy missing. Preprocessing limited."
            )

    # Adjust output format compatibility
    effective_output_format = output_format
    if strategy != "docling" and output_format not in _OCR_COMPATIBLE_FORMATS:
        logger.warning(
            f"Output format '{output_format}' is not directly supported by strategy '{strategy}'. Defaulting to 'markdown'."
        )
        effective_output_format = "markdown"

    # --- Prepare Input ---
    input_path_obj: Optional[Path] = None
    is_temp_file = False
    input_name: str = "input_data"

    try:
        input_path_obj, is_temp_file = _get_input_path_or_temp(document_path, document_data)
        input_name = input_path_obj.name

        with _handle_temp_file(input_path_obj, is_temp_file) as current_input_path:
            input_suffix = current_input_path.suffix.lower()
            is_pdf = input_suffix == ".pdf"
            is_image = input_suffix in [".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".gif"]
            is_office = input_suffix in [
                ".docx",
                ".pptx",
                ".xlsx",
                ".zip",
            ]  # Treat zip as potential office
            is_text = input_suffix in [".txt", ".md", ".html", ".xml", ".json"]

            # Validate strategy vs input type & adjust strategy if needed
            if not is_pdf and strategy in ["direct_text", "hybrid_direct_ocr"]:
                if is_image:
                    logger.warning(f"Strategy '{strategy}' needs PDF. Input is image. Using 'ocr'.")
                    strategy = "ocr"
                else:
                    raise ToolInputError(
                        f"Strategy '{strategy}' requires PDF input, got '{input_suffix}'."
                    )
            if not is_pdf and not is_image and strategy == "ocr":
                raise ToolInputError(
                    f"OCR strategy needs PDF/Image, got '{input_suffix}'. Use 'docling' or handle as text."
                )
            if is_office and strategy != "docling":
                if input_suffix == ".docx" and _DOCX_AVAILABLE:
                    logger.warning("Input is DOCX without 'docling'. Using fallback.")
                    strategy = "fallback_docx"
                # Add other office fallbacks here if needed
                else:
                    raise ToolInputError(
                        f"Office file ('{input_suffix}') requires 'docling' strategy or specific fallback library."
                    )
            if is_text and strategy != "docling":
                logger.info(f"Input is text ('{input_suffix}'). Using direct text handling.")
                strategy = "fallback_text"

            # --- Parse Page Range ---
            pages_to_process: Optional[List[int]] = None
            total_doc_pages = 0
            if page_range:
                try:
                    pages_set: Set[int] = set()
                    parts = page_range.split(",")
                    for part in parts:
                        part = part.strip()
                        if "-" in part:
                            start_str, end_str = part.split("-", 1)
                            start, end = int(start_str), int(end_str)
                            if start < 1 or end < start:
                                raise ValueError(f"Invalid range: {start}-{end}")
                            pages_set.update(range(start - 1, end))
                        else:
                            page_num = int(part)
                            if page_num < 1:
                                raise ValueError(f"Page number must be positive: {page_num}")
                            pages_set.add(page_num - 1)
                    if not pages_set:
                        raise ValueError("No valid pages selected.")
                    pages_to_process = sorted(list(pages_set))
                    logger.debug(
                        f"Parsed page range: {page_range} -> 0-based indices: {pages_to_process}"
                    )
                except ValueError as e:
                    raise ToolInputError(
                        f"Invalid page_range format: '{page_range}'. Error: {e}",
                        param_name="page_range",
                    ) from e

            # --- Result Structure Defaults ---
            result_content: Union[str, Dict] = ""
            doc_metadata: Dict[str, Any] = {}
            raw_text_pages: List[str] = []
            final_raw_text: Optional[str] = None
            quality_metrics: Optional[Dict] = None
            strategy_used = strategy

            # ======================== EXTRACTION STRATEGIES ========================

            if strategy == "docling":
                logger.info(f"Using 'docling' strategy for {input_name}")
                _ocr_check_dep("docling", _DOCLING_AVAILABLE, "Docling strategy")
                device_str = accelerator_device.lower()
                if device_str not in _ACCEL_MAP:
                    logger.warning(f"Invalid device '{device_str}', using 'auto'.")
                    device_str = "auto"
                device = _ACCEL_MAP[device_str]
                conv = _get_docling_converter(device, num_threads)
                loop = asyncio.get_running_loop()
                with _span("docling_conversion"):
                    docling_result = await loop.run_in_executor(
                        None, conv.convert, current_input_path
                    )
                if not docling_result or not docling_result.document:
                    raise ToolError("CONVERSION_FAILED", details={"reason": "Docling empty result"})
                doc_obj = docling_result.document
                doc_metadata = _get_docling_metadata(doc_obj)
                total_doc_pages = doc_metadata.get("num_pages", 0)

                if effective_output_format == "markdown":
                    result_content = doc_obj.export_to_markdown()
                elif effective_output_format == "text":
                    result_content = doc_obj.export_to_text()
                elif effective_output_format == "html":
                    result_content = doc_obj.export_to_html()
                elif effective_output_format == "json":
                    result_content = _json(doc_obj.export_to_dict())
                elif effective_output_format == "doctags":
                    result_content = doc_obj.export_to_doctags()
                else:
                    logger.warning(
                        f"Unsupported format '{effective_output_format}' for Docling, using markdown."
                    )
                    result_content = doc_obj.export_to_markdown()
                    effective_output_format = "markdown"

                if save_to_file:
                    fp = (
                        Path(output_path)
                        if output_path
                        else _tmp_path(str(current_input_path), effective_output_format)
                    )
                    fp.parent.mkdir(parents=True, exist_ok=True)
                    img_mode = (
                        _ImageRefModeType.PLACEHOLDER
                        if effective_output_format in ["text", "json"]
                        else _ImageRefModeType.REFERENCED
                    )
                    save_func_map = {
                        "markdown": functools.partial(
                            doc_obj.save_as_markdown, image_mode=img_mode
                        ),
                        "text": functools.partial(doc_obj.save_as_markdown(strict_text=True)),
                        "html": functools.partial(doc_obj.save_as_html, image_mode=img_mode),
                        "json": functools.partial(doc_obj.save_as_json, image_mode=img_mode),
                        "doctags": functools.partial(doc_obj.save_as_doctags),
                    }
                    save_func = save_func_map.get(effective_output_format)
                    if save_func and callable(save_func):
                        with _span(f"docling_save_{effective_output_format}"):
                            save_func(fp)
                        logger.info(f"Saved Docling output ({effective_output_format}) to {fp}")
                        doc_metadata["saved_output_path"] = str(fp)
                    else:
                        fp.write_text(str(result_content), encoding="utf-8")
                        logger.info(f"Saved Docling output (generic text write) to {fp}")
                        doc_metadata["saved_output_path"] = str(fp)

            elif strategy.startswith("fallback_"):
                fallback_type = strategy.split("_", 1)[1]
                logger.info(f"Using fallback strategy for: {fallback_type}")
                fallback_result: Optional[Dict[str, Any]] = None
                if fallback_type == "docx":
                    fallback_result = await _fallback_convert_docx(current_input_path)
                elif fallback_type == "pdf":
                    fallback_result = await _fallback_convert_pdf(current_input_path)
                elif fallback_type == "text":
                    fallback_result = await _fallback_convert_text(current_input_path)
                if fallback_result:
                    raw_text_pages = [fallback_result.get("content", "")]
                    doc_metadata = fallback_result.get("metadata", {})
                    total_doc_pages = doc_metadata.get("num_pages", 1)
                    strategy_used = f"fallback_{fallback_type}"
                else:
                    raise ToolError(
                        "CONVERSION_FAILED",
                        details={"reason": f"Fallback '{fallback_type}' failed."},
                    )

            else:  # Text/OCR strategies
                run_ocr = False
                run_direct = False
                if strategy == "direct_text":
                    run_direct = True
                elif strategy == "ocr":
                    run_ocr = True
                elif strategy == "hybrid_direct_ocr":
                    if not is_pdf:
                        run_ocr = True
                        strategy_used = "ocr"
                        logger.info("Input is image, using 'ocr'.")
                    else:
                        run_direct = True
                extract_start_page = pages_to_process[0] if pages_to_process else 0
                extract_page_count = len(pages_to_process) if pages_to_process else 0

                if run_direct:
                    logger.info(f"Attempting 'direct_text' strategy for {input_name}")
                    try:
                        with _span("direct_text_extraction"):
                            (
                                extracted_pages,
                                has_meaningful_text,
                            ) = await asyncio.to_thread(  # Use helper defined above
                                _ocr_extract_text_from_pdf_direct,
                                current_input_path,
                                start_page=extract_start_page,
                                max_pages=extract_page_count,
                            )
                        total_doc_pages = len(
                            extracted_pages
                        )  # Page count reflects extracted range
                        if strategy == "hybrid_direct_ocr" and not has_meaningful_text:
                            logger.warning("Direct text minimal. Falling back to OCR.")
                            run_ocr = True
                            strategy_used = "ocr"
                        elif not has_meaningful_text and strategy == "direct_text":
                            raise ToolError(
                                "DIRECT_EXTRACTION_FAILED",
                                details={"reason": "No meaningful text found."},
                            )
                        else:
                            raw_text_pages = extracted_pages
                            logger.info(f"Direct text success: {len(raw_text_pages)} pages.")
                    except ToolError as e:
                        if strategy == "hybrid_direct_ocr":
                            logger.warning(f"Direct failed ({e.error_code}). Falling back to OCR.")
                            run_ocr = True
                            strategy_used = "ocr"
                        else:
                            raise e
                    except Exception as e_direct:
                        logger.error(f"Unexpected direct text error: {e_direct}", exc_info=True)
                        if strategy == "hybrid_direct_ocr":
                            logger.warning("Direct failed. Falling back to OCR.")
                            run_ocr = True
                            strategy_used = "ocr"
                        else:
                            raise ToolError(
                                "DIRECT_EXTRACTION_FAILED", details={"error": str(e_direct)}
                            ) from e_direct

                if run_ocr:
                    logger.info(f"Using 'ocr' strategy for {input_name}")
                    strategy_used = "ocr"
                    ocr_lang = ocr_options.get("language", "eng")
                    ocr_dpi = ocr_options.get("dpi", 300)
                    ocr_prep_opts = ocr_options.get("preprocessing")
                    images: List["PILImage.Image"] = []
                    if is_pdf:
                        convert_func = _ocr_convert_pdf_to_images  # Use helper defined above
                        with _span("pdf_to_images"):
                            images = await asyncio.to_thread(
                                convert_func,
                                current_input_path,
                                start_page=extract_start_page,
                                max_pages=extract_page_count,
                                dpi=ocr_dpi,
                            )
                        total_doc_pages = len(images)
                    elif is_image:
                        _ocr_check_dep("Pillow", _PIL_AVAILABLE, "Image loading")
                        if Image is None:
                            raise ToolError(
                                "INTERNAL_ERROR", details={"reason": "PIL.Image is None"}
                            )
                        with _span(f"load_image_{input_name}"):
                            img = Image.open(current_input_path)  # type: ignore
                        images = [img.convert("RGB")]
                        total_doc_pages = 1
                        img.close()  # Close after converting
                    if not images:
                        raise ToolError("OCR_FAILED", details={"reason": "No images for OCR."})

                    processed_pages_text: List[str] = [""] * len(images)

                    async def _process_ocr_page_worker(
                        idx: int, img: "PILImage.Image"
                    ) -> Tuple[int, str]:
                        try:
                            loop = asyncio.get_running_loop()
                            with _span(f"ocr_page_{idx}_preprocess"):
                                prep_img = await loop.run_in_executor(
                                    None, _ocr_preprocess_image, img, ocr_prep_opts
                                )
                            with _span(f"ocr_page_{idx}_tesseract"):
                                text = await loop.run_in_executor(
                                    None,
                                    _ocr_run_tesseract,
                                    prep_img,
                                    ocr_lang,
                                    ocr_options.get("tesseract_config", ""),
                                )  # Use helper defined above
                            if prep_img != img:
                                prep_img.close()  # Close preprocessed image if different
                            return idx, text
                        except Exception as page_err:
                            logger.error(
                                f"OCR page {idx + extract_start_page} error: {page_err}",
                                exc_info=True,
                            )
                            return idx, f"[Page {idx + extract_start_page + 1} OCR Error]"
                        finally:
                            img.close()  # Close the original image passed to worker

                    tasks = [_process_ocr_page_worker(i, img) for i, img in enumerate(images)]
                    page_results = await asyncio.gather(*tasks)
                    for idx, text in page_results:
                        processed_pages_text[idx] = text
                    raw_text_pages = processed_pages_text
                    logger.info(f"OCR extraction successful for {len(raw_text_pages)} pages.")

            # --- Stage 2 & 3 (Post-processing for non-Docling) ---
            if strategy != "docling":
                if not raw_text_pages:
                    raise ToolError(
                        "EXTRACTION_FAILED",
                        details={"reason": f"Strategy '{strategy_used}' yielded no text."},
                    )
                final_raw_text = "\n\n".join(raw_text_pages).strip()
                if section_filter and final_raw_text:
                    try:
                        pat = re.compile(section_filter, re.I | re.M)
                        blocks = re.split(r"(\n\s*\n)", final_raw_text)
                        kept_content = ""
                        for i in range(0, len(blocks), 2):
                            block = blocks[i]
                            separator = blocks[i + 1] if i + 1 < len(blocks) else ""
                        if block and pat.search(block):
                            kept_content += block + separator
                        final_raw_text = kept_content.strip()
                        if not final_raw_text:
                            logger.warning(
                                f"Section filter '{section_filter}' removed all content."
                            )
                        else:
                            logger.info(f"Applied section filter: '{section_filter}'")
                    except Exception as e_filter:
                        logger.warning(f"Failed to apply section filter: {e_filter}")

                if enhance_with_llm and final_raw_text:
                    logger.info("Applying LLM enhancement.")
                    result_content = ""
                    with _span("llm_text_enhancement"):
                        chunks = _ocr_split_text_into_chunks(
                            final_raw_text
                        )  # Use helper defined above
                        if chunks:
                            enhancement_tasks = [
                                _ocr_enhance_text_chunk(
                                    chunk,
                                    output_format=effective_output_format,
                                    remove_headers=ocr_options.get("remove_headers", False),
                                )
                                for chunk in chunks
                            ]  # Use helper defined above
                            enhanced_chunks = await asyncio.gather(*enhancement_tasks)
                            result_content = "\n\n".join(enhanced_chunks).strip()
                        else:
                            logger.warning("Text empty pre-LLM.")
                else:
                    result_content = final_raw_text or ""
                if not doc_metadata or doc_metadata.get("is_fallback"):
                    doc_metadata = _get_basic_metadata(str(result_content), total_doc_pages)
                if enhance_with_llm and final_raw_text and ocr_options.get("assess_quality", False):
                    logger.info("Performing OCR quality assessment.")
                    with _span("ocr_quality_assessment"):
                        quality_metrics = await _ocr_assess_text_quality(
                            final_raw_text, str(result_content)
                        )  # Use helper defined above

            # ======================== POST-PROCESSING & RETURN ========================
            final_content = result_content
            if save_to_file and strategy != "docling":  # Docling saving handled above
                fp = (
                    Path(output_path)
                    if output_path
                    else _tmp_path(input_name, effective_output_format)
                )
                fp.parent.mkdir(parents=True, exist_ok=True)
                try:
                    content_to_save = (
                        _json(final_content)
                        if isinstance(final_content, dict)
                        else str(final_content)
                    )
                    fp.write_text(content_to_save, encoding="utf-8")
                    logger.info(
                        f"Saved output ({effective_output_format}, strategy: {strategy_used}) to {fp}"
                    )
                    doc_metadata["saved_output_path"] = str(fp)
                except Exception as e_save:
                    logger.error(f"Failed to save output file to {fp}: {e_save}", exc_info=True)
                    doc_metadata["save_error"] = f"Failed to save: {e_save}"

            elapsed = round(time.time() - t0, 3)
            response: Dict[str, Any] = {
                "success": True,
                "content": final_content,
                "output_format": effective_output_format,
                "processing_time": elapsed,
                "document_metadata": doc_metadata,
                "extraction_strategy_used": strategy_used,
            }
            if final_raw_text is not None and strategy != "docling":
                response["raw_text"] = final_raw_text
            if quality_metrics is not None:
                response["ocr_quality_metrics"] = quality_metrics
            if "saved_output_path" in doc_metadata:
                response["file_path"] = doc_metadata["saved_output_path"]
            logger.info(
                f"Completed conversion '{input_name}' -> {effective_output_format} (strategy: {strategy_used}) in {elapsed}s"
            )
            return response
    except Exception as e:
        logger.error(f"Error in convert_document for '{input_name}': {e}", exc_info=True)
        if isinstance(e, (ToolInputError, ToolError)):
            raise e
        raise ToolError("CONVERSION_FAILED", details={"input": input_name, "error": str(e)}) from e


# <<< Part 1 code goes here >>>

###############################################################################
# Chunking Helpers (Internal)                                                 #
###############################################################################


async def _internal_token_chunks(doc: str, size: int, overlap: int) -> List[str]:
    """Chunk document by tokens, respecting sentence boundaries (Internal Helper)."""
    enc = _get_tiktoken_encoder()
    if not enc:
        logger.warning("Tiktoken not available, falling back to character chunking.")
        char_size = size * 4
        char_overlap = overlap * 4
        return await _internal_char_chunks(doc, char_size, char_overlap)
    if not doc:
        return []
    try:
        tokens = enc.encode(doc, disallowed_special=())
    except Exception as e:
        logger.error(f"Tiktoken encoding failed: {e}. Falling back to char.", exc_info=True)
        return await _internal_char_chunks(doc, size * 4, overlap * 4)
    if not tokens:
        return []
    chunks: List[str] = []
    current_pos = 0
    n_tokens = len(tokens)
    try:
        sentence_end_tokens = {enc.encode(p)[0] for p in (".", "?", "!", "\n")}
    except Exception as e:
        encoding_name = getattr(enc, "name", "unknown")
        if encoding_name == "cl100k_base":
            sentence_end_tokens = {13, 30, 106, 198}
        else:
            try:
                sentence_end_tokens = {enc.encode("\n")[0]}
            except Exception:  # If even newline encoding fails
                logger.error(
                    f"Cannot encode even newline token for encoding '{encoding_name}'. Using empty set for sentence ends."
                )
                sentence_end_tokens = set()
        logger.warning(
            f"Could not encode sentence ends: {e}. Using fallback tokens: {sentence_end_tokens}"
        )
    while current_pos < n_tokens:
        end_pos = min(current_pos + size, n_tokens)
        best_split_pos = end_pos
        if end_pos < n_tokens:
            lookback_distance = min(overlap, size // 4, end_pos - current_pos)
            search_start = max(current_pos, end_pos - lookback_distance)
            for k in range(end_pos - 1, search_start - 1, -1):
                if tokens[k] in sentence_end_tokens:
                    best_split_pos = k + 1
                    break
        chunk_token_ids = tokens[current_pos:best_split_pos]
        if not chunk_token_ids:
            if current_pos >= n_tokens:
                break
            current_pos += 1
            continue
        try:
            chunk_text = enc.decode(chunk_token_ids).strip()
            if chunk_text:
                chunks.append(chunk_text)
        except Exception as decode_err:  # Keep variable for logging
            logger.error(
                f"Tiktoken decode failed for {current_pos}:{best_split_pos}: {decode_err}",
                exc_info=False,
            )
        next_start_pos = best_split_pos - overlap
        current_pos = max(current_pos + 1, next_start_pos)
        if current_pos <= best_split_pos - size:
            current_pos = best_split_pos
    return chunks


async def _internal_char_chunks(doc: str, size: int, overlap: int) -> List[str]:
    """Chunk document by characters, respecting sentence/paragraph boundaries (Internal Helper)."""
    if not doc:
        return []
    chunks: List[str] = []
    current_pos = 0
    n_chars = len(doc)
    sentence_ends = (". ", "? ", "! ", "\n\n")
    softer_breaks = ("\n", "; ", ": ", ", ", ".)", "?)", "!)", "\t", " ")
    while current_pos < n_chars:
        end_pos = min(current_pos + size, n_chars)
        best_split_pos = end_pos
        if end_pos < n_chars:
            lookback_window_start = max(current_pos, end_pos - int(size * 0.2), end_pos - 150)
            best_found_pos = -1
            for marker in sentence_ends:
                found_pos = doc.rfind(marker, lookback_window_start, end_pos)
                if found_pos != -1:
                    best_found_pos = max(best_found_pos, found_pos + len(marker))
            if best_found_pos == -1:
                for marker in softer_breaks:
                    found_pos = doc.rfind(marker, lookback_window_start, end_pos)
                    if found_pos != -1:
                        best_found_pos = max(best_found_pos, found_pos + len(marker))
            if best_found_pos > current_pos:
                best_split_pos = best_found_pos
        actual_chunk_text = doc[current_pos:best_split_pos].strip()
        if actual_chunk_text:
            chunks.append(actual_chunk_text)
        next_start_pos = best_split_pos - overlap
        current_pos = max(current_pos + 1, next_start_pos)
        if current_pos <= best_split_pos - size:
            current_pos = best_split_pos
    return chunks


async def _internal_paragraph_chunks(doc: str, size: int, overlap: int) -> List[str]:
    """Chunk document by paragraphs, combining small ones (Internal Helper)."""
    if not doc:
        return []
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n+", doc) if p.strip()]
    if not paragraphs:
        return []
    chunks = []
    current_chunk_paragraphs: List[str] = []
    current_chunk_len = 0

    def get_len(text: str) -> int:
        return len(text)

    def is_markdown_table(text: str) -> bool:
        lines = text.strip().split("\n")
        return (
            len(lines) >= 2
            and all(line.strip().startswith("|") for line in lines[:2])
            and "|" in lines[0]
            and re.search(r"\|.*?(-{3,}|:{1,2}-{1,}:?).*?\|", lines[1]) is not None
        )

    for p in paragraphs:
        p_len = get_len(p)
        potential_new_len = (
            current_chunk_len + (get_len("\n\n") if current_chunk_paragraphs else 0) + p_len
        )
        is_table = is_markdown_table(p)
        if current_chunk_paragraphs and potential_new_len > size and not is_table:
            chunks.append("\n\n".join(current_chunk_paragraphs))
            current_chunk_paragraphs = [p]
            current_chunk_len = p_len
        elif p_len > size and not is_table:
            logger.warning(
                f"Paragraph (len {p_len}) starting '{p[:50]}...' exceeds size {size}. Splitting."
            )
            if current_chunk_paragraphs:
                chunks.append("\n\n".join(current_chunk_paragraphs))
            sub_chunks = await _internal_char_chunks(p, size, overlap)
            chunks.extend(sub_chunks)
            current_chunk_paragraphs = []
            current_chunk_len = 0
        else:
            current_chunk_paragraphs.append(p)
            current_chunk_len = potential_new_len
    if current_chunk_paragraphs:
        chunks.append("\n\n".join(current_chunk_paragraphs))
    logger.info(f"Chunked into {len(chunks)} paragraphs/groups.")
    return chunks


async def _internal_section_chunks(doc: str, size: int, overlap: int) -> List[str]:
    """Chunk document by identified sections (Internal Helper). Falls back to paragraphs."""
    try:
        # Call main tool function (defined in Part 3)
        section_result = await identify_sections(document=doc)
        if (
            isinstance(section_result, dict)
            and section_result.get("success")
            and isinstance(section_result.get("sections"), list)
        ):
            sections = section_result["sections"]
        else:
            logger.warning(
                "identify_sections failed/unexpected format. Falling back to paragraphs."
            )
            return await _internal_paragraph_chunks(doc, size, overlap)
        if not sections:
            logger.info("No sections identified, using paragraph fallback.")
            return await _internal_paragraph_chunks(doc, size, overlap)
        section_texts: List[str] = []
        for s in sections:
            title = s.get("title", "").strip()
            text = s.get("text", "").strip()
            if text:
                use_title = title and title.lower() not in [
                    "introduction",
                    "main content",
                    "body",
                    "abstract",
                    "summary",
                ]
                full_section_text = f"## {title}\n\n{text}" if use_title else text
                section_texts.append(full_section_text.strip())

        def contains_markdown_table(text: str) -> bool:
            lines = text.strip().split("\n")
            return (
                len(lines) >= 2
                and all(line.strip().startswith("|") for line in lines[:2])
                and "|" in lines[0]
                and re.search(r"\|.*?(-{3,}|:{1,2}-{1,}:?).*?\|", lines[1]) is not None
            )

        final_chunks = []
        for text in section_texts:
            text_len = len(text)
            has_table = contains_markdown_table(text)
            should_split = text_len > size * 1.1 and (not has_table or text_len > size * 2)
            if should_split:
                logger.warning(
                    f"Section chunk (len {text_len}) starting '{text[:50]}...' exceeds size {size}. Sub-chunking."
                )
                sub_chunks = await _internal_paragraph_chunks(text, size, overlap)
                final_chunks.extend(sub_chunks)
            elif text:
                final_chunks.append(text)
        return final_chunks
    except Exception as e:
        logger.error(f"Section chunking failed: {e}. Falling back to paragraphs.", exc_info=True)
        return await _internal_paragraph_chunks(doc, size, overlap)


###############################################################################
# Standalone Tool Functions (Continued)                                       #
###############################################################################


# ------------------------ Chunking Tool Function (Merged) --------------------
@with_tool_metrics
@with_error_handling
async def chunk_document(
    document: str,
    *,
    chunk_size: int = 1000,
    chunk_method: str = "paragraph",
    chunk_overlap: int = 0,
    chunk_strategy: Optional[str] = None,  # Keep alias for compatibility
) -> Dict[str, Any]:
    """
    Split document text into chunks using various strategies (Standalone Tool Function).

    Args:
        document: Text content to chunk.
        chunk_size: Target maximum size of each chunk (meaning depends on method: tokens or characters).
        chunk_method: Chunking method ('token', 'character', 'section', 'paragraph').
        chunk_overlap: Number of tokens/characters to overlap between chunks (for token/char methods).
                       Overlap logic for paragraph/section is heuristic/simplified.
        chunk_strategy: Alias for chunk_method (for backward compatibility).

    Returns:
        Dictionary containing list of chunked text strings.
        Example: {"chunks": ["chunk 1 text...", "chunk 2 text..."], "success": True}
    """
    # Use module logger directly
    _logger = logger

    if not document or not isinstance(document, str):
        _logger.warning("Chunking called with empty or invalid document input.")
        return {"chunks": [], "success": True}

    size = max(100, int(chunk_size))
    overlap = max(0, min(int(chunk_overlap), size // 3))
    method = (chunk_strategy or chunk_method or "paragraph").lower()

    # Map to internal helpers
    chunker_map = {
        "token": _internal_token_chunks,
        "character": _internal_char_chunks,
        "section": _internal_section_chunks,  # Relies on identify_sections tool
        "paragraph": _internal_paragraph_chunks,
    }

    strat_func = chunker_map.get(method)
    if not strat_func:
        _logger.warning(f"Unknown chunk_method '{method}'. Defaulting to 'paragraph'.")
        strat_func = _internal_paragraph_chunks
        method = "paragraph"

    _logger.info(f"Chunking document using method='{method}', size={size}, overlap={overlap}")
    chunks: List[str] = []

    try:
        t0_chunk = time.time()
        chunks = await strat_func(document, size, overlap)
        elapsed_chunk = time.time() - t0_chunk
        _logger.info(f"Chunking completed in {elapsed_chunk:.3f}s")
    except Exception as e:
        _logger.error(f"Error during chunking operation ({method}): {e}", exc_info=True)
        raise ToolError("CHUNKING_FAILED", details={"method": method, "error": str(e)}) from e

    final_chunks = [c for c in chunks if isinstance(c, str) and c]
    _logger.info(f"Generated {len(final_chunks)} chunks.")
    return {"chunks": final_chunks, "success": True}


###############################################################################
# HTML Processing Helpers & Tools                                             #
###############################################################################


# --- HTML Extraction Helpers ---
def _extract_readability(html_txt: str) -> str:
    """Extract main content using readability-lxml (Standalone)."""
    if not _READABILITY_AVAILABLE or not readability:
        logger.warning("Readability-lxml not installed. Cannot use readability extraction.")
        return ""
    try:
        # Adjust readability settings for better extraction
        # Use setdefault to avoid modifying the original regexes if called multiple times
        # Ensure the default regexes exist before modifying
        default_unlikely = readability.htmls.DEFAULT_REGEXES.get(
            "unlikelyCandidates", re.compile(r"$^")
        )  # Default to matching nothing
        readability.htmls.DEFAULT_REGEXES["unlikelyCandidates"] = re.compile(
            default_unlikely.pattern
            + "|aside|footer|nav|sidebar|footnote|advertisement|related|recommend|share|social|comment|meta",
            re.I,
        )

        default_positive = readability.htmls.DEFAULT_REGEXES.get("positive", re.compile(r"$^"))
        readability.htmls.DEFAULT_REGEXES["positive"] = re.compile(
            default_positive.pattern + "|article|main|content|post|entry|body", re.I
        )

        default_negative = readability.htmls.DEFAULT_REGEXES.get("negative", re.compile(r"$^"))
        readability.htmls.DEFAULT_REGEXES["negative"] = re.compile(
            default_negative.pattern + "|widget|menu|legal|promo|disclaimer", re.I
        )

        doc = readability.Document(html_txt)
        summary_html = doc.summary(html_partial=True)
        return summary_html
    except Exception as e:
        logger.warning(f"Readability extraction failed: {e}", exc_info=True)
        return ""


def _extract_trafilatura(html_txt: str) -> str:
    """Extract main content using trafilatura (Standalone)."""
    if not _TRAFILATURA_AVAILABLE or not trafilatura:
        logger.warning("Trafilatura not installed. Cannot use trafilatura extraction.")
        return ""
    try:
        extracted = trafilatura.extract(
            html_txt,
            include_comments=False,
            include_tables=True,
            favor_precision=True,
            deduplicate=True,
            target_language=None,
            include_formatting=True,
            output_format="html",
        )
        return extracted or ""
    except Exception as e:
        logger.warning(f"Trafilatura extraction failed: {e}", exc_info=True)
        return ""


# --- HTML Processing Tools ---


@with_tool_metrics
@with_error_handling
async def clean_and_format_text_as_markdown(
    text: str,
    force_markdown_conversion: bool = False,
    extraction_method: str = "auto",  # auto, readability, trafilatura, none
    preserve_tables: bool = True,
    preserve_links: bool = True,
    preserve_images: bool = False,
    max_line_length: int = 0,  # 0 means no wrapping
) -> Dict[str, Any]:
    """
    Convert plain text or HTML to clean Markdown, optionally extracting main content (Standalone Tool).
    """
    start_time = time.time()
    if not text or not isinstance(text, str):
        raise ToolInputError("Input text must be a non-empty string", param_name="text")

    # Detect content type (function defined in Part 3)
    content_type_result = await detect_content_type(text)
    input_type = content_type_result.get("content_type", "unknown")
    input_confidence = content_type_result.get("confidence", 0.0)

    was_html = (input_type == "html" and input_confidence > 0.3) or (
        input_type != "markdown" and input_type != "code" and _is_html_fragment(text)
    )

    extraction_method_used = "none"
    processed_text = text

    logger.debug(f"Input content type detected as: {input_type}, treating as HTML: {was_html}")

    if was_html or force_markdown_conversion:
        was_html = True
        actual_extraction = extraction_method.lower()
        if actual_extraction == "auto":
            if _TRAFILATURA_AVAILABLE:
                actual_extraction = "trafilatura"
            elif _READABILITY_AVAILABLE:
                actual_extraction = "readability"
            else:
                actual_extraction = "none"
            logger.debug(f"Auto-selected extraction method: {actual_extraction}")

        extraction_method_used = actual_extraction

        if actual_extraction != "none":
            extracted_html = ""
            logger.info(f"Attempting HTML content extraction using: {actual_extraction}")
            try:
                if actual_extraction == "readability":
                    extracted_html = _extract_readability(processed_text)
                elif actual_extraction == "trafilatura":
                    extracted_html = _extract_trafilatura(processed_text)

                if extracted_html and len(extracted_html.strip()) > 50:
                    processed_text = extracted_html
                    logger.info(f"Successfully extracted content using {actual_extraction}")
                else:
                    logger.warning(
                        f"{actual_extraction.capitalize()} extraction yielded minimal content. Using original."
                    )
                    extraction_method_used = f"{actual_extraction} (failed)"
            except Exception as e_extract:
                logger.error(
                    f"Error during {actual_extraction} extraction: {e_extract}", exc_info=True
                )
                extraction_method_used = f"{actual_extraction} (error)"

        try:
            logger.info(
                f"Converting HTML (extracted: {extraction_method_used != 'none'}) to Markdown..."
            )
            md_text = _html_to_md_core(
                processed_text,
                links=preserve_links,
                imgs=preserve_images,
                tbls=preserve_tables,
                width=0,  # Disable html2text wrapping here
            )
            md_text = _sanitize(md_text)
            md_text = _improve(md_text)
            processed_text = md_text
        except Exception as e_conv:
            logger.error(f"Error converting HTML to Markdown: {e_conv}", exc_info=True)
            processed_text = _sanitize(processed_text)
            logger.warning("HTML to Markdown conversion failed, returning sanitized input.")

    elif input_type == "markdown" and not force_markdown_conversion:
        logger.debug("Input detected as Markdown, applying cleanup.")
        processed_text = _sanitize(text)
        processed_text = _improve(processed_text)

    elif input_type == "text" or input_type == "unknown":
        logger.debug(f"Input detected as {input_type}, applying basic text formatting.")
        processed_text = re.sub(r"\n{2,}", "<TEMP_PARA_BREAK>", text)
        processed_text = re.sub(r"\n", " ", processed_text)
        processed_text = processed_text.replace("<TEMP_PARA_BREAK>", "\n\n")
        processed_text = _sanitize(processed_text)
        processed_text = _improve(processed_text)

    # Apply line wrapping if requested (using textwrap module)
    if max_line_length > 0:
        try:
            wrapped_lines = []
            current_block = ""
            in_code_block = False
            for line in processed_text.split("\n"):
                line_stripped = line.strip()
                if line_stripped.startswith("```"):
                    in_code_block = not in_code_block
                    if current_block:
                        wrapped_lines.extend(
                            textwrap.wrap(
                                current_block.strip(),
                                width=max_line_length,
                                break_long_words=False,
                                break_on_hyphens=False,
                            )
                        )
                    current_block = ""
                    wrapped_lines.append(line)
                elif (
                    in_code_block
                    or line_stripped.startswith(("#", ">", "- ", "* ", "+ "))
                    or re.match(r"^\d+\.\s", line_stripped)
                    or line_stripped == ""
                    or line_stripped.startswith("|")
                ):
                    if current_block:
                        wrapped_lines.extend(
                            textwrap.wrap(
                                current_block.strip(),
                                width=max_line_length,
                                break_long_words=False,
                                break_on_hyphens=False,
                            )
                        )
                    current_block = ""
                    wrapped_lines.append(line)
                else:
                    current_block += line + " "
            if current_block:
                wrapped_lines.extend(
                    textwrap.wrap(
                        current_block.strip(),
                        width=max_line_length,
                        break_long_words=False,
                        break_on_hyphens=False,
                    )
                )
            processed_text = "\n".join(wrapped_lines)
        except Exception as e_wrap:
            logger.error(f"Error during line wrapping: {e_wrap}")

    processing_time = time.time() - start_time
    return {
        "success": True,
        "markdown_text": processed_text.strip(),
        "original_content_type": input_type,
        "was_html": was_html,
        "extraction_method_used": extraction_method_used,
        "processing_time": processing_time,
    }


@with_tool_metrics
@with_error_handling
async def detect_content_type(text: str) -> Dict[str, Any]:
    """
    Detect if text is primarily HTML, Markdown, code, or plain text (Standalone Tool).
    """
    t0 = time.time()
    if not text or not isinstance(text, str):
        raise ToolInputError("Input text must be a non-empty string", param_name="text")

    sample_size = 4000
    if len(text) > sample_size * 2:
        sample = text[:sample_size] + "\n" + text[-sample_size:]
    else:
        sample = text

    scores = {"html": 0.0, "markdown": 0.0, "code": 0.0, "text": 1.0}
    detection_criteria: Dict[str, List[str]] = {"html": [], "markdown": [], "code": [], "text": []}
    max_score = 0.0

    for type_name, patterns in _CONTENT_PATTERNS.items():
        type_score = 0.0
        for pattern, weight in patterns:
            matches = pattern.findall(sample)
            if matches:
                type_score += weight * 0.2
                density_score = min(1.0, len(matches) / 10.0)
                type_score += weight * density_score * 0.8
                detection_criteria[type_name].append(
                    f"Pattern matched ({len(matches)}x): {pattern.pattern[:50]}..."
                )

        scores[type_name] = min(scores[type_name] + type_score, 5.0)
        max_score = max(max_score, scores[type_name])

    if scores["html"] > 0.1 and any(
        p[0].pattern in ["<html", "<head", "<body", "<!DOCTYPE"] for p in _CONTENT_PATTERNS["html"]
    ):
        scores["html"] *= 1.5

    if max_score < 0.5:
        if (
            len(re.findall(r"\b(the|a|is|was|in|on|at)\b", sample, re.I)) > 10
            and len(re.findall(r"[.?!]\s", sample)) > 3
        ):
            detection_criteria["text"].append("Natural language indicators found")
            scores["text"] += 0.5

    if max_score > 0.2:
        scores["text"] *= 0.8

    primary_type = max(scores, key=lambda k: scores[k])

    confidence = min(1.0, scores[primary_type] / max(1.0, max_score * 0.8))
    sorted_scores = sorted(scores.values(), reverse=True)
    if len(sorted_scores) > 1 and sorted_scores[0] > sorted_scores[1] * 2:
        confidence = min(1.0, confidence * 1.2)
    confidence = min(0.95, confidence) if scores[primary_type] < 3.0 else confidence
    confidence = min(1.0, confidence)

    processing_time = time.time() - t0
    return {
        "success": True,
        "content_type": primary_type,
        "confidence": round(confidence, 3),
        "detection_criteria": detection_criteria[primary_type],
        "all_scores": {k: round(v, 2) for k, v in scores.items()},
        "processing_time": round(processing_time, 3),
    }


@with_tool_metrics
@with_error_handling
async def batch_format_texts(
    texts: List[str],
    force_markdown_conversion: bool = False,
    extraction_method: str = "auto",
    max_concurrency: int = 5,
    preserve_tables: bool = True,
    preserve_links: bool = True,
    preserve_images: bool = False,
) -> Dict[str, Any]:
    """Applies 'clean_and_format_text_as_markdown' to texts concurrently (Standalone Tool)."""
    if not texts or not isinstance(texts, list):
        raise ToolInputError("Input must be a non-empty list", param_name="texts")
    if not all(isinstance(t, str) for t in texts):
        raise ToolInputError("All items in 'texts' list must be strings", param_name="texts")
    max_concurrency = max(1, max_concurrency)
    sem = asyncio.Semaphore(max_concurrency)
    tasks = []

    async def _process_one_standalone(idx: int, txt: str):
        async with sem:
            logger.debug(f"Starting batch formatting for text index {idx}")
            result_dict = {"original_index": idx}
            try:
                res = await clean_and_format_text_as_markdown(
                    text=txt,
                    force_markdown_conversion=force_markdown_conversion,
                    extraction_method=extraction_method,
                    preserve_tables=preserve_tables,
                    preserve_links=preserve_links,
                    preserve_images=preserve_images,
                    max_line_length=0,
                )
                result_dict.update(res)
                result_dict["success"] = bool(res.get("success", False))
                if result_dict["success"]:
                    logger.debug(f"Successfully batch formatted text index {idx}")
            except ToolInputError as e_input:
                logger.warning(f"Input error formatting text index {idx}: {e_input}")
                result_dict.update(
                    {
                        "error": str(e_input),
                        "success": False,
                        "error_type": "ToolInputError",
                        "error_code": e_input.error_code,
                    }
                )
            except ToolError as e_tool:
                logger.warning(
                    f"Processing error formatting text index {idx}: {e_tool.error_code} - {str(e_tool)}"
                )
                result_dict.update(
                    {
                        "error": str(e_tool),
                        "success": False,
                        "error_code": e_tool.error_code,
                        "error_type": "ToolError",
                    }
                )
            except Exception as e:
                logger.error(f"Unexpected error formatting text index {idx}: {e}", exc_info=True)
                result_dict.update({"error": str(e), "success": False, "error_type": "Exception"})
            return result_dict

    tic = time.perf_counter()
    logger.info(
        f"Starting batch formatting for {len(texts)} texts with concurrency {max_concurrency}..."
    )
    for i, t in enumerate(texts):
        tasks.append(_process_one_standalone(i, t))
    all_results = await asyncio.gather(*tasks)
    toc = time.perf_counter()
    logger.info(f"Batch formatting completed in {toc - tic:.3f}s")

    all_results.sort(key=lambda r: r.get("original_index", -1))
    final_results = []
    success_count = 0
    failure_count = 0
    for r in all_results:
        if r.get("success"):
            success_count += 1
        else:
            failure_count += 1
        r.pop("original_index", None)
        final_results.append(r)
    return {
        "results": final_results,
        "total_processing_time": round(toc - tic, 3),
        "success_count": success_count,
        "failure_count": failure_count,
        "success": True,
    }


@with_tool_metrics
@with_error_handling
async def optimize_markdown_formatting(
    markdown: str,
    normalize_headings: bool = False,
    fix_lists: bool = True,
    fix_links: bool = True,
    add_line_breaks: bool = True,
    compact_mode: bool = False,
    max_line_length: int = 0,
) -> Dict[str, Any]:
    """
    Clean up and standardize existing Markdown text (Standalone Tool).
    Note: `normalize_headings` is currently a basic implementation.
    """
    t0 = time.time()
    if not markdown or not isinstance(markdown, str):
        raise ToolInputError("Input markdown must be a non-empty string", param_name="markdown")

    content_type_result = await detect_content_type(markdown)
    input_type = content_type_result.get("content_type", "unknown")
    actual_markdown = markdown
    conversion_note = ""

    if input_type == "html":
        logger.info("Input detected as HTML, converting to Markdown before optimizing.")
        conversion_note = "⚠️ Input was detected as HTML and automatically converted. "
        try:
            conversion_result = await clean_and_format_text_as_markdown(
                text=markdown, force_markdown_conversion=True, extraction_method="none"
            )
            if conversion_result.get("success", False):
                actual_markdown = conversion_result.get("markdown_text", "")
            else:
                raise ToolError(
                    "FORMAT_CONVERSION_FAILED",
                    details={"error": conversion_result.get("error", "Unknown")},
                )
        except Exception as e_conv:
            logger.error(f"Failed to convert HTML input for optimization: {e_conv}", exc_info=True)
            return {
                "success": False,
                "error": f"Input HTML conversion failed: {e_conv}",
                "error_code": "FORMAT_CONVERSION_FAILED",
            }
    elif input_type != "markdown":
        logger.warning(
            f"Input detected as {input_type}, applying Markdown optimization rules anyway."
        )
        conversion_note = f"⚠️ Input detected as {input_type}, not Markdown. "

    optimized = actual_markdown
    changes = []

    if normalize_headings:
        logger.warning("normalize_headings is not fully implemented.")
        changes.append("Attempted heading normalization (basic)")

    if fix_lists:
        optimized = _sanitize(optimized)  # Handles basic list marker normalization
        changes.append("Standardized list formatting via sanitize")

    if fix_links:
        optimized = re.sub(r"\[\s*([^\]]+?)\s*\]\s*\(\s*([^)]+?)\s*\)", r"[\1](\2)", optimized)
        changes.append("Cleaned link formatting")

    if compact_mode:
        optimized = re.sub(r"\n{3,}", "\n\n", optimized)
        optimized = re.sub(r"[ \t]+$", "", optimized, flags=re.MULTILINE)
        changes.append("Applied compact formatting")
    elif add_line_breaks:
        optimized = _improve(optimized)
        changes.append("Added standard line breaks")

    if max_line_length > 0:
        try:
            wrapped_lines = []
            current_block = ""
            in_code_block = False
            for line in optimized.split("\n"):
                line_stripped = line.strip()
                if line_stripped.startswith("```"):
                    in_code_block = not in_code_block
                    if current_block:
                        wrapped_lines.extend(
                            textwrap.wrap(
                                current_block.strip(),
                                width=max_line_length,
                                break_long_words=False,
                                break_on_hyphens=False,
                            )
                        )
                    current_block = ""
                    wrapped_lines.append(line)
                elif (
                    in_code_block
                    or line_stripped.startswith(("#", ">", "- ", "* ", "+ "))
                    or re.match(r"^\d+\.\s", line_stripped)
                    or line_stripped == ""
                    or line_stripped.startswith("|")
                ):
                    if current_block:
                        wrapped_lines.extend(
                            textwrap.wrap(
                                current_block.strip(),
                                width=max_line_length,
                                break_long_words=False,
                                break_on_hyphens=False,
                            )
                        )
                    current_block = ""
                    wrapped_lines.append(line)
                else:
                    current_block += line + " "
            if current_block:
                wrapped_lines.extend(
                    textwrap.wrap(
                        current_block.strip(),
                        width=max_line_length,
                        break_long_words=False,
                        break_on_hyphens=False,
                    )
                )
            optimized = "\n".join(wrapped_lines)
            changes.append(f"Wrapped lines at {max_line_length} chars")
        except Exception as e_wrap:
            logger.error(f"Error during line wrapping: {e_wrap}")

    return {
        "success": True,
        "optimized_markdown": optimized.strip(),
        "changes_summary": conversion_note
        + (", ".join(changes) if changes else "No specific optimizations applied."),
        "processing_time": time.time() - t0,
    }


# <<< Part 1 and 2 code goes here >>>

###############################################################################
# Document Analysis Tools (Standalone)                                        #
###############################################################################


@with_tool_metrics
@with_error_handling
async def identify_sections(document: str) -> Dict[str, Any]:
    """Identifies logical sections in a document using regex patterns (Standalone Tool)."""
    if not document or not isinstance(document, str):
        logger.warning("identify_sections called with empty or invalid input.")
        return {"sections": [], "success": True}

    domain_rules = _get_active_domain_rules()
    bound_rx = domain_rules.get("bound_rx")
    custom_sect_rx = domain_rules.get("custom_sect_rx", [])
    if not bound_rx or not isinstance(bound_rx, re.Pattern):
        raise ToolError(
            "INITIALIZATION_ERROR", details={"reason": "Section boundary regex not loaded/compiled"}
        )

    sections_found: List[Dict[str, Any]] = []
    last_section_end = 0
    try:
        matches = list(bound_rx.finditer(document))
        if not matches:
            logger.info(
                "No regex-based section boundaries found. Treating document as single section."
            )
            if document.strip():
                sections_found.append(
                    {
                        "title": "Main Content",
                        "text": document.strip(),
                        "position": 0,
                        "start_char": 0,
                        "end_char": len(document),
                    }
                )
        else:
            logger.info(f"Found {len(matches)} potential section boundaries based on regex.")
            first_match_start = matches[0].start()
            if first_match_start > 0:
                initial_text = document[last_section_end:first_match_start].strip()
                if initial_text:
                    sections_found.append(
                        {
                            "title": "Introduction",
                            "text": initial_text,
                            "position": 0,
                            "start_char": last_section_end,
                            "end_char": first_match_start,
                        }
                    )
                    last_section_end = first_match_start
            for i, match in enumerate(matches):
                title_raw = match.group(0).strip()
                title_start_char = match.start()
                title_end_char = match.end()
                section_content_start = title_end_char
                section_content_end = (
                    matches[i + 1].start() if i < len(matches) - 1 else len(document)
                )
                section_text = document[section_content_start:section_content_end].strip()
                section_title = title_raw
                if custom_sect_rx:
                    for pat, label in custom_sect_rx:
                        if isinstance(pat, re.Pattern) and pat.search(title_raw):
                            section_title = label
                            logger.debug(
                                f"Applied custom label '{label}' to section '{title_raw}'."
                            )
                            break
                if section_text:
                    sections_found.append(
                        {
                            "title": section_title,
                            "text": section_text,
                            "position": len(sections_found),
                            "start_char": title_start_char,
                            "end_char": section_content_end,
                        }
                    )
                else:
                    logger.debug(
                        f"Skipping section '{section_title}' (no content)."
                    )  # Corrected typo
                last_section_end = section_content_end
    except Exception as e:
        logger.error(f"Error during section identification: {e}", exc_info=True)
        raise ToolError("SECTION_IDENTIFICATION_FAILED", details={"error": str(e)}) from e
    return {"sections": sections_found, "success": True}


@with_tool_metrics
@with_error_handling
async def extract_entities(
    document: str, entity_types: Optional[List[str]] = None
) -> Dict[str, Any]:
    """Extracts named entities from document text using an LLM (Standalone Tool)."""
    if not document or not isinstance(document, str):
        raise ToolInputError("Input document must be non-empty string", param_name="document")
    max_context = 3800
    context = document[:max_context] + ("..." if len(document) > max_context else "")
    if len(document) > max_context:
        logger.warning(f"Doc truncated to ~{max_context} chars for entity extraction.")
    entity_focus = (
        f"Extract only: {', '.join(entity_types)}."
        if entity_types
        else "Extract common types (PERSON, ORG, LOC, DATE, MONEY...)."
    )
    prompt = f"""Analyze text, extract entities. {entity_focus} Output ONLY valid JSON object (keys=TYPE, values=list of unique strings).
Text:
\"\"\"
{context}
\"\"\"
JSON Output:
"""
    logger.info(f"Requesting entity extraction. Focus: {entity_types or 'common'}")
    llm_response_raw = ""
    try:
        llm_response_raw = await _standalone_llm_call(
            prompt=prompt, max_tokens=1500, temperature=0.1
        )
        logger.debug(f"LLM raw response for entities:\n{llm_response_raw}")
        json_str = llm_response_raw
        json_match = re.search(r"```(?:json)?\s*([\s\S]+?)\s*```", json_str)
        if json_match:
            json_str = json_match.group(1).strip()
        start_brace = json_str.find("{")
        end_brace = json_str.rfind("}")
        if start_brace != -1 and end_brace != -1 and start_brace < end_brace:
            json_str = json_str[start_brace : end_brace + 1]
        else:
            logger.warning("Could not find JSON object boundaries for entities.")
        try:
            entities_dict = json.loads(json_str)
        except json.JSONDecodeError as json_e:
            logger.warning(f"Initial JSON parse failed: {json_e}. Trying lenient find.")
            match = re.search(r"\{.*\}", llm_response_raw, re.DOTALL)
            if match:
                try:
                    entities_dict = json.loads(match.group(0))
                except json.JSONDecodeError as final_json_e:
                    raise ValueError(f"Could not parse JSON: {final_json_e}") from final_json_e
            else:
                raise ValueError("No JSON object found") from json_e
        if not isinstance(entities_dict, dict):
            raise ValueError("LLM response is not JSON object.")
        validated_entities: Dict[str, List[str]] = {}
        for key, value in entities_dict.items():
            entity_type = str(key).upper().strip()
            if not entity_type:
                continue
            sanitized_values: Set[str] = set()
            items_to_process = value if isinstance(value, list) else [value]
            for item in items_to_process:
                text_val = None
                if isinstance(item, str) and item.strip():
                    text_val = item.strip()
                elif (
                    isinstance(item, dict)
                    and isinstance(item.get("text"), str)
                    and item["text"].strip()
                ):
                    text_val = item["text"].strip()
                if text_val:
                    text_val = re.sub(r"^[.,!?;:'\"\(] +|[ .,!?;:'\"\) ]+$", "", text_val)
                if text_val:
                    sanitized_values.add(text_val)
            if sanitized_values:
                validated_entities[entity_type] = sorted(list(sanitized_values))
        logger.info(f"Successfully extracted entities for types: {list(validated_entities.keys())}")
        return {
            "entities": validated_entities,
            "success": True,
            "raw_llm_response": llm_response_raw,
        }
    except (json.JSONDecodeError, ValueError) as e:
        logger.error(f"Failed to parse LLM response for entities: {e}")
        return {
            "entities": {},
            "error": f"Parse fail: {e}",
            "raw_llm_response": llm_response_raw,
            "success": False,
            "error_code": "LLM_INVALID_RESPONSE",
        }
    except ToolError as e:
        logger.error(f"LLM call failed during entity extraction: {e}", exc_info=False)
        return {
            "entities": {},
            "error": f"LLM fail: {str(e)}",
            "raw_llm_response": llm_response_raw,
            "success": False,
            "error_code": e.error_code,
        }
    except Exception as e:
        logger.error(f"Unexpected entity extraction error: {e}", exc_info=True)
        return {
            "entities": {},
            "error": f"Unexpected: {e}",
            "raw_llm_response": llm_response_raw,
            "success": False,
            "error_code": "ENTITY_EXTRACTION_FAILED",
        }


@with_tool_metrics
@with_error_handling
async def generate_qa_pairs(document: str, num_questions: int = 5) -> Dict[str, Any]:
    """Generates question-answer pairs based on the document content using an LLM (Standalone Tool)."""
    if not document or not isinstance(document, str):
        raise ToolInputError("Input must be non-empty string", param_name="document")
    if not isinstance(num_questions, int) or num_questions <= 0:
        raise ToolInputError("num_questions must be positive int")
    max_context = 3800
    context = document[:max_context] + ("..." if len(document) > max_context else "")
    if len(document) > max_context:
        logger.warning(f"Doc truncated to ~{max_context} chars for QA generation.")
    prompt = f"""Based ONLY on text, generate {num_questions} relevant QA pairs. Output ONLY JSON list of objects (keys "question", "answer").
Text:
\"\"\"
{context}
\"\"\"
JSON Output:
"""
    logger.info(f"Requesting {num_questions} QA pairs.")
    llm_response_raw = ""
    try:
        llm_max_tokens = num_questions * 150 + 200
        llm_response_raw = await _standalone_llm_call(
            prompt=prompt, max_tokens=llm_max_tokens, temperature=0.4
        )
        logger.debug(f"LLM raw response for QA pairs:\n{llm_response_raw}")
        json_str = llm_response_raw
        json_match = re.search(r"```(?:json)?\s*([\s\S]+?)\s*```", json_str, re.I)
        if json_match:
            json_str = json_match.group(1).strip()
        start_bracket = json_str.find("[")
        end_bracket = json_str.rfind("]")
        if start_bracket != -1 and end_bracket != -1 and start_bracket < end_bracket:
            json_str = json_str[start_bracket : end_bracket + 1]
        else:
            logger.warning("Could not find JSON list boundaries for QA.")
        try:
            qa_list = json.loads(json_str)
        except json.JSONDecodeError as json_e:
            logger.warning(f"JSON parse for QA failed: {json_e}. Trying regex.")
            pairs = []
            extracted = re.findall(
                r'\{\s*"question":\s*"(.*?)",\s*"answer":\s*"(.*?)"\s*\}',
                llm_response_raw,
                re.DOTALL | re.I,
            )
            pairs = [
                {
                    "question": q.strip().replace('\\"', '"'),
                    "answer": a.strip().replace('\\"', '"'),
                }
                for q, a in extracted
                if q.strip() and a.strip()
            ]
            if pairs:
                logger.info(f"Regex fallback extracted {len(pairs)} pairs.")
                return {
                    "qa_pairs": pairs[:num_questions],
                    "success": True,
                    "warning": "Used regex fallback.",
                    "raw_llm_response": llm_response_raw,
                }
            else:
                raise ValueError("Regex fallback found no QA pairs.") from json_e
        if not isinstance(qa_list, list):
            raise ValueError("LLM response is not JSON list.")
        validated_pairs: List[Dict[str, str]] = []
        for item in qa_list:
            if isinstance(item, dict):
                q = item.get("question")
                a = item.get("answer")
                if isinstance(q, str) and q.strip() and isinstance(a, str) and a.strip():
                    validated_pairs.append({"question": q.strip(), "answer": a.strip()})
                else:
                    logger.warning(f"Skipping invalid QA item: {item}")
            else:
                logger.warning(f"Skipping non-dict item in QA list: {item}")
        if not validated_pairs:
            logger.warning("LLM response parsed but no valid QA pairs found.")
        else:
            logger.info(f"Successfully generated {len(validated_pairs)} valid QA pairs.")
        return {
            "qa_pairs": validated_pairs[:num_questions],
            "success": True,
            "raw_llm_response": llm_response_raw,
        }
    except (json.JSONDecodeError, ValueError) as e:
        logger.error(f"Failed to parse LLM response for QA pairs: {e}")
        return {
            "qa_pairs": [],
            "error": f"Parse fail: {e}",
            "raw_llm_response": llm_response_raw,
            "success": False,
            "error_code": "LLM_INVALID_RESPONSE",
        }
    except ToolError as e:
        logger.error(f"LLM call failed during QA generation: {e}", exc_info=False)
        return {
            "qa_pairs": [],
            "error": f"LLM fail: {str(e)}",
            "raw_llm_response": llm_response_raw,
            "success": False,
            "error_code": e.error_code,
        }
    except Exception as e:
        logger.error(f"Unexpected QA generation error: {e}", exc_info=True)
        return {
            "qa_pairs": [],
            "error": f"Unexpected: {e}",
            "raw_llm_response": llm_response_raw,
            "success": False,
            "error_code": "QA_GENERATION_FAILED",
        }


@with_tool_metrics
@with_error_handling
async def summarize_document(
    document: str, max_length: int = 150, focus: Optional[str] = None
) -> Dict[str, Any]:
    """
    Generates a concise summary of the document text using an LLM (Standalone Tool).
    """
    _logger = logger
    if not document or not isinstance(document, str):
        raise ToolInputError("Input document must be a non-empty string", param_name="document")
    if not isinstance(max_length, int) or max_length <= 10:
        raise ToolInputError("max_length must be a positive integer > 10", param_name="max_length")

    llm_caller = _standalone_llm_call
    max_context = 8000
    context = document[:max_context] + ("..." if len(document) > max_context else "")
    if len(document) > max_context:
        _logger.warning(f"Document truncated to ~{max_context} chars for summarization.")

    focus_instruction = (
        f" Focus particularly on aspects related to: {focus}." if focus and focus.strip() else ""
    )
    prompt = f"""Generate a concise, coherent summary of the following text, about {max_length} words long.{focus_instruction}
Capture the main points and key information accurately based ONLY on the provided text. Do not add external information or opinions.
Output ONLY the summary text itself, without any introductory phrases like "Here is the summary:".

Text:
\"\"\"
{context}
\"\"\"

Summary:
"""
    _logger.info(
        f"Requesting summary from LLM (max_length≈{max_length}, focus='{focus or 'none'}')."
    )
    llm_response_raw = ""
    try:
        llm_max_tokens = max(50, min(4000, int(max_length / 0.6)))
        summary_text = await llm_caller(prompt=prompt, max_tokens=llm_max_tokens, temperature=0.5)
        llm_response_raw = summary_text
        summary_text = re.sub(
            r"^(Here is a summary:|Summary:|The text discusses|This document is about)\s*:?\s*",
            "",
            summary_text,
            flags=re.I,
        ).strip()
        word_count = len(summary_text.split())
        _logger.info(f"Generated summary with {word_count} words (target: {max_length}).")
        if word_count < max_length * 0.5:
            _logger.warning(f"Summary ({word_count} words) shorter than requested ({max_length}).")
        elif word_count > max_length * 1.5:
            _logger.warning(f"Summary ({word_count} words) longer than requested ({max_length}).")
        return {
            "summary": summary_text,
            "word_count": word_count,
            "success": True,
            "raw_llm_response": llm_response_raw,
        }
    except ToolError as te:
        _logger.error(f"ToolError during summarization: {te.error_code} - {str(te)}", exc_info=True)
        return {
            "summary": "",
            "word_count": 0,
            "error": str(te),
            "success": False,
            "raw_llm_response": llm_response_raw,
            "error_code": te.error_code,
        }
    except Exception as e:
        _logger.error(f"Unexpected error during summarization: {str(e)}", exc_info=True)
        return {
            "summary": "",
            "word_count": 0,
            "error": f"Unexpected error: {e}",
            "success": False,
            "raw_llm_response": llm_response_raw,
            "error_code": "SUMMARIZATION_FAILED",
        }


@with_tool_metrics
@with_error_handling
async def extract_metrics(document: str) -> Dict[str, Any]:
    """Extracts numeric metrics based on domain patterns (Standalone Tool)."""
    if not document or not isinstance(document, str):
        raise ToolInputError("Input must be non-empty string", param_name="document")
    domain_rules = _get_active_domain_rules()
    metric_rx_list = domain_rules.get("metric_rx")
    if not metric_rx_list or not isinstance(metric_rx_list, list):
        logger.warning(
            f"No metric patterns found for domain '{domain_rules.get('active_domain', 'unknown')}'."
        )
        return {"metrics": {}, "success": True}
    extracted_metrics: Dict[str, List[float]] = {}
    logger.info(
        f"Starting metric extraction for domain '{domain_rules.get('active_domain')}' ({len(metric_rx_list)} types)."
    )
    for metric_name, pattern in metric_rx_list:
        if not isinstance(pattern, re.Pattern):
            logger.warning(f"Skipping invalid pattern type for metric '{metric_name}'.")
            continue
        found_values: Set[float] = set()
        try:
            matches = pattern.findall(document)
            if matches:
                logger.debug(f"Found {len(matches)} potential matches for metric '{metric_name}'")
            for match_groups in matches:
                val_str = None
                if isinstance(match_groups, tuple) and len(match_groups) >= 2:
                    val_str = str(match_groups[1]).strip()
                elif isinstance(match_groups, str):
                    val_str = match_groups.strip()
                if val_str is None:
                    continue
                val_str_cleaned = re.sub(r"[$,€£\s,]", "", val_str)
                if val_str_cleaned.endswith("."):
                    val_str_cleaned = val_str_cleaned[:-1]
                if not val_str_cleaned or val_str_cleaned == "-":
                    continue
                try:
                    found_values.add(float(val_str_cleaned))
                except ValueError:
                    logger.debug(
                        f"Could not convert value '{val_str_cleaned}' for metric '{metric_name}'."
                    )
        except Exception as e:
            logger.error(f"Error processing regex for metric '{metric_name}': {e}", exc_info=True)
        if found_values:
            unique_values = sorted(list(found_values))
            extracted_metrics[metric_name] = unique_values
            logger.info(
                f"Extracted {len(unique_values)} unique value(s) for metric '{metric_name}': {unique_values}"
            )
    return {"metrics": extracted_metrics, "success": True}


@with_tool_metrics
@with_error_handling
async def flag_risks(document: str) -> Dict[str, Any]:
    """Flags potential risks using domain patterns (Standalone Tool)."""
    if not document or not isinstance(document, str):
        raise ToolInputError("Input must be non-empty string", param_name="document")
    domain_rules = _get_active_domain_rules()
    risk_rx_dict = domain_rules.get("risk_rx")
    if not risk_rx_dict or not isinstance(risk_rx_dict, dict):
        logger.warning(
            f"No risk patterns found for domain '{domain_rules.get('active_domain', 'unknown')}'."
        )
        return {"risks": {}, "success": True}
    flagged_risks: Dict[str, Dict[str, Any]] = {}
    logger.info(
        f"Starting risk flagging for domain '{domain_rules.get('active_domain')}' ({len(risk_rx_dict)} types)."
    )
    context_window = 50
    max_samples = 3
    for risk_type, pattern in risk_rx_dict.items():
        if not isinstance(pattern, re.Pattern):
            logger.warning(f"Skipping invalid pattern type for risk '{risk_type}'.")
            continue
        match_contexts: List[str] = []
        match_count = 0
        try:
            for match in pattern.finditer(document):
                match_count += 1
                if len(match_contexts) < max_samples:
                    start, end = match.start(), match.end()
                    ctx_start = max(0, start - context_window)
                    ctx_end = min(len(document), end + context_window)
                    snippet = document[ctx_start:ctx_end].replace("\n", " ").strip()
                    prefix = "..." if ctx_start > 0 else ""
                    suffix = "..." if ctx_end < len(document) else ""
                    hl_start = start - ctx_start + len(prefix)
                    hl_end = end - ctx_start + len(prefix)
                    formatted_snippet = f"{prefix}{snippet[:hl_start]}**{snippet[hl_start:hl_end]}**{snippet[hl_end:]}{suffix}"
                    match_contexts.append(formatted_snippet)
            if match_count > 0:
                logger.info(f"Flagged risk '{risk_type}' {match_count} time(s).")
                flagged_risks[risk_type] = {"count": match_count, "sample_contexts": match_contexts}
        except Exception as e:
            logger.error(f"Error processing regex for risk '{risk_type}': {e}", exc_info=True)
    return {"risks": flagged_risks, "success": True}


@with_tool_metrics
@with_error_handling
async def canonicalise_entities(entities_input: Dict[str, Any]) -> Dict[str, Any]:
    """Normalizes and attempts to merge similar entities using fuzzy matching (Standalone Tool)."""
    entities_list: List[Dict[str, Any]] = []
    raw_entities = entities_input.get("entities")
    if isinstance(raw_entities, dict):
        for etype, text_list in raw_entities.items():
            entity_type_str = str(etype).upper().strip()
            if isinstance(text_list, list) and entity_type_str:
                for text in text_list:
                    if isinstance(text, str) and text.strip():
                        entities_list.append({"text": text.strip(), "type": entity_type_str})
            else:
                logger.warning(
                    f"Expected list for entity type '{etype}', got {type(text_list)}. Skipping."
                )
    elif isinstance(raw_entities, list):
        for item in raw_entities:
            if (
                isinstance(item, dict)
                and isinstance(item.get("text"), str)
                and item["text"].strip()
                and isinstance(item.get("type"), str)
                and item["type"].strip()
            ):
                entities_list.append(
                    {
                        "text": item["text"].strip(),
                        "type": item["type"].upper().strip(),
                        "metadata": {k: v for k, v in item.items() if k not in ["text", "type"]},
                    }
                )
            else:
                logger.warning(f"Skipping invalid item in entity list: {item}")
    else:
        raise ToolInputError(
            'Input dict must contain "entities" key with either Dict[str, List[str]] or List[Dict[str, Any]].',
            param_name="entities_input",
        )
    if not entities_list:
        logger.info("No entities provided for canonicalization.")
        return {"canonicalized": {}, "success": True}
    entities_by_type: Dict[str, List[Dict[str, Any]]] = {}
    for entity in entities_list:
        etype = entity.get("type", "UNKNOWN")
        entities_by_type.setdefault(etype, []).append(entity)
    canonicalized_output: Dict[str, List[Dict[str, Any]]] = {}
    similarity_threshold = 85
    for entity_type, entity_group in entities_by_type.items():
        logger.debug(f"Canonicalising {len(entity_group)} entities of type '{entity_type}'...")
        entity_group.sort(key=lambda x: len(x.get("text", "")), reverse=True)
        merged_entities: List[Dict[str, Any]] = []
        processed_indices = set()
        for i in range(len(entity_group)):
            if i in processed_indices:
                continue
            current_entity = entity_group[i]
            canonical_form = current_entity.get("text", "")
            if not canonical_form:
                processed_indices.add(i)
                continue
            cluster_variants_data = [current_entity]
            processed_indices.add(i)
            for j in range(i + 1, len(entity_group)):
                if j in processed_indices:
                    continue
                other_entity = entity_group[j]
                other_text = other_entity.get("text", "")
                if not other_text:
                    processed_indices.add(j)
                    continue
                score = fuzz.token_sort_ratio(canonical_form.lower(), other_text.lower())
                if score >= similarity_threshold:
                    cluster_variants_data.append(other_entity)
                    processed_indices.add(j)
                    logger.debug(
                        f"  Merging '{other_text}' into '{canonical_form}' (score: {score:.0f})"
                    )
            if cluster_variants_data:
                canonical_text = cluster_variants_data[0].get("text", "")
                variant_texts = sorted(
                    list({ent.get("text", "") for ent in cluster_variants_data if ent.get("text")})
                )
                merged_metadata = {}
                scores = [
                    ent.get("metadata", {}).get("score")
                    for ent in cluster_variants_data
                    if ent.get("metadata", {}).get("score") is not None
                ]
                if scores:
                    merged_metadata["scores"] = scores
                merged_entities.append(
                    {
                        "text": canonical_text,
                        "count": len(cluster_variants_data),
                        "type": entity_type,
                        "variants": variant_texts,
                        "metadata": merged_metadata,
                    }
                )
        merged_entities.sort(key=lambda x: (-x.get("count", 0), x.get("text", "")))
        canonicalized_output[entity_type] = merged_entities
        logger.info(
            f"Canonicalised type '{entity_type}': {len(entity_group)} input -> {len(merged_entities)} unique entities."
        )
    return {"canonicalized": canonicalized_output, "success": True}


###############################################################################
# OCR-Specific Tools (Standalone)                                             #
###############################################################################


@with_tool_metrics
@with_retry(max_retries=2, retry_delay=1.5)
@with_error_handling
async def ocr_image(
    image_path: Optional[str] = None,
    image_data: Optional[str] = None,  # Base64 encoded string
    ocr_options: Optional[Dict] = None,
    enhance_with_llm: bool = True,
    output_format: str = "markdown",
) -> Dict[str, Any]:
    """
    Performs OCR on a single image and optionally enhances the text with an LLM (Standalone Tool).

    Args:
        image_path: Path to the image file (e.g., PNG, JPG). Mutually exclusive with image_data.
        image_data: Base64-encoded image data string. Mutually exclusive with image_path.
        ocr_options: Dictionary of options for OCR/Enhancement:
            - language (str): Tesseract language(s). Default: "eng".
            - preprocessing (dict): Image preprocessing options.
            - remove_headers (bool): Attempt header/footer removal (less effective on single images). Default: False.
            - assess_quality (bool): Run LLM quality assessment. Default: False.
            - detect_tables (bool): Attempt to detect tables in the image (used for metadata). Default: True.
            - tesseract_config (str): Additional Tesseract config options (e.g., '--psm 6'). Default: "".
        enhance_with_llm: If True (default), enhance the raw OCR text using an LLM.
        output_format: Target format ('markdown' or 'text').

    Returns:
        Dictionary with OCR results (see convert_document return structure).
    """
    t0 = time.time()
    ocr_opts = ocr_options or {}
    output_format = output_format.lower()
    if output_format not in _OCR_COMPATIBLE_FORMATS:
        logger.warning(
            f"Output format '{output_format}' not ideal for image OCR. Using 'markdown'."
        )
        output_format = "markdown"

    # --- Dependency Checks ---
    _ocr_check_dep("Pillow", _PIL_AVAILABLE, "Image OCR")
    _ocr_check_dep("pytesseract", _PYTESSERACT_AVAILABLE, "Image OCR")
    can_use_cv2 = _CV2_AVAILABLE and _NUMPY_AVAILABLE
    if ocr_opts.get("preprocessing") and not can_use_cv2:
        logger.warning(
            "Preprocessing options provided but OpenCV/NumPy missing. Preprocessing limited."
        )
    if ocr_opts.get("detect_tables", True) and not can_use_cv2:
        logger.warning(
            "Table detection requires OpenCV/NumPy. Disabling table detection for metadata."
        )
        ocr_opts["detect_tables"] = False

    # --- Input Handling ---
    if not image_path and not image_data:
        raise ToolInputError("Either 'image_path' or 'image_data' must be provided.")
    if image_path and image_data:
        raise ToolInputError("Provide either 'image_path' or 'image_data', not both.")

    img: Optional["PILImage.Image"] = None
    preprocessed_img: Optional["PILImage.Image"] = None
    input_name = "image_data"

    try:
        if image_path:
            img_path_obj = _ocr_validate_file_path(image_path)
            input_name = img_path_obj.name
            with _span(f"load_image_{input_name}"):
                img = Image.open(img_path_obj)  # type: ignore
        elif image_data:
            if not isinstance(image_data, str):
                raise ToolInputError("image_data must be a base64 encoded string.")
            try:
                if image_data.startswith("data:image"):
                    image_data = image_data.split(";base64,", 1)[1]
                img_bytes = base64.b64decode(image_data)
                with _span("load_image_bytes"):
                    img = Image.open(io.BytesIO(img_bytes))  # type: ignore
                input_name = f"base64_input_{_hash(image_data[:100])}"
            except (base64.binascii.Error, ValueError, TypeError) as e_b64:
                raise ToolInputError(
                    f"Invalid base64 image data: {e_b64}", param_name="image_data"
                ) from e_b64
            except Exception as e_img_open:
                raise ToolError(
                    "IMAGE_LOAD_FAILED",
                    details={"error": f"Failed to open image from bytes: {e_img_open}"},
                ) from e_img_open

        if img is None or not _PIL_AVAILABLE:  # Added check for _PIL_AVAILABLE
            raise ToolError(
                "IMAGE_LOAD_FAILED",
                details={"reason": "Image object is None or Pillow unavailable."},
            )

        img = img.convert("RGB")

        # --- OCR Pipeline ---
        loop = asyncio.get_running_loop()
        with _span("image_preprocessing"):
            preprocessed_img = await loop.run_in_executor(
                None, _ocr_preprocess_image, img, ocr_opts.get("preprocessing")
            )

        ocr_lang = ocr_opts.get("language", "eng")
        ocr_config_str = ocr_opts.get("tesseract_config", "")
        with _span("tesseract_ocr"):
            raw_text = await loop.run_in_executor(
                None, _ocr_run_tesseract, preprocessed_img, ocr_lang, ocr_config_str
            )

        # --- LLM Enhancement ---
        final_content = raw_text
        quality_metrics = None
        if enhance_with_llm and raw_text.strip():
            with _span("llm_image_text_enhancement"):
                remove_headers = ocr_opts.get("remove_headers", False)
                final_content = await _ocr_enhance_text_chunk(
                    raw_text, output_format=output_format, remove_headers=remove_headers
                )

            if ocr_opts.get("assess_quality", False):
                with _span("ocr_quality_assessment"):
                    quality_metrics = await _ocr_assess_text_quality(raw_text, final_content)
        else:
            final_content = raw_text

        # --- Metadata ---
        tables_detected = False
        if ocr_opts.get("detect_tables", True) and can_use_cv2:
            # Run detection on the preprocessed image
            detected_regions = await loop.run_in_executor(
                None, _ocr_detect_tables, preprocessed_img
            )
            tables_detected = len(detected_regions) > 0
        elif ocr_opts.get("detect_tables", True):
            logger.warning("Table detection requested but OpenCV/Numpy unavailable.")

        doc_metadata = {
            "num_pages": 1,
            "has_tables": tables_detected,
            "has_figures": True,
            "has_sections": bool(re.search(r"^#{1,6}\s+", final_content, re.M))
            if output_format == "markdown"
            else False,
            "image_width": img.width,
            "image_height": img.height,
            "ocr_language": ocr_lang,
        }

        # --- Construct Response ---
        elapsed = round(time.time() - t0, 3)
        response = {
            "success": True,
            "content": final_content,
            "output_format": output_format,
            "processing_time": elapsed,
            "document_metadata": doc_metadata,
            "extraction_strategy_used": "ocr",
        }
        if enhance_with_llm:
            response["raw_text"] = raw_text
        if quality_metrics:
            response["ocr_quality_metrics"] = quality_metrics

        logger.info(f"Completed OCR for '{input_name}' in {elapsed}s")
        return response

    except Exception as e:
        logger.error(f"Error during image OCR for '{input_name}': {e}", exc_info=True)
        if isinstance(e, (ToolInputError, ToolError)):
            raise e
        raise ToolError("IMAGE_OCR_FAILED", details={"input": input_name, "error": str(e)}) from e
    finally:
        if img:
            img.close()
        if preprocessed_img and preprocessed_img != img:
            preprocessed_img.close()


@with_tool_metrics
@with_retry(max_retries=2, retry_delay=1.0)
@with_error_handling
async def enhance_ocr_text(
    text: str,
    output_format: str = "markdown",
    enhancement_options: Optional[Dict] = None,
) -> Dict[str, Any]:
    """
    Enhances existing OCR text using an LLM to correct errors and improve formatting (Standalone Tool).

    Args:
        text: The raw OCR text to enhance.
        output_format: Target format ('markdown' or 'text').
        enhancement_options: Dictionary of options:
            - remove_headers (bool): Attempt to remove headers/footers. Default: False.
            - assess_quality (bool): Run LLM quality assessment comparing input vs output. Default: False.

    Returns:
        Dictionary containing enhanced text and metadata.
    """
    t0 = time.time()
    if not text or not isinstance(text, str):
        raise ToolInputError("Input 'text' must be a non-empty string", param_name="text")

    options = enhancement_options or {}
    output_format = output_format.lower()
    if output_format not in _OCR_COMPATIBLE_FORMATS:
        logger.warning(
            f"Output format '{output_format}' not ideal for text enhancement. Using 'markdown'."
        )
        output_format = "markdown"

    try:
        final_content = ""
        quality_metrics = None

        with _span("llm_text_enhancement"):
            max_direct_process_len = 15000
            if len(text) > max_direct_process_len:
                logger.info(f"Splitting large text ({len(text)} chars) for enhancement.")
                chunks = _ocr_split_text_into_chunks(text)  # Helper defined earlier
            else:
                chunks = [text]

            if not chunks:
                logger.warning("Input text resulted in zero chunks for enhancement.")
            else:
                enhancement_tasks = [
                    _ocr_enhance_text_chunk(
                        chunk,
                        output_format=output_format,
                        remove_headers=options.get("remove_headers", False),
                    )
                    for chunk in chunks
                ]  # Helper defined earlier
                enhanced_chunks = await asyncio.gather(*enhancement_tasks)
                final_content = "\n\n".join(enhanced_chunks).strip()

        if options.get("assess_quality", False):
            with _span("ocr_quality_assessment"):
                quality_metrics = await _ocr_assess_text_quality(
                    text, final_content
                )  # Helper defined earlier

        elapsed = round(time.time() - t0, 3)
        response = {
            "success": True,
            "content": final_content,
            "output_format": output_format,
            "processing_time": elapsed,
            "raw_text": text,
        }
        if quality_metrics:
            response["ocr_quality_metrics"] = quality_metrics
        logger.info(f"Completed OCR text enhancement in {elapsed}s")
        return response

    except Exception as e:
        logger.error(f"Error during OCR text enhancement: {e}", exc_info=True)
        if isinstance(e, (ToolInputError, ToolError)):
            raise e
        raise ToolError("TEXT_ENHANCEMENT_FAILED", details={"error": str(e)}) from e


@with_tool_metrics
@with_retry(max_retries=2, retry_delay=1.0)
@with_error_handling
async def analyze_pdf_structure(
    file_path: Optional[str] = None,
    document_data: Optional[bytes] = None,
    extract_metadata: bool = True,
    extract_outline: bool = True,
    extract_fonts: bool = False,
    extract_images: bool = False,
    estimate_ocr_needs: bool = True,
) -> Dict[str, Any]:
    """
    Analyzes PDF structure (metadata, outline, fonts, images, OCR needs) without full text extraction (Standalone Tool).
    Requires either PyMuPDF or PDFPlumber.
    """
    t0 = time.time()
    if not _PYMUPDF_AVAILABLE and not _PDFPLUMBER_AVAILABLE:
        raise ToolError("DEPENDENCY_MISSING", details={"dependency": "PyMuPDF or PDFPlumber"})

    pdf_lib = "pymupdf" if _PYMUPDF_AVAILABLE else "pdfplumber"
    logger.info(f"Analyzing PDF structure using {pdf_lib}.")

    input_path_obj: Optional[Path] = None
    is_temp_file = False
    input_name = "input_data"
    try:
        input_path_obj, is_temp_file = _get_input_path_or_temp(file_path, document_data)
        input_name = input_path_obj.name
        if not input_path_obj.suffix.lower() == ".pdf":
            raise ToolInputError(f"Input must be PDF, got: {input_path_obj.suffix}")

        with _handle_temp_file(input_path_obj, is_temp_file) as current_input_path:
            result: Dict[str, Any] = {
                "success": False,
                "file_info": input_name,
                "analysis_engine": pdf_lib,
                "processing_time": 0.0,
            }
            loop = asyncio.get_running_loop()

            if pdf_lib == "pymupdf":

                def _analyze_with_pymupdf_sync():
                    analysis_data = {}
                    _ocr_check_dep("PyMuPDF", _PYMUPDF_AVAILABLE, "PDF Analysis")
                    if pymupdf is None:
                        raise ToolError("INTERNAL_ERROR", details={"reason": "pymupdf is None"})
                    with pymupdf.open(current_input_path) as doc:  # type: ignore
                        analysis_data["page_count"] = len(doc)
                        if extract_metadata:
                            analysis_data["metadata"] = {
                                k: doc.metadata.get(k, "")
                                for k in [
                                    "title",
                                    "author",
                                    "subject",
                                    "keywords",
                                    "creator",
                                    "producer",
                                    "creationDate",
                                    "modDate",
                                ]
                            }
                        if extract_outline:
                            toc = doc.get_toc()
                            analysis_data["outline"] = (
                                _ocr_process_toc(toc) if toc else []
                            )  # Helper defined earlier
                        if extract_fonts:
                            fonts: Set[str] = set()
                            embedded_fonts: Set[str] = set()
                            limit = min(10, len(doc))
                            for i in range(limit):
                                for font_info in doc.get_page_fonts(i):
                                    fonts.add(font_info[3])
                                    embedded_fonts.add(font_info[3]) if font_info[4] else None
                            analysis_data["font_info"] = {
                                "total_fonts": len(fonts),
                                "embedded_fonts": len(embedded_fonts),
                                "font_names": sorted(list(fonts)),
                            }
                        if extract_images:
                            img_count = 0
                            img_types: Dict[str, int] = {}
                            total_size = 0
                            limit = min(5, len(doc))
                            for i in range(limit):
                                for img in doc.get_page_images(i, full=True):
                                    img_count += 1
                                    xref = img[0]
                                    try:
                                        img_info = doc.extract_image(xref)
                                        img_ext = img_info["ext"]
                                        img_size = len(img_info["image"])
                                    except Exception:
                                        img_ext = "unknown"
                                        img_size = 0
                                    img_types[img_ext] = img_types.get(img_ext, 0) + 1
                                    total_size += img_size
                            est_total = (
                                int(img_count * (len(doc) / max(1, limit)))
                                if limit > 0
                                else img_count
                            )
                            avg_size_kb = (
                                int(total_size / max(1, img_count) / 1024) if img_count > 0 else 0
                            )
                            analysis_data["image_info"] = {
                                "sampled_images": img_count,
                                "estimated_total_images": est_total,
                                "image_types": img_types,
                                "average_size_kb": avg_size_kb,
                            }
                        if estimate_ocr_needs:
                            text_pages = 0
                            sample_size = min(10, len(doc))
                            min_chars = 50
                            for i in range(sample_size):
                                if len(doc[i].get_text("text").strip()) > min_chars:
                                    text_pages += 1
                            text_ratio = text_pages / max(1, sample_size)
                            needs_ocr = text_ratio < 0.8
                            confidence = (
                                "high" if text_ratio < 0.2 or text_ratio > 0.95 else "medium"
                            )
                            reason = (
                                "Likely scanned or image-based."
                                if needs_ocr and confidence == "high"
                                else "Likely contains extractable text."
                                if not needs_ocr and confidence == "high"
                                else "Mix of text/image pages likely."
                            )
                            analysis_data["ocr_assessment"] = {
                                "needs_ocr": needs_ocr,
                                "confidence": confidence,
                                "reason": reason,
                                "text_coverage_ratio": round(text_ratio, 2),
                            }
                    return analysis_data

                result.update(await loop.run_in_executor(None, _analyze_with_pymupdf_sync))
            elif pdf_lib == "pdfplumber":

                def _analyze_with_pdfplumber_sync():
                    analysis_data = {}
                    _ocr_check_dep("pdfplumber", _PDFPLUMBER_AVAILABLE, "PDF Analysis")
                    if pdfplumber is None:
                        raise ToolError("INTERNAL_ERROR", details={"reason": "pdfplumber is None"})
                    with pdfplumber.open(current_input_path) as pdf:  # type: ignore
                        analysis_data["page_count"] = len(pdf.pages)
                        if extract_metadata:
                            analysis_data["metadata"] = {
                                k: pdf.metadata.get(k.capitalize(), "")
                                for k in [
                                    "title",
                                    "author",
                                    "subject",
                                    "keywords",
                                    "creator",
                                    "producer",
                                    "creationDate",
                                    "modDate",
                                ]
                            }
                        if extract_outline:
                            analysis_data["outline"] = {
                                "error": "Outline extraction not supported by pdfplumber."
                            }
                        if extract_fonts:
                            analysis_data["font_info"] = {
                                "error": "Font extraction not supported by pdfplumber."
                            }
                        if extract_images:
                            analysis_data["image_info"] = {
                                "error": "Image info not supported by pdfplumber."
                            }
                        if estimate_ocr_needs:
                            text_pages = 0
                            sample_size = min(10, len(pdf.pages))
                            min_chars = 50
                            for i in range(sample_size):
                                if len((pdf.pages[i].extract_text() or "").strip()) > min_chars:
                                    text_pages += 1
                            text_ratio = text_pages / max(1, sample_size)
                            needs_ocr = text_ratio < 0.8
                            confidence = (
                                "high" if text_ratio < 0.2 or text_ratio > 0.95 else "medium"
                            )
                            reason = (
                                "Likely scanned or image-based."
                                if needs_ocr and confidence == "high"
                                else "Likely contains extractable text."
                                if not needs_ocr and confidence == "high"
                                else "Mix of text/image pages likely."
                            )
                            analysis_data["ocr_assessment"] = {
                                "needs_ocr": needs_ocr,
                                "confidence": confidence,
                                "reason": reason,
                                "text_coverage_ratio": round(text_ratio, 2),
                            }
                    return analysis_data

                result.update(await loop.run_in_executor(None, _analyze_with_pdfplumber_sync))

            result["success"] = True
            result["processing_time"] = round(time.time() - t0, 3)
            logger.info(
                f"PDF structure analysis for '{input_name}' completed in {result['processing_time']:.3f}s using {pdf_lib}"
            )
            return result
    except Exception as e:
        logger.error(f"Error during PDF structure analysis for '{input_name}': {e}", exc_info=True)
        if isinstance(e, (ToolInputError, ToolError)):
            raise e
        raise ToolError(
            "PDF_ANALYSIS_FAILED", details={"input": input_name, "error": str(e)}
        ) from e


@with_tool_metrics
@with_error_handling
async def extract_tables(
    document_path: Optional[str] = None,
    document_data: Optional[bytes] = None,
    *,
    table_mode: str = "csv",
    output_dir: Optional[str] = None,
    accelerator_device: str = "auto",
    num_threads: int = 4,
) -> Dict[str, Any]:
    """
    Extracts tables found in a document using Docling and returns them (Standalone Tool).
    NOTE: This tool currently *requires* the 'docling' extraction strategy implicitly.
    """
    _ocr_check_dep("docling", _DOCLING_AVAILABLE, "Table Extraction (extract_tables tool)")

    valid_modes = {"csv", "json", "pandas"}
    table_mode = table_mode.lower()
    if table_mode not in valid_modes:
        raise ToolInputError(
            f"table_mode must be one of {', '.join(valid_modes)}", param_name="table_mode"
        )
    if table_mode == "pandas":
        _ocr_check_dep("pandas", _PANDAS_AVAILABLE, "extract_tables(mode='pandas')")
        if pd is None:
            raise ToolError(
                "INTERNAL_ERROR", details={"reason": "Pandas check passed but pd is None."}
            )

    input_path_obj: Optional[Path] = None
    is_temp_file = False
    input_name = "input_data"
    try:
        input_path_obj, is_temp_file = _get_input_path_or_temp(document_path, document_data)
        input_name = input_path_obj.name
        logger.info(f"Starting Docling table extraction from {input_name}, mode='{table_mode}'")

        with _handle_temp_file(input_path_obj, is_temp_file) as current_input_path:
            try:
                device_str = accelerator_device.lower()
                if device_str not in _ACCEL_MAP:
                    logger.warning(f"Invalid device '{device_str}', using 'auto'.")
                    device_str = "auto"
                device = _ACCEL_MAP[device_str]
                conv = _get_docling_converter(device, num_threads)
                loop = asyncio.get_running_loop()
                with _span("docling_table_conversion"):
                    result = await loop.run_in_executor(None, conv.convert, current_input_path)
                if result and result.document:
                    doc_obj = result.document
                    logger.info("Docling conversion successful.")
                else:
                    raise ToolError(
                        "CONVERSION_FAILED", details={"reason": "Docling returned empty result"}
                    )
            except Exception as e:
                logger.error(f"Error during Docling conversion: {e}", exc_info=True)
                raise ToolError(
                    "CONVERSION_FAILED", details={"doc": str(current_input_path), "error": str(e)}
                ) from e

            if not doc_obj:
                return {
                    "tables": [],
                    "saved_files": [],
                    "success": False,
                    "error": "Conversion failed.",
                    "error_code": "CONVERSION_FAILED",
                }

            tables_raw_data: List[List[List[str]]] = []
            try:
                with _span("docling_table_extraction"):
                    if hasattr(doc_obj, "get_tables") and callable(doc_obj.get_tables):
                        tables_raw_data = doc_obj.get_tables() or []
                    elif hasattr(doc_obj, "pages") and isinstance(doc_obj.pages, list):
                        logger.warning("Using page iteration for tables.")
                        for page in doc_obj.pages:
                            if (
                                hasattr(page, "content")
                                and page.content
                                and callable(getattr(page.content, "has_tables", None))
                                and page.content.has_tables()
                            ):
                                page_tables = page.content.get_tables()
                                if page_tables and isinstance(page_tables, list):
                                    tables_raw_data.extend(
                                        pt for pt in page_tables if isinstance(pt, list)
                                    )
                    else:
                        logger.error("Cannot extract tables: Missing get_tables/pages.")
                sanitized_tables = []
                for tbl in tables_raw_data:
                    if isinstance(tbl, list) and all(isinstance(row, list) for row in tbl):
                        sanitized_tables.append(
                            [[str(cell) if cell is not None else "" for cell in row] for row in tbl]
                        )
                    else:
                        logger.warning(f"Skipping malformed table: {type(tbl)}")
                tables_raw_data = sanitized_tables
            except Exception as e:
                logger.error(f"Error accessing tables: {e}", exc_info=True)

            if not tables_raw_data:
                logger.warning(f"No tables found in {input_name}.")
                return {"tables": [], "saved_files": [], "success": True}
            logger.info(f"Extracted {len(tables_raw_data)} raw tables.")

            output_tables: List[Any] = []
            saved_files: List[str] = []
            output_dir_path = Path(output_dir) if output_dir else None
            if output_dir_path:
                output_dir_path.mkdir(parents=True, exist_ok=True)

            with _span("table_formatting_saving"):
                for i, raw_table in enumerate(tables_raw_data):
                    processed_table: Any = None
                    file_ext = ""
                    try:
                        if table_mode == "csv":
                            output = StringIO()
                            writer = csv.writer(output, quoting=csv.QUOTE_MINIMAL)
                            writer.writerows(raw_table)
                            processed_table = output.getvalue()
                            file_ext = "csv"
                            save_content = processed_table
                        elif table_mode == "json":
                            processed_table = raw_table
                            file_ext = "json"
                            save_content = _json(processed_table)
                        elif table_mode == "pandas":
                            df = pd.DataFrame(raw_table)
                            if not df.empty and len(df) > 1:
                                first_row = df.iloc[0]
                                is_header = (
                                    sum(
                                        1
                                        for cell in first_row
                                        if not str(cell).replace(".", "", 1).strip("-").isdigit()
                                    )
                                    > len(first_row) / 2
                                )
                                if is_header:
                                    df.columns = first_row
                                    df = df[1:].reset_index(drop=True)
                            processed_table = df
                            file_ext = "csv"
                            save_content = df
                        output_tables.append(processed_table)
                        if output_dir_path and file_ext and save_content is not None:
                            base_name = Path(input_name).stem
                            fp = output_dir_path / f"{base_name}_table_{i + 1}.{file_ext}"
                            try:
                                if isinstance(save_content, str):
                                    fp.write_text(save_content, encoding="utf-8")
                                elif isinstance(save_content, pd.DataFrame):
                                    save_content.to_csv(fp, index=False, encoding="utf-8")
                                saved_files.append(str(fp))
                                logger.debug(f"Saved table {i + 1} to {fp}")
                            except Exception as e_save:
                                logger.error(
                                    f"Failed to save table {i + 1} to {fp}: {e_save}", exc_info=True
                                )
                    except Exception as e_format:
                        logger.error(
                            f"Failed to format table {i} into '{table_mode}': {e_format}",
                            exc_info=True,
                        )
            logger.info(f"Processed {len(output_tables)} tables into '{table_mode}'.")
            return {"tables": output_tables, "saved_files": saved_files, "success": True}
    except Exception as e:
        logger.error(f"Error in extract_tables for '{input_name}': {e}", exc_info=True)
        if isinstance(e, (ToolInputError, ToolError)):
            raise e
        raise ToolError(
            "TABLE_EXTRACTION_FAILED", details={"input": input_name, "error": str(e)}
        ) from e


###############################################################################
# Batch Processing Tool (Standalone)                                          #
###############################################################################


# Map operation names to the standalone functions
# Placed here to ensure all target functions are defined above
_OP_MAP: Dict[str, Callable[..., Awaitable[Any]]] = {
    "convert_document": convert_document,
    "ocr_image": ocr_image,
    "enhance_ocr_text": enhance_ocr_text,
    "clean_and_format_text_as_markdown": clean_and_format_text_as_markdown,
    "optimize_markdown_formatting": optimize_markdown_formatting,
    "detect_content_type": detect_content_type,
    "chunk_document": chunk_document,
    "summarize_document": summarize_document,
    "extract_entities": extract_entities,
    "generate_qa_pairs": generate_qa_pairs,
    "identify_sections": identify_sections,
    "extract_metrics": extract_metrics,
    "flag_risks": flag_risks,
    "canonicalise_entities": canonicalise_entities,
    "analyze_pdf_structure": analyze_pdf_structure,
    "extract_tables": extract_tables,
    "batch_format_texts": batch_format_texts,
}

# Assume necessary imports and _OP_MAP are defined above


@with_tool_metrics
@with_error_handling  # Catch errors setting up the batch itself
async def process_document_batch(
    inputs: List[Dict[str, Any]],
    operations: List[Dict[str, Any]],
    max_concurrency: int = 5,
) -> List[Dict[str, Any]]:
    """
    Processes a list of input items through a sequence of operations concurrently (Standalone Tool).

    Args:
        inputs: List of input dictionaries. Each dict represents an item (e.g., {"document_path": "..."}).
        operations: List of operation specifications. Each dict defines:
            - operation (str): Name of the tool function to call (from _OP_MAP).
            - output_key (str): Key to store the operation's result under in the item's state.
            - params (dict): Fixed parameters for the operation.
            Optional:
            - input_key (str): Key in item state holding primary input (default conventions apply).
            - input_keys_map (dict): Map function parameters to item state keys.
            - promote_output (str): Key in result dict to promote to top-level "content".
        max_concurrency: Max parallel items per operation step.

    Returns:
        List of dictionaries, representing the final state of each input item.
    """
    # --- Input Validation ---
    if not isinstance(inputs, list):
        raise ToolInputError("'inputs' must be a list.")
    if not isinstance(operations, list):
        raise ToolInputError("'operations' must be a list.")
    if not all(isinstance(item, dict) for item in inputs):
        raise ToolInputError("All items in 'inputs' must be dictionaries.")
    if not all(isinstance(op, dict) for op in operations):
        raise ToolInputError("All items in 'operations' must be dictionaries.")
    max_concurrency = max(1, max_concurrency)
    if not inputs:
        logger.warning("Input list is empty.")
        return []

    # --- Initialize Results State ---
    results_state: List[Dict[str, Any]] = []
    for i, item in enumerate(inputs):
        state_item = item.copy()
        state_item["_original_index"] = i
        state_item["_error_log"] = []
        state_item["_status"] = "pending"
        results_state.append(state_item)

    logger.info(f"Starting batch processing: {len(inputs)} items, {len(operations)} operations.")

    # --- Define Worker Outside the Loop ---
    # This worker now takes all potentially changing parameters explicitly
    async def _apply_op_to_item_worker(
        item_state: Dict[str, Any],
        semaphore: asyncio.Semaphore,
        op_func: Callable,
        op_name: str,
        step_label: str,
        op_output_key: str,
        op_input_key: Optional[str],
        op_params: Dict,
        op_input_map: Dict,
        op_promote: Optional[str],
    ) -> Dict[str, Any]:
        item_idx = item_state["_original_index"]
        if item_state["_status"] == "failed":
            return item_state  # Don't process failed items

        async with semaphore:
            logger.debug(f"Applying {step_label} to item {item_idx}")
            call_kwargs = {}
            primary_input_arg_name = None
            input_source_key = None

            try:
                # 1. Determine Primary Input Source Key
                if op_input_key and op_input_key in item_state:
                    input_source_key = op_input_key
                else:
                    potential_keys = []
                    if op_name.startswith(
                        ("convert_document", "ocr_image", "analyze_pdf_structure", "extract_tables")
                    ):
                        potential_keys = [
                            "document_path",
                            "image_path",
                            "file_path",
                            "document_data",
                            "image_data",
                        ]
                    elif op_name == "canonicalise_entities":
                        potential_keys = ["entities_input"]
                    elif op_name == "batch_format_texts":
                        potential_keys = ["texts"]
                    else:
                        potential_keys = ["content", "document", "text"]

                    for key in potential_keys:
                        if key in item_state:
                            input_source_key = key
                            break
                    if not input_source_key:
                        if "content" in item_state:
                            input_source_key = "content"
                        elif "document" in item_state:
                            input_source_key = "document"
                        else:
                            raise ToolInputError(
                                f"Cannot determine input for op '{op_name}' for item {item_idx}."
                            )

                primary_input_value = item_state[input_source_key]

                # 2. Determine Primary Input Argument Name
                primary_param_map = {
                    "document_path": "document_path",
                    "image_path": "image_path",
                    "file_path": "file_path",
                    "document_data": "document_data",
                    "image_data": "image_data",
                    "text": "text",
                    "entities_input": "entities_input",
                    "texts": "texts",
                    "document": "document",
                    "content": "document",
                }
                primary_input_arg_name = primary_param_map.get(input_source_key)
                if not primary_input_arg_name:
                    # Inspect function signature to find the likely primary argument
                    try:
                        func_vars = op_func.__code__.co_varnames[: op_func.__code__.co_argcount]
                        primary_input_arg_name = (
                            "document"
                            if "document" in func_vars
                            else "text"
                            if "text" in func_vars
                            else func_vars[0]
                            if func_vars
                            else "input"
                        )
                    except (
                        AttributeError
                    ):  # Handle cases where introspection fails (e.g., built-ins)
                        primary_input_arg_name = "document"  # Default guess
                    logger.warning(
                        f"Assuming primary arg for op '{op_name}' is '{primary_input_arg_name}'."
                    )

                call_kwargs[primary_input_arg_name] = primary_input_value

                # 3. Handle Mapped Inputs
                if isinstance(op_input_map, dict):
                    for param_name, state_key in op_input_map.items():
                        if state_key not in item_state:
                            raise ToolInputError(
                                f"Mapped key '{state_key}' not found for item {item_idx}.",
                                param_name=state_key,
                            )
                        if param_name != primary_input_arg_name:
                            call_kwargs[param_name] = item_state[state_key]
                        elif call_kwargs[primary_input_arg_name] != item_state[state_key]:
                            logger.warning(
                                f"Mapped input '{param_name}' overrides primary input for item {item_idx}."
                            )
                            call_kwargs[primary_input_arg_name] = item_state[state_key]

                # 4. Add Fixed Params
                if isinstance(op_params, dict):
                    for p_name, p_value in op_params.items():
                        if p_name == primary_input_arg_name and p_name in call_kwargs:
                            logger.warning(
                                f"Fixed param '{p_name}' overrides dynamic input for item {item_idx}."
                            )
                        call_kwargs[p_name] = p_value

                # --- Execute Operation ---
                logger.debug(
                    f"Calling {op_name} for item {item_idx} with args: {list(call_kwargs.keys())}"
                )
                op_result = await op_func(**call_kwargs)  # Call the standalone function

                # --- Process Result ---
                if not isinstance(op_result, dict):
                    raise ToolError(
                        "INVALID_RESULT_FORMAT",
                        details={"operation": op_name, "result_type": type(op_result).__name__},
                    )
                item_state[op_output_key] = op_result  # Store full result

                # Promote output if requested
                if op_promote and isinstance(op_promote, str):
                    if op_promote in op_result:
                        item_state["content"] = op_result[op_promote]
                        logger.debug(f"Promoted '{op_promote}' to 'content' for item {item_idx}")
                    else:
                        logger.warning(
                            f"Cannot promote key '{op_promote}' for item {item_idx}: key not found in result."
                        )

                # Update status based on success flag
                if not op_result.get("success", False):
                    err_msg = op_result.get("error", f"Op '{op_name}' failed.")
                    err_code = op_result.get("error_code", "PROCESSING_ERROR")
                    log_entry = f"{step_label} Failed: [{err_code}] {err_msg}"
                    item_state["_error_log"].append(log_entry)
                    item_state["_status"] = "failed"
                    logger.warning(f"Op '{op_name}' failed for item {item_idx}: {err_msg}")
                elif item_state["_status"] != "failed":
                    item_state["_status"] = "processed"

            # --- Error Handling for Worker ---
            except ToolInputError as tie:
                error_msg = f"{step_label} Input Error: [{tie.error_code}] {str(tie)}"
                logger.error(f"{error_msg} for item {item_idx}", exc_info=False)
                item_state["_error_log"].append(error_msg)
                item_state["_status"] = "failed"
                item_state[op_output_key] = {
                    "error": str(tie),
                    "error_code": tie.error_code,
                    "success": False,
                }
            except ToolError as te:
                error_msg = f"{step_label} Tool Error: [{te.error_code}] {str(te)}"
                logger.error(f"{error_msg} for item {item_idx}", exc_info=True)
                item_state["_error_log"].append(error_msg)
                item_state["_status"] = "failed"
                item_state[op_output_key] = {
                    "error": str(te),
                    "error_code": te.error_code,
                    "success": False,
                }
            except Exception as e:
                error_msg = f"{step_label} Unexpected Error: {type(e).__name__}: {str(e)}"
                logger.error(f"{error_msg} for item {item_idx}", exc_info=True)
                item_state["_error_log"].append(error_msg)
                item_state["_status"] = "failed"
                item_state[op_output_key] = {
                    "error": str(e),
                    "error_type": type(e).__name__,
                    "success": False,
                }
            return item_state

    # --- Apply Operations Sequentially ---
    for op_index, op_spec in enumerate(operations):
        op_name = op_spec.get("operation")
        op_output_key = op_spec.get("output_key")
        op_params = op_spec.get("params", {})
        op_input_key = op_spec.get("input_key")
        op_input_map = op_spec.get("input_keys_map", {})
        op_promote = op_spec.get("promote_output")

        # --- Validate Operation Spec (robust checks) ---
        if not op_name or not isinstance(op_name, str) or op_name not in _OP_MAP:
            error_msg = f"Invalid/unknown operation '{op_name}' at step {op_index + 1}."
            logger.error(error_msg + " Skipping step for all items.")
            for item_state in results_state:
                if item_state["_status"] != "failed":
                    item_state["_error_log"].append(error_msg + " (Skipped)")
                    item_state["_status"] = "failed"
            continue
        if not op_output_key or not isinstance(op_output_key, str):
            error_msg = f"Missing/invalid 'output_key' for '{op_name}' at step {op_index + 1}."
            logger.error(error_msg + " Skipping step for all items.")
            for item_state in results_state:
                if item_state["_status"] != "failed":
                    item_state["_error_log"].append(error_msg + " (Skipped)")
                    item_state["_status"] = "failed"
            continue
        if not isinstance(op_params, dict):
            error_msg = f"Invalid 'params' (must be dict) for '{op_name}' at step {op_index + 1}."
            logger.error(error_msg + " Skipping step for all items.")
            for item_state in results_state:
                if item_state["_status"] != "failed":
                    item_state["_error_log"].append(error_msg + " (Skipped)")
                    item_state["_status"] = "failed"
            continue

        # Get the actual function from the map
        current_op_func = _OP_MAP[op_name]
        current_step_label = f"Step {op_index + 1}/{len(operations)}: '{op_name}'"
        logger.info(f"--- Starting {current_step_label} (Concurrency: {max_concurrency}) ---")

        # --- Run Tasks for Current Step ---
        step_semaphore = asyncio.Semaphore(max_concurrency)
        step_tasks = [
            # Call the single worker function, passing the current loop's values
            _apply_op_to_item_worker(
                item_state=item_state,
                semaphore=step_semaphore,
                op_func=current_op_func,
                op_name=op_name,
                step_label=current_step_label,
                op_output_key=op_output_key,
                op_input_key=op_input_key,
                op_params=op_params,
                op_input_map=op_input_map,
                op_promote=op_promote,
            )
            for item_state in results_state
        ]
        updated_states = await asyncio.gather(*step_tasks)
        results_state = updated_states  # Update the main state list

        # Log summary after step
        step_processed_count = sum(1 for s in results_state if s.get("_status") == "processed")
        step_fail_count = sum(1 for s in results_state if s.get("_status") == "failed")
        logger.info(
            f"--- Finished {current_step_label} (Processed: {step_processed_count}, Failed: {step_fail_count}) ---"
        )

    # --- Final Cleanup ---
    final_results = []
    for item_state in results_state:
        final_item = item_state.copy()
        final_item.pop("_original_index", None)
        # Keep _status and _error_log for visibility
        final_results.append(final_item)

    logger.info(f"Batch processing finished for {len(inputs)} items.")
    return final_results
