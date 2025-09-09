"""
Word文档解析工具，用于解析Word文档内容
"""

import os
import traceback
import io
import base64
import imghdr
import tempfile
import subprocess
import shutil
from typing import Dict, List, Any, Tuple, Optional
import docx
from docx.document import Document
from docx.parts.document import DocumentPart
from docx.package import Package
import mcp.types as types
from . import BaseTool, ToolRegistry

@ToolRegistry.register
class WordTool(BaseTool):
    """
    用于解析Word文档的工具，提取文本内容、表格和图片信息
    支持.docx和.doc(Word 97-2003)格式
    """
    
    name = "parse_word"
    description = "解析Word文档内容，提取文本、表格和图片信息"
    input_schema = {
        "type": "object",
        "required": ["file_path"],
        "properties": {
            "file_path": {
                "type": "string",
                "description": "Word文档的本地路径，例如'/path/to/document.docx'",
            }
        },
    }
    
    async def execute(self, arguments: Dict[str, Any]) -> List[types.TextContent | types.ImageContent | types.EmbeddedResource]:
        """
        解析Word文档
        
        Args:
            arguments: 参数字典，必须包含'file_path'键
            
        Returns:
            解析结果列表
        """
        if "file_path" not in arguments:
            return [types.TextContent(
                type="text",
                text="错误: 缺少必要参数 'file_path'"
            )]
        
        # 处理文件路径，支持挂载目录的转换
        file_path = self.process_file_path(arguments["file_path"])
        
        return await self._parse_word_document(file_path)
    
    def _get_image_mime_type(self, image_bytes: bytes) -> str:
        """
        获取图片的MIME类型
        """
        image_type = imghdr.what(None, image_bytes)
        if image_type:
            return f"image/{image_type}"
        return "image/png"  # 默认返回PNG类型
    
    def _encode_image_base64(self, image_bytes: bytes) -> str:
        """
        将图片编码为base64格式
        """
        return base64.b64encode(image_bytes).decode('utf-8')
    
    def _is_valid_image(self, image_bytes: bytes) -> bool:
        """
        检查数据是否为有效的图片
        
        Args:
            image_bytes: 图片二进制数据
            
        Returns:
            是否为有效图片
        """
        # 检查常见图片格式的文件头特征
        if len(image_bytes) < 12:
            return False  # 文件太小，不可能是有效图片
            
        # 使用imghdr识别图片类型
        image_type = imghdr.what(None, image_bytes)
        if not image_type:
            return False
            
        # 进一步验证常见图片格式的文件头特征
        file_signatures = {
            'jpeg': [bytes([0xFF, 0xD8, 0xFF])],  # JPEG
            'png': [bytes([0x89, 0x50, 0x4E, 0x47, 0x0D, 0x0A, 0x1A, 0x0A])],  # PNG
            'gif': [bytes([0x47, 0x49, 0x46, 0x38, 0x37, 0x61]), bytes([0x47, 0x49, 0x46, 0x38, 0x39, 0x61])],  # GIF
            'bmp': [bytes([0x42, 0x4D])],  # BMP
            'webp': [bytes([0x52, 0x49, 0x46, 0x46]) + b'....WEBP'],  # WEBP (使用模式匹配)
        }
        
        # 检查文件头是否匹配任何已知图片格式
        if image_type in file_signatures:
            for signature in file_signatures[image_type]:
                if len(signature) <= len(image_bytes):
                    # 对于WEBP这种需要模式匹配的格式特殊处理
                    if image_type == 'webp':
                        if image_bytes.startswith(b'RIFF') and b'WEBP' in image_bytes[0:12]:
                            return True
                    # 直接比较字节序列
                    elif image_bytes.startswith(signature):
                        return True
            return False
        
        # 未知格式但imghdr认为是图片，需要更严格的验证
        try:
            from PIL import Image
            img = Image.open(io.BytesIO(image_bytes))
            img.verify()  # 验证图片完整性
            return True
        except Exception:
            return False
    
    def _extract_images_from_word(self, doc: Document) -> List[Tuple[str, bytes]]:
        """
        从Word文档中提取图片，过滤掉嵌入的外部文档
        
        Args:
            doc: Word文档对象
            
        Returns:
            图片列表，每项包含图片ID和二进制数据
        """
        images = []
        document_part = doc.part
        rels = document_part.rels
        
        for rel in rels.values():
            try:
                # 只处理图片类型的关系
                if "image" in rel.reltype:
                    image_part = rel.target_part
                    image_bytes = image_part.blob
                    image_id = rel.rId
                    
                    # 验证是否为真实图片，过滤掉嵌入的外部文档
                    if self._is_valid_image(image_bytes):
                        images.append((image_id, image_bytes))
            except Exception:
                continue
                    
        return images
    
    def _is_libreoffice_installed(self) -> bool:
        """
        检查系统是否安装了LibreOffice
        
        Returns:
            布尔值，表示LibreOffice是否可用
        """
        try:
            # 尝试执行LibreOffice命令，检查是否安装
            result = subprocess.run(
                ["soffice", "--version"], 
                stdout=subprocess.PIPE, 
                stderr=subprocess.PIPE,
                timeout=5  # 设置超时，避免长时间等待
            )
            return result.returncode == 0
        except (subprocess.SubprocessError, FileNotFoundError, OSError):
            return False
    
    def _convert_doc_to_docx(self, file_path: str) -> Optional[str]:
        """
        使用LibreOffice将.doc文件转换为.docx格式
        
        Args:
            file_path: .doc文档路径
            
        Returns:
            转换后的.docx文件路径，如果转换失败则返回None
            
        Raises:
            Exception: 当转换过程中出现错误时
        """
        # 创建临时目录用于存放转换结果
        temp_dir = tempfile.mkdtemp(prefix="word_convert_")
        
        try:
            # 获取文件名（不含路径）
            file_name = os.path.basename(file_path)
            # 基础文件名（不含扩展名）
            base_name = os.path.splitext(file_name)[0]
            # 预期的转换结果文件路径
            output_docx = os.path.join(temp_dir, f"{base_name}.docx")
            
            # 调用LibreOffice进行格式转换
            process = subprocess.run(
                [
                    "soffice",
                    "--headless",
                    "--convert-to", "docx",
                    "--outdir", temp_dir,
                    file_path
                ],
                capture_output=True,
                text=True,
                timeout=60  # 设置一分钟超时
            )
            
            # 检查转换是否成功
            if process.returncode != 0:
                error_msg = process.stderr.strip() or "未知错误"
                raise Exception(f"LibreOffice转换失败: {error_msg}")
            
            # 验证输出文件是否存在
            if not os.path.exists(output_docx):
                raise Exception("转换后的文件不存在")
                
            # 检查文件大小，确保不是空文件
            if os.path.getsize(output_docx) == 0:
                raise Exception("转换结果为空文件")
            
            return output_docx
            
        except subprocess.TimeoutExpired:
            raise Exception("文档转换超时，可能是文档过大或复杂")
        except Exception as e:
            raise Exception(f"转换.doc文件时出错: {str(e)}")
        
    async def _parse_word_document(self, file_path: str) -> List[types.TextContent | types.ImageContent | types.EmbeddedResource]:
        """
        解析Word文档内容，支持.docx和.doc格式
        
        Args:
            file_path: Word文档路径
            
        Returns:
            Word文档内容列表
        """
        results = []
        temp_docx_path = None
        
        # 检查文件是否存在
        if not os.path.exists(file_path):
            return [types.TextContent(
                type="text",
                text=f"错误: 文件不存在: {file_path}\n请检查路径是否正确，并确保文件可访问。"
            )]
        
        # 检查文件扩展名
        if not file_path.lower().endswith(('.docx', '.doc')):
            return [types.TextContent(
                type="text",
                text=f"错误: 不支持的文件格式: {file_path}\n仅支持.docx和.doc格式的Word文档。"
            )]
        
        try:
            # 添加文件信息
            file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
            
            # 处理.doc格式（Word 97-2003文档）
            if file_path.lower().endswith('.doc'):
                results.append(types.TextContent(
                    type="text",
                    text=f"# Word文档解析 (Word 97-2003 格式)\n\n文件大小: {file_size_mb:.2f} MB"
                ))
                
                # 检查LibreOffice是否可用
                if not self._is_libreoffice_installed():
                    return [types.TextContent(
                        type="text",
                        text="错误: 无法解析Word 97-2003 (.doc)格式。\n"
                             "系统未安装LibreOffice，无法进行格式转换。\n"
                             "请安装LibreOffice后重试，或将文档另存为.docx格式。"
                    )]
                
                try:
                    # 显示转换提示
                    results.append(types.TextContent(
                        type="text",
                        text="正在使用LibreOffice转换文档格式，请稍候..."
                    ))
                    
                    # 转换.doc到.docx
                    temp_docx_path = self._convert_doc_to_docx(file_path)
                    
                    # 更新文件路径为转换后的文件
                    file_path = temp_docx_path
                    
                    results.append(types.TextContent(
                        type="text",
                        text="文档格式转换完成，继续解析...\n"
                    ))
                except Exception as e:
                    return results + [types.TextContent(
                        type="text",
                        text=f"错误: {str(e)}\n"
                             f"建议:\n"
                             f"1. 确保已正确安装LibreOffice且可通过命令行访问\n"
                             f"2. 尝试手动将文档转换为.docx格式后重试\n"
                             f"3. 检查文档是否加密或损坏"
                    )]
            else:
                results.append(types.TextContent(
                    type="text",
                    text=f"# Word文档解析\n\n文件大小: {file_size_mb:.2f} MB"
                ))
            
            # 打开Word文档
            doc = docx.Document(file_path)
            
            # 提取文档属性
            properties = {}
            if hasattr(doc.core_properties, 'title') and doc.core_properties.title:
                properties['标题'] = doc.core_properties.title
            if hasattr(doc.core_properties, 'author') and doc.core_properties.author:
                properties['作者'] = doc.core_properties.author
            if hasattr(doc.core_properties, 'created') and doc.core_properties.created:
                properties['创建时间'] = str(doc.core_properties.created)
            if hasattr(doc.core_properties, 'modified') and doc.core_properties.modified:
                properties['修改时间'] = str(doc.core_properties.modified)
            if hasattr(doc.core_properties, 'comments') and doc.core_properties.comments:
                properties['备注'] = doc.core_properties.comments
            
            # 添加文档属性信息
            if properties:
                properties_text = "## 文档属性\n\n"
                for key, value in properties.items():
                    properties_text += f"- {key}: {value}\n"
                results.append(types.TextContent(
                    type="text",
                    text=properties_text
                ))
            
            # 提取文档内容
            content_text = "## 文档内容\n\n"
            
            # 处理段落
            paragraphs_count = len(doc.paragraphs)
            content_text += f"### 段落 (共{paragraphs_count}个)\n\n"
            
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():  # 只处理非空段落
                    content_text += f"{para.text}\n\n"
            
            # 处理表格
            tables_count = len(doc.tables)
            if tables_count > 0:
                content_text += f"### 表格 (共{tables_count}个)\n\n"
                
                for i, table in enumerate(doc.tables):
                    content_text += f"#### 表格 {i+1}\n\n"
                    
                    # 创建Markdown表格
                    rows = []
                    for row in table.rows:
                        cells = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
                        rows.append(cells)
                    
                    if rows:
                        # 表头
                        content_text += "| " + " | ".join(rows[0]) + " |\n"
                        # 分隔线
                        content_text += "| " + " | ".join(["---"] * len(rows[0])) + " |\n"
                        # 表格内容
                        for row in rows[1:]:
                            content_text += "| " + " | ".join(row) + " |\n"
                        
                        content_text += "\n"
            
            # 添加文档内容
            results.append(types.TextContent(
                type="text",
                text=content_text
            ))
            
            # 提取图片信息和内容
            try:
                # 提取文档中的所有图片，并过滤掉嵌入的外部文档
                images = self._extract_images_from_word(doc)
                
                if images:
                    image_info = f"## 图片信息\n\n文档中包含 {len(images)} 张图片。\n\n"
                    results.append(types.TextContent(
                        type="text",
                        text=image_info
                    ))
                    
                    # 返回图片内容
                    for i, (image_id, image_bytes) in enumerate(images):
                        try:
                            # 获取图片MIME类型
                            mime_type = self._get_image_mime_type(image_bytes)
                            
                            # 将图片添加到结果中
                            image_base64 = self._encode_image_base64(image_bytes)
                            results.append(types.TextContent(
                                type="text",
                                text=f"### 图片 {i+1}\n\n"
                            ))
                            results.append(types.ImageContent(
                                type="image",
                                data=image_base64,
                                mimeType=mime_type
                            ))
                        except Exception as e:
                            # 记录图片处理错误但不中断
                            results.append(types.TextContent(
                                type="text",
                                text=f"注意: 图片 {i+1} 处理失败: {str(e)}"
                            ))
                else:
                    results.append(types.TextContent(
                        type="text",
                        text="## 图片信息\n\n文档中未包含图片或嵌入对象均不是有效图片。"
                    ))
            except Exception as img_error:
                results.append(types.TextContent(
                    type="text",
                    text=f"警告: 提取图片信息时出错: {str(img_error)}"
                ))
            
            # 添加处理完成的提示
            results.append(types.TextContent(
                type="text",
                text="Word文档处理完成！"
            ))
            
            return results
        except Exception as e:
            error_details = traceback.format_exc()
            return [types.TextContent(
                type="text",
                text=f"错误: 解析Word文档失败: {str(e)}\n"
                     f"可能的原因:\n"
                     f"1. 文件格式不兼容或已损坏\n"
                     f"2. 文件受密码保护\n"
                     f"3. 文件包含不支持的内容\n\n"
                     f"详细错误信息: {error_details}"
            )]
        finally:
            # 清理临时文件
            if temp_docx_path and os.path.exists(temp_docx_path):
                try:
                    # 删除临时文件
                    temp_dir = os.path.dirname(temp_docx_path)
                    shutil.rmtree(temp_dir, ignore_errors=True)
                except Exception:
                    # 忽略清理过程中的错误
                    pass 