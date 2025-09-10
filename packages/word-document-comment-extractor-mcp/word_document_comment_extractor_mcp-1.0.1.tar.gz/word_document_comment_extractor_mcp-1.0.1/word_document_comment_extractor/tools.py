"""
Word文档评论提取工具函数

提供三个核心的文档评论提取功能
"""

import os
import json
from typing import Dict, List, Optional, Any
from .utils import (
    ensure_docx_extension,
    extract_all_comments,
    filter_comments_by_author,
    get_comments_for_paragraph_by_index
)


async def get_all_comments(filename: str) -> str:
    """
    Extract all comments from a Word document.
    
    Args:
        filename: Path to the Word document
        
    Returns:
        JSON string containing all comments with metadata
    """
    from docx import Document
    
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        raise FileNotFoundError(f'Document {filename} does not exist')
    
    try:
        # Load the document
        doc = Document(filename)
        
        # Extract all comments
        comments = extract_all_comments(doc)
        
        # Return results
        return json.dumps({
            'success': True,
            'comments': comments,
            'total_comments': len(comments)
        }, indent=2)
        
    except Exception as e:
        return json.dumps({
            'success': False,
            'error': f'Failed to extract comments: {str(e)}'
        }, indent=2)


async def get_comments_by_author(filename: str, author: str) -> str:
    """
    Extract comments from a specific author in a Word document.
    
    Args:
        filename: Path to the Word document
        author: Name of the comment author to filter by
        
    Returns:
        JSON string containing filtered comments
    """
    from docx import Document
    
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        raise FileNotFoundError(f'Document {filename} does not exist')

    if not author or not author.strip():
        raise ValueError('Author name cannot be empty')
    
    try:
        # Load the document
        doc = Document(filename)
        
        # Extract all comments
        all_comments = extract_all_comments(doc)
        
        # Filter by author
        author_comments = filter_comments_by_author(all_comments, author)
        
        # Return results
        return json.dumps({
            'success': True,
            'author': author,
            'comments': author_comments,
            'total_comments': len(author_comments)
        }, indent=2)
        
    except Exception as e:
        return json.dumps({
            'success': False,
            'error': f'Failed to extract comments: {str(e)}'
        }, indent=2)


async def get_comments_for_paragraph(filename: str, paragraph_index: int) -> str:
    """
    Extract comments for a specific paragraph in a Word document.
    
    Args:
        filename: Path to the Word document
        paragraph_index: Index of the paragraph (0-based)
        
    Returns:
        JSON string containing comments for the specified paragraph
    """
    from docx import Document
    
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        raise FileNotFoundError(f'Document {filename} does not exist')

    if paragraph_index < 0:
        raise ValueError('Paragraph index must be non-negative')
    
    try:
        # Load the document
        doc = Document(filename)
        
        # Check if paragraph index is valid
        if paragraph_index >= len(doc.paragraphs):
            return json.dumps({
                'success': False,
                'error': f'Paragraph index {paragraph_index} is out of range. Document has {len(doc.paragraphs)} paragraphs.'
            }, indent=2)
        
        # Extract all comments
        all_comments = extract_all_comments(doc)
        
        # Filter for the specific paragraph
        para_comments = get_comments_for_paragraph_by_index(all_comments, paragraph_index)
        
        # Get the paragraph text for context
        paragraph_text = doc.paragraphs[paragraph_index].text
        
        # Return results
        return json.dumps({
            'success': True,
            'paragraph_index': paragraph_index,
            'paragraph_text': paragraph_text,
            'comments': para_comments,
            'total_comments': len(para_comments)
        }, indent=2)
        
    except Exception as e:
        return json.dumps({
            'success': False,
            'error': f'Failed to extract comments: {str(e)}'
        }, indent=2)
