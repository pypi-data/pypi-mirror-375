"""Advanced audio transcription and enhancement tools for Ultimate MCP Server.

This module provides tools for high-quality audio transcription, pre-processing,
and intelligent transcript enhancement with advanced features like speaker 
diarization, custom vocabulary support, and semantic structuring.
"""
import asyncio
import concurrent.futures
import datetime
import json
import os
import re
import shutil
import subprocess
import tempfile
import time
from dataclasses import dataclass
from enum import Enum
from pathlib import Path
from typing import Any, Dict, List, Optional

import aiofiles
import httpx
from docx import Document
from pydantic import BaseModel, Field
from pydantic.functional_validators import field_validator

from ultimate_mcp_server.constants import Provider, TaskType
from ultimate_mcp_server.core.providers.base import get_provider
from ultimate_mcp_server.exceptions import (
    ProviderError,
    ResourceError,
    ToolError,
    ToolInputError,
)
from ultimate_mcp_server.services.cache import with_cache
from ultimate_mcp_server.tools.base import with_error_handling, with_retry, with_tool_metrics
from ultimate_mcp_server.tools.completion import chat_completion, generate_completion
from ultimate_mcp_server.utils import get_logger
from ultimate_mcp_server.utils.text import count_tokens

logger = get_logger("ultimate_mcp_server.tools.audio")

# --- Constants and Enums ---

class AudioEnhancementProfile(str, Enum):
    """Predefined audio enhancement profiles for different recording types."""
    CONFERENCE_CALL = "conference_call"  # Optimized for online meetings
    INTERVIEW = "interview"  # Optimized for interview recordings
    LECTURE = "lecture"  # Optimized for lectures/presentations
    NOISY = "noisy"  # Heavy noise reduction for noisy environments
    PHONE_CALL = "phone_call"  # Optimized for telephone audio
    VOICEMAIL = "voicemail"  # Optimized for voicemail recordings
    CUSTOM = "custom"  # User-defined settings

# Expected model sizes in bytes, with some tolerance (minimum size)
WHISPER_MODEL_SIZES = {
    "large-v3": 2900000000, # ~2.9GB
    "large-v3-turbo": 1500000000, # ~1.5GB
}

class TranscriptionQuality(str, Enum):
    """Quality settings for transcription, balancing speed vs accuracy."""
    DRAFT = "draft"  # Fastest, less accurate
    STANDARD = "standard"  # Balanced speed/accuracy
    ENHANCED = "enhanced"  # More accurate, slower
    MAXIMUM = "maximum"  # Most accurate, slowest


class EnhancementStyle(str, Enum):
    """Transcript enhancement styles for different use cases."""
    RAW = "raw"  # No enhancement, just cleaned
    READABLE = "readable"  # Basic readability improvements
    POLISHED = "polished"  # Well-formatted with proper punctuation
    VERBATIM = "verbatim"  # Preserve all speech patterns, hesitations
    STRUCTURED = "structured"  # Add semantic structure (paragraphs, sections)


class OutputFormat(str, Enum):
    """Available output formats for transcripts."""
    JSON = "json"  # Full JSON with all metadata
    TEXT = "text"  # Plain text
    SRT = "srt"  # SubRip subtitle format
    VTT = "vtt"  # WebVTT subtitle format
    DOCX = "docx"  # Microsoft Word format
    MARKDOWN = "markdown"  # Markdown with formatting


# --- Schema Validation Models ---

class AudioEnhancementParams(BaseModel):
    """Parameters for audio enhancement."""
    profile: AudioEnhancementProfile = Field(
        default=AudioEnhancementProfile.CONFERENCE_CALL,
        description="Predefined audio enhancement profile"
    )
    volume: float = Field(
        default=1.5, 
        ge=0.1, 
        le=5.0,
        description="Volume adjustment factor"
    )
    noise_reduction: int = Field(
        default=10, 
        ge=0, 
        le=30,
        description="Noise reduction strength (0-30)"
    )
    highpass: int = Field(
        default=200, 
        ge=50, 
        le=500,
        description="Highpass filter frequency in Hz"
    )
    lowpass: int = Field(
        default=3000, 
        ge=1000, 
        le=20000,
        description="Lowpass filter frequency in Hz"
    )
    normalize: bool = Field(
        default=True,
        description="Apply dynamic audio normalization"
    )
    compression: bool = Field(
        default=True,
        description="Apply dynamic range compression"
    )
    dereverberation: bool = Field(
        default=False,
        description="Apply dereverberation filter"
    )
    custom_filters: Optional[str] = Field(
        default=None,
        description="Additional custom FFmpeg filters"
    )
    output_channels: int = Field(
        default=2, 
        ge=1, 
        le=2,
        description="Number of output channels (1=mono, 2=stereo)"
    )
    output_sample_rate: int = Field(
        default=16000, 
        ge=8000, 
        le=48000,
        description="Output sample rate in Hz"
    )
    
    @field_validator('custom_filters')
    def validate_custom_filters(cls, v):
        """Validate that custom filters don't contain dangerous commands."""
        if v:
            # Check for shell escape attempts
            dangerous_patterns = [';', '&&', '||', '`', '$', '\\', '>', '<', '|', '*', '?', '~', '#']
            for pattern in dangerous_patterns:
                if pattern in v:
                    raise ValueError(f"Custom filter contains disallowed character: {pattern}")
        return v


class WhisperParams(BaseModel):
    """Parameters for Whisper transcription."""
    model: str = Field(
        default="large-v3-turbo",
        description="Whisper model name"
    )
    language: Optional[str] = Field(
        default=None,
        description="Language code (auto-detect if None)"
    )
    quality: TranscriptionQuality = Field(
        default=TranscriptionQuality.STANDARD,
        description="Transcription quality level"
    )
    beam_size: int = Field(
        default=5, 
        ge=1, 
        le=10,
        description="Beam size for beam search"
    )
    processors: int = Field(
        default=2, 
        ge=1, 
        le=8,
        description="Number of processors to use"
    )
    word_timestamps: bool = Field(
        default=True,
        description="Generate timestamps for each word"
    )
    translate: bool = Field(
        default=False,
        description="Translate non-English to English"
    )
    diarize: bool = Field(
        default=False,
        description="Attempt speaker diarization"
    )
    highlight_words: bool = Field(
        default=False,
        description="Highlight words with lower confidence"
    )
    max_context: int = Field(
        default=-1, 
        ge=-1,
        description="Maximum number of text tokens to consider from previous history"
    )
    custom_vocab: Optional[List[str]] = Field(
        default=None,
        description="Custom vocabulary terms to improve recognition"
    )


class TranscriptEnhancementParams(BaseModel):
    """Parameters for transcript enhancement."""
    style: EnhancementStyle = Field(
        default=EnhancementStyle.READABLE,
        description="Enhancement style"
    )
    provider: str = Field(
        default=Provider.ANTHROPIC.value,
        description="LLM provider for enhancement"
    )
    model: Optional[str] = Field(
        default=None,
        description="Specific model to use (provider default if None)"
    )
    identify_speakers: bool = Field(
        default=False,
        description="Attempt to identify and label speakers"
    )
    add_paragraphs: bool = Field(
        default=True,
        description="Add paragraph breaks at natural points"
    )
    fix_spelling: bool = Field(
        default=True,
        description="Fix spelling errors"
    )
    fix_grammar: bool = Field(
        default=True,
        description="Fix basic grammatical errors"
    )
    sections: bool = Field(
        default=False,
        description="Add section headings based on topic changes"
    )
    max_chunk_size: int = Field(
        default=6500, 
        ge=1000, 
        le=100000,
        description="Maximum chunk size in characters"
    )
    format_numbers: bool = Field(
        default=True,
        description="Format numbers consistently (e.g., '25' instead of 'twenty-five')"
    )
    custom_instructions: Optional[str] = Field(
        default=None,
        description="Additional custom instructions for enhancement"
    )


class TranscriptionOptions(BaseModel):
    """Complete options for audio transcription."""
    enhance_audio: bool = Field(
        default=True,
        description="Whether to preprocess audio with FFmpeg"
    )
    enhance_transcript: bool = Field(
        default=True,
        description="Whether to enhance the transcript with LLM"
    )
    parallel_processing: bool = Field(
        default=True,
        description="Process chunks in parallel when possible"
    )
    max_workers: int = Field(
        default=4, 
        ge=1,
        description="Maximum number of parallel workers"
    )
    output_formats: List[OutputFormat] = Field(
        default=[OutputFormat.JSON, OutputFormat.TEXT],
        description="Output formats to generate"
    )
    save_enhanced_audio: bool = Field(
        default=False,
        description="Save the enhanced audio file"
    )
    keep_artifacts: bool = Field(
        default=False,
        description="Keep temporary files and artifacts"
    )
    audio_params: AudioEnhancementParams = Field(
        default_factory=AudioEnhancementParams,
        description="Audio enhancement parameters"
    )
    whisper_params: WhisperParams = Field(
        default_factory=WhisperParams,
        description="Whisper transcription parameters"
    )
    enhancement_params: TranscriptEnhancementParams = Field(
        default_factory=TranscriptEnhancementParams,
        description="Transcript enhancement parameters"
    )
    language_detection: bool = Field(
        default=False,  # Disable language detection by default
        description="Automatically detect language before transcription"
    )


class Segment(BaseModel):
    """A segment of transcript with timing information."""
    start: float = Field(..., description="Start time in seconds")
    end: float = Field(..., description="End time in seconds")
    text: str = Field(..., description="Segment text")
    speaker: Optional[str] = Field(None, description="Speaker identifier")
    words: Optional[List[Dict[str, Any]]] = Field(None, description="Word-level data")
    confidence: Optional[float] = Field(None, description="Confidence score")


class AudioInfo(BaseModel):
    """Audio file information."""
    duration: float = Field(..., description="Duration in seconds")
    channels: int = Field(..., description="Number of audio channels")
    sample_rate: int = Field(..., description="Sample rate in Hz")
    format: str = Field(..., description="Audio format")
    codec: Optional[str] = Field(None, description="Audio codec")
    bit_depth: Optional[int] = Field(None, description="Bit depth")
    bitrate: Optional[int] = Field(None, description="Bitrate in bits/second")
    size_bytes: Optional[int] = Field(None, description="File size in bytes")


# --- Data Classes ---

@dataclass
class ProcessingContext:
    """Context for the transcription process."""
    file_path: str
    temp_dir: str
    original_filename: str
    base_filename: str
    options: TranscriptionOptions
    enhanced_audio_path: Optional[str] = None
    processing_times: Dict[str, float] = None
    language_code: Optional[str] = None
    
    def __post_init__(self):
        if self.processing_times is None:
            self.processing_times = {}


# --- Tool Functions ---

@with_cache(ttl=24 * 60 * 60)  # Cache results for 24 hours
@with_tool_metrics
@with_retry(max_retries=1, retry_delay=1.0)
@with_error_handling
async def transcribe_audio(
    file_path: str,
    options: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    """Transcribes an audio file to text with advanced preprocessing and enhancement.
    
    This tool performs a multi-stage process:
    1. Analyzes the audio file to determine optimal processing parameters
    2. Enhances audio quality with adaptive filtering and preprocessing 
    3. Performs high-quality transcription with customizable settings
    4. Intelligently enhances and structures the transcript for readability
    5. Optionally identifies speakers and adds semantic structure
    
    Args:
        file_path: Path to the input audio file (.mp3, .m4a, .wav, etc.)
        options: Optional dictionary with transcription options including:
            - enhance_audio: Whether to preprocess audio (default True)
            - enhance_transcript: Whether to enhance transcript (default True)
            - parallel_processing: Process chunks in parallel (default True)
            - output_formats: List of output formats (default ["json", "text"])
            - audio_params: Audio enhancement parameters
            - whisper_params: Whisper transcription parameters 
            - enhancement_params: Transcript enhancement parameters
    
    Returns:
        A dictionary containing:
        {
            "raw_transcript": "Original unmodified transcript from Whisper",
            "enhanced_transcript": "LLM-enhanced transcript with improved formatting",
            "segments": [
                {
                    "start": 0.0,
                    "end": 10.5,
                    "text": "Segment text content", 
                    "speaker": "Speaker 1",  # If speaker diarization is enabled
                    "words": [...]  # Word-level data if available
                },
                ...
            ],
            "metadata": {
                "language": "en",
                "duration": 120.5,
                "title": "Automatically detected title",
                "topics": ["Topic 1", "Topic 2"]  # If topic extraction is enabled
            },
            "audio_info": {
                "duration": 120.5,
                "channels": 2,
                "sample_rate": 44100,
                "format": "wav",
                "codec": "pcm_s16le",
                "bit_depth": 16,
                "bitrate": 1411000,
                "size_bytes": 31000000
            },
            "processing_time": {
                "audio_analysis": 0.5,
                "audio_enhancement": 5.2,
                "language_detection": 1.1,
                "transcription": 45.3,
                "transcript_enhancement": 10.2,
                "total": 62.3
            },
            "artifacts": {
                "enhanced_audio": "/path/to/enhanced.wav",  # If save_enhanced_audio is True
                "output_files": {
                    "json": "/path/to/transcript.json",
                    "text": "/path/to/transcript.txt",
                    "srt": "/path/to/transcript.srt"
                }
            },
            "tokens": {
                "input": 5000,
                "output": 3200,
                "total": 8200
            },
            "cost": 0.00185,
            "success": true
        }
    
    Raises:
        ToolInputError: If the file path is invalid or unsupported
        ToolError: If transcription or enhancement fails
        ResourceError: If required dependencies are not available
    """
    # Start timing total processing
    start_time = time.time()
    
    # --- Input Validation ---
    try:
        if not file_path or not isinstance(file_path, str):
            raise ToolInputError("File path must be a non-empty string.")
        
        file_path = os.path.abspath(os.path.expanduser(file_path))
        if not os.path.exists(file_path):
            raise ToolInputError(f"File not found: {file_path}")
        
        if not os.access(file_path, os.R_OK):
            raise ToolInputError(f"File not readable: {file_path}")
            
        # Validate file is an audio file
        _, ext = os.path.splitext(file_path)
        if ext.lower() not in ['.mp3', '.wav', '.m4a', '.mp4', '.flac', '.ogg', '.aac', '.wma', '.opus']:
            raise ToolInputError(f"Unsupported file format: {ext}. Please provide an audio file.")
        
        # Parse and validate options
        parsed_options = parse_options(options or {})
        
    except Exception as e:
        if isinstance(e, ToolInputError):
            raise
        raise ToolInputError(f"Failed to validate input: {str(e)}") from e
    
    # --- Initialize Processing Context ---
    try:
        temp_dir = tempfile.mkdtemp(prefix="llm_audio_")
        original_filename = os.path.basename(file_path)
        base_filename = os.path.splitext(original_filename)[0]
        
        context = ProcessingContext(
            file_path=file_path,
            temp_dir=temp_dir,
            original_filename=original_filename,
            base_filename=base_filename,
            options=parsed_options,
        )
        
        logger.info(
            f"Starting audio transcription process for {original_filename}",
            emoji_key="audio",
            temp_dir=temp_dir
        )
    except Exception as e:
        raise ToolError(f"Failed to initialize processing context: {str(e)}") from e
    
    try:
        # --- Check Dependencies ---
        await check_dependencies(context)
        
        # --- Process Audio ---
        result = await process_audio_file(context)
        
        # --- Calculate Total Time ---
        total_time = time.time() - start_time
        context.processing_times["total"] = total_time
        result["processing_time"] = context.processing_times
        
        logger.success(
            f"Audio transcription completed in {total_time:.2f}s",
            emoji_key="success",
            file=context.original_filename,
            duration=result.get("audio_info", {}).get("duration", 0)
        )
        
        return result
    
    except Exception as e:
        logger.error(
            f"Audio transcription failed: {str(e)}",
            emoji_key="error",
            exc_info=True,
            file=context.original_filename
        )
        # Clean up temporary directory unless keep_artifacts is True
        if context.options.keep_artifacts:
            logger.info(f"Keeping artifacts in {temp_dir}")
        else:
            try:
                shutil.rmtree(temp_dir)
            except Exception as cleanup_err:
                logger.warning(f"Failed to clean up temporary directory: {cleanup_err}")
        
        if isinstance(e, (ToolError, ToolInputError, ResourceError)):
            raise
        
        # Return a structured error response instead of just raising an exception
        return {
            "raw_transcript": "",
            "enhanced_transcript": "",
            "segments": [],
            "metadata": {},
            "audio_info": {},
            "processing_time": context.processing_times if hasattr(context, "processing_times") else {},
            "success": False,
            "error": f"Audio transcription failed: {str(e)}"
        }
    
    finally:
        # Clean up temporary directory unless keep_artifacts is True
        if not context.options.keep_artifacts:
            try:
                shutil.rmtree(temp_dir)
            except Exception as cleanup_err:
                logger.warning(f"Failed to clean up temporary directory: {cleanup_err}")


# --- Main Processing Functions ---

async def process_audio_file(context: ProcessingContext) -> Dict[str, Any]:
    """Process an audio file through the complete transcription pipeline."""
    # Get detailed audio information
    audio_analysis_start = time.time()
    audio_info = await get_detailed_audio_info(context.file_path)
    context.processing_times["audio_analysis"] = time.time() - audio_analysis_start
    
    logger.info(
        f"Audio analysis: {audio_info.get('duration', 0):.1f}s duration, "
        f"{audio_info.get('sample_rate', 0)} Hz, "
        f"{audio_info.get('channels', 0)} channels", 
        emoji_key="audio"
    )
    
    # Update parameters based on audio analysis if needed
    _update_parameters_from_audio_info(context, audio_info)
    
    # --- Audio Enhancement ---
    enhanced_audio_path = context.file_path
    if context.options.enhance_audio:
        audio_enhance_start = time.time()
        logger.info("Enhancing audio quality", emoji_key="audio", profile=context.options.audio_params.profile.value)
        
        enhanced_audio_path = await enhance_audio(context, audio_info)
        if not enhanced_audio_path:
            logger.warning("Audio enhancement failed, falling back to original file", emoji_key="warning")
            enhanced_audio_path = context.file_path
            
        context.processing_times["audio_enhancement"] = time.time() - audio_enhance_start
    else:
        context.processing_times["audio_enhancement"] = 0
        
    context.enhanced_audio_path = enhanced_audio_path
    
    if not os.path.exists(enhanced_audio_path):
        logger.warning(f"Enhanced audio path does not exist: {enhanced_audio_path}", emoji_key="warning")
        return {
            "raw_transcript": "",
            "enhanced_transcript": "",
            "segments": [],
            "metadata": {},
            "audio_info": audio_info,
            "processing_time": context.processing_times,
            "success": False,
            "error": "Enhanced audio file not found"
        }
    
    # --- Skip Language Detection ---
    # Language detection is disabled - always use Whisper's built-in detection
    context.processing_times["language_detection"] = 0
    if context.options.whisper_params.language:
        logger.info(f"Using specified language: {context.options.whisper_params.language}", emoji_key="language")
    else:
        logger.info("Using Whisper's built-in language detection", emoji_key="language")
    
    # --- Transcribe Audio ---
    transcribe_start = time.time()
    model = context.options.whisper_params.model
    quality = context.options.whisper_params.quality.value
    logger.info(
        f"Transcribing audio with model '{model}' (quality: {quality})",
        emoji_key="transcribe"
    )
    
    try:
        transcript_result = await transcribe_with_whisper(context)
        context.processing_times["transcription"] = time.time() - transcribe_start
        
        raw_transcript = transcript_result.get("text", "")
        segments = transcript_result.get("segments", [])
        
        # Check if transcript is empty
        if not raw_transcript:
            logger.warning("Whisper returned an empty transcript", emoji_key="warning")
        else:
            transcript_length = len(raw_transcript)
            segments_count = len(segments)
            logger.info(
                f"Transcription complete: {transcript_length} characters, {segments_count} segments",
                emoji_key="success"
            )
        
        # Extract metadata if available
        metadata = transcript_result.get("metadata", {})
        if context.language_code and "language" not in metadata:
            metadata["language"] = context.language_code
            
    except Exception as e:
        logger.error(f"Transcription failed: {str(e)}", emoji_key="error", exc_info=True)
        context.processing_times["transcription"] = time.time() - transcribe_start
        return {
            "raw_transcript": "",
            "enhanced_transcript": "",
            "segments": [],
            "metadata": {"language": context.language_code} if context.language_code else {},
            "audio_info": audio_info,
            "processing_time": context.processing_times,
            "success": False,
            "error": f"Transcription failed: {str(e)}"
        }
    
    # --- Enhance Transcript ---
    enhanced_transcript = raw_transcript
    enhancement_cost = 0.0
    enhancement_tokens = {"input": 0, "output": 0, "total": 0}
    
    if context.options.enhance_transcript and raw_transcript:
        enhance_start = time.time()
        logger.info(
            f"Enhancing transcript with style: {context.options.enhancement_params.style.value}",
            emoji_key="enhance",
            provider=context.options.enhancement_params.provider
        )
        
        try:
            enhancement_result = await enhance_transcript(context, raw_transcript, metadata)
            
            enhanced_transcript = enhancement_result["transcript"]
            enhancement_cost = enhancement_result["cost"]
            enhancement_tokens = enhancement_result["tokens"]
            
            if "topics" in enhancement_result and enhancement_result["topics"]:
                metadata["topics"] = enhancement_result["topics"]
                
            if "title" in enhancement_result and enhancement_result["title"]:
                metadata["title"] = enhancement_result["title"]
                
            context.processing_times["transcript_enhancement"] = time.time() - enhance_start
            
            if not enhanced_transcript:
                logger.warning("Enhancement returned an empty transcript, falling back to raw transcript", emoji_key="warning")
                enhanced_transcript = raw_transcript
            else:
                enhancement_length = len(enhanced_transcript)
                logger.info(f"Enhancement complete: {enhancement_length} characters", emoji_key="success")
                
        except Exception as e:
            logger.error(f"Transcript enhancement failed: {e}", emoji_key="error", exc_info=True)
            context.processing_times["transcript_enhancement"] = time.time() - enhance_start
            # Fall back to raw transcript
            enhanced_transcript = raw_transcript
    else:
        if not raw_transcript:
            logger.warning("Skipping transcript enhancement because raw transcript is empty", emoji_key="warning")
        elif not context.options.enhance_transcript:
            logger.info("Transcript enhancement disabled by options", emoji_key="info")
            
        context.processing_times["transcript_enhancement"] = 0
    
    # --- Generate Output Files ---
    artifact_paths = await generate_output_files(context, raw_transcript, enhanced_transcript, segments, metadata)
    
    # --- Prepare Result ---
    result = {
        "raw_transcript": raw_transcript,
        "enhanced_transcript": enhanced_transcript,
        "segments": segments,
        "metadata": metadata,
        "audio_info": audio_info,
        "tokens": enhancement_tokens,
        "cost": enhancement_cost,
        "artifacts": artifact_paths,
        "success": bool(raw_transcript or enhanced_transcript)
    }
    
    return result


def parse_options(options: Dict[str, Any]) -> TranscriptionOptions:
    """Parse and validate transcription options."""
    # Convert string output formats to enum values
    if "output_formats" in options:
        if isinstance(options["output_formats"], list):
            formats = []
            for fmt in options["output_formats"]:
                if isinstance(fmt, str):
                    try:
                        formats.append(OutputFormat(fmt.lower()))
                    except ValueError:
                        logger.warning(f"Invalid output format: {fmt}, ignoring")
                elif isinstance(fmt, OutputFormat):
                    formats.append(fmt)
            if formats:  # Only update if we have valid formats
                options["output_formats"] = formats
    
    # Handle nested parameter objects
    for key in ["audio_params", "whisper_params", "enhancement_params"]:
        if key in options and options[key]:
            # If a dictionary is provided, keep it for Pydantic
            if not isinstance(options[key], dict):
                # Convert non-dict to dict by serializing/deserializing if possible
                try:
                    options[key] = json.loads(json.dumps(options[key]))
                except Exception:
                    # If can't convert, remove the invalid value
                    logger.warning(f"Invalid format for {key}, using defaults")
                    options.pop(key)
    
    # Set audio profile parameters if a profile is specified
    if "audio_params" in options and "profile" in options["audio_params"]:
        profile = options["audio_params"]["profile"]
        if isinstance(profile, str):
            try:
                profile = AudioEnhancementProfile(profile.lower())
                # Update audio parameters based on the selected profile
                options["audio_params"].update(_get_audio_profile_params(profile))
            except ValueError:
                logger.warning(f"Invalid audio profile: {profile}, using default")
    
    # Set whisper quality parameters if quality is specified
    if "whisper_params" in options and "quality" in options["whisper_params"]:
        quality = options["whisper_params"]["quality"]
        if isinstance(quality, str):
            try:
                quality = TranscriptionQuality(quality.lower())
                # Update whisper parameters based on the selected quality
                options["whisper_params"].update(_get_whisper_quality_params(quality))
            except ValueError:
                logger.warning(f"Invalid transcription quality: {quality}, using default")
    
    # Parse with Pydantic model
    try:
        return TranscriptionOptions(**options)
    except Exception as e:
        logger.warning(f"Error parsing options: {e}, using defaults with valid values", emoji_key="warning")
        # Create with default options
        return TranscriptionOptions()


def _get_audio_profile_params(profile: AudioEnhancementProfile) -> Dict[str, Any]:
    """Get audio enhancement parameters for a specific profile."""
    profiles = {
        AudioEnhancementProfile.CONFERENCE_CALL: {
            "volume": 1.5,
            "noise_reduction": 10,
            "highpass": 200,
            "lowpass": 3000,
            "compression": True,
            "normalize": True,
            "dereverberation": False
        },
        AudioEnhancementProfile.INTERVIEW: {
            "volume": 1.3,
            "noise_reduction": 8,
            "highpass": 150,
            "lowpass": 8000,
            "compression": True,
            "normalize": True,
            "dereverberation": False
        },
        AudioEnhancementProfile.LECTURE: {
            "volume": 1.4,
            "noise_reduction": 6,
            "highpass": 120,
            "lowpass": 8000,
            "compression": True,
            "normalize": True,
            "dereverberation": True
        },
        AudioEnhancementProfile.NOISY: {
            "volume": 1.8,
            "noise_reduction": 20,
            "highpass": 250,
            "lowpass": 3000,
            "compression": True,
            "normalize": True,
            "dereverberation": True
        },
        AudioEnhancementProfile.PHONE_CALL: {
            "volume": 2.0,
            "noise_reduction": 15,
            "highpass": 300,
            "lowpass": 3400,
            "compression": True,
            "normalize": True,
            "dereverberation": False
        },
        AudioEnhancementProfile.VOICEMAIL: {
            "volume": 2.0,
            "noise_reduction": 12,
            "highpass": 250,
            "lowpass": 3000,
            "compression": True,
            "normalize": True,
            "dereverberation": False
        }
    }
    
    return profiles.get(profile, {})


def _get_whisper_quality_params(quality: TranscriptionQuality) -> Dict[str, Any]:
    """Get whisper parameters for a specific quality level."""
    quality_params = {
        TranscriptionQuality.DRAFT: {
            "beam_size": 1,
            "processors": 1,
            "word_timestamps": False,
            "highlight_words": False
        },
        TranscriptionQuality.STANDARD: {
            "beam_size": 5,
            "processors": 2,
            "word_timestamps": True,
            "highlight_words": False
        },
        TranscriptionQuality.ENHANCED: {
            "beam_size": 8,
            "processors": 2,
            "word_timestamps": True,
            "highlight_words": True
        },
        TranscriptionQuality.MAXIMUM: {
            "beam_size": 10,
            "processors": 4,
            "word_timestamps": True,
            "highlight_words": True
        }
    }
    
    return quality_params.get(quality, {})


def _update_parameters_from_audio_info(context: ProcessingContext, audio_info: Dict[str, Any]) -> None:
    """Update processing parameters based on audio file analysis."""
    # If mono audio, adjust enhancement params
    audio_params = context.options.audio_params
    updated_params = False

    if audio_info.get("channels", 0) == 1:
        # Set output channels to match input if not explicitly set
        if "output_channels" not in context.options.audio_params.dict():
            # Create a copy with the updated parameter
            audio_params = AudioEnhancementParams(
                **{**audio_params.dict(), "output_channels": 1}
            )
            updated_params = True
    
    # If low-quality audio, adjust enhancement profile
    sample_rate = audio_info.get("sample_rate", 0)
    if sample_rate < 16000 and context.options.audio_params.profile == AudioEnhancementProfile.CONFERENCE_CALL:
        logger.info(f"Detected low sample rate ({sample_rate} Hz), adjusting enhancement profile", emoji_key="audio")
        # Use phone_call profile for low sample rate audio
        params = _get_audio_profile_params(AudioEnhancementProfile.PHONE_CALL)
        # Create a copy with the updated parameters
        audio_params = AudioEnhancementParams(
            **{**audio_params.dict(), **params}
        )
        updated_params = True
    
    # If params were updated, create a new options object with the updated audio_params
    if updated_params:
        context.options = TranscriptionOptions(
            **{**context.options.dict(), "audio_params": audio_params}
        )
    
    # If very short audio (<10 seconds), adjust transcription quality
    duration = audio_info.get("duration", 0)
    if duration < 10 and context.options.whisper_params.quality != TranscriptionQuality.MAXIMUM:
        logger.info(f"Short audio detected ({duration:.1f}s), increasing transcription quality", emoji_key="audio")
        # Create a new whisper_params with enhanced quality
        whisper_params = WhisperParams(
            **{**context.options.whisper_params.dict(), "quality": TranscriptionQuality.ENHANCED}
        )
        # Update options with new whisper_params
        context.options = TranscriptionOptions(
            **{**context.options.dict(), "whisper_params": whisper_params}
        )


# --- Dependency and Audio Processing Functions ---

async def download_whisper_model(model_name: str, output_path: str) -> bool:
    """Download a Whisper model from Hugging Face using httpx.
    
    Args:
        model_name: Name of the model to download (e.g. 'large-v3')
        output_path: Path where to save the model file
    
    Returns:
        True if download was successful, False otherwise
    """
    url = f"https://huggingface.co/ggerganov/whisper.cpp/resolve/main/ggml-{model_name}.bin"
    logger.info(f"Downloading Whisper model '{model_name}' from {url}", emoji_key="download")
    
    # Expected model size
    min_size_bytes = WHISPER_MODEL_SIZES.get(model_name, 100000000)  # Default to 100MB minimum if unknown
    expected_size_bytes = min_size_bytes  # Initialize with minimum size
    
    try:
        async with httpx.AsyncClient(timeout=None) as client:
            # Make HEAD request first to check if the URL is valid and get expected size
            try:
                head_response = await client.head(url)
                if head_response.status_code != 200:
                    logger.warning(f"Model URL is not accessible: HTTP {head_response.status_code} for {url}", emoji_key="warning")
                    return False
                
                # If Content-Length is available, use it to check expected size
                content_length = int(head_response.headers.get("content-length", 0))
                if content_length > 0:
                    expected_size_mb = content_length / (1024 * 1024)
                    logger.info(f"Expected model size: {expected_size_mb:.1f} MB", emoji_key="info")
                    
                    # Update expected size if it's larger than our preset minimum
                    if content_length > min_size_bytes:
                        expected_size_bytes = int(content_length * 0.95)  # Allow 5% tolerance
            except Exception as e:
                logger.warning(f"Failed to validate model URL: {url} - Error: {str(e)}", emoji_key="warning")
                # Continue anyway, the GET might still work
            
            # Stream the response to handle large files
            try:
                async with client.stream("GET", url) as response:
                    if response.status_code != 200:
                        logger.warning(f"Failed to download model: HTTP {response.status_code} for {url}", emoji_key="warning")
                        if response.status_code == 404:
                            logger.warning(f"Model '{model_name}' not found on Hugging Face repository", emoji_key="warning")
                        return False
                    
                    # Get content length if available
                    content_length = int(response.headers.get("content-length", 0))
                    if content_length == 0:
                        logger.warning("Content length is zero or not provided", emoji_key="warning")
                    
                    downloaded = 0
                    
                    # Open file for writing
                    try:
                        with open(output_path, "wb") as f:
                            # Display progress
                            last_log_time = time.time()
                            async for chunk in response.aiter_bytes(chunk_size=8192):
                                f.write(chunk)
                                downloaded += len(chunk)
                                
                                # Log progress every 5 seconds
                                now = time.time()
                                if now - last_log_time > 5:
                                    if content_length > 0:
                                        percent = downloaded / content_length * 100
                                        logger.info(f"Download progress: {percent:.1f}% ({downloaded/(1024*1024):.1f} MB)", emoji_key="download")
                                    else:
                                        logger.info(f"Downloaded {downloaded/(1024*1024):.1f} MB", emoji_key="download")
                                    last_log_time = now
                    except IOError as e:
                        logger.warning(f"Failed to write to file {output_path}: {str(e)}", emoji_key="warning")
                        return False
            except httpx.RequestError as e:
                logger.warning(f"HTTP request error while downloading model: {str(e)}", emoji_key="warning")
                return False
        
        # Verify the file was downloaded
        if not os.path.exists(output_path):
            logger.warning(f"Downloaded file doesn't exist at {output_path}", emoji_key="warning")
            return False
            
        actual_size = os.path.getsize(output_path)
        if actual_size == 0:
            logger.warning(f"Downloaded file is empty at {output_path}", emoji_key="warning")
            return False
        
        # Verify file size meets expectations
        actual_size_mb = actual_size / (1024 * 1024)
        if actual_size < expected_size_bytes:
            logger.warning(
                f"Model file size ({actual_size_mb:.1f} MB) is smaller than expected minimum size "
                f"({expected_size_bytes/(1024*1024):.1f} MB). File may be corrupted or incomplete.", 
                emoji_key="warning"
            )
            return False
            
        logger.info(f"Successfully downloaded model to {output_path} ({actual_size_mb:.1f} MB)", emoji_key="success")
        return True
            
    except Exception as e:
        logger.warning(f"Error downloading whisper model: {e}", emoji_key="warning", exc_info=True)
        return False

async def check_dependencies(context: ProcessingContext) -> bool:
    """Verifies that required dependencies are installed and accessible."""
    # Check ffmpeg
    try:
        result = await asyncio.create_subprocess_exec(
            "ffmpeg", "-version",
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE
        )
        stdout, stderr = await result.communicate()
        
        if result.returncode != 0:
            raise ResourceError(
                "ffmpeg is not installed or not accessible. Please install it with 'apt install ffmpeg'."
            )
        
        # Extract ffmpeg version for logging
        version_match = re.search(r'ffmpeg version (\S+)', stdout.decode('utf-8', errors='ignore'))
        version = version_match.group(1) if version_match else "unknown"
        logger.debug(f"Found ffmpeg version {version}", emoji_key="dependency")
        
    except FileNotFoundError as e:
        raise ResourceError(
            "ffmpeg is not installed. Please install it with 'apt install ffmpeg'."
        ) from e
    
    # Check whisper.cpp
    whisper_path = os.path.expanduser("~/whisper.cpp")
    
    # Use user-supplied model
    model = context.options.whisper_params.model
    model_path = os.path.join(whisper_path, "models", f"ggml-{model}.bin")
    
    if not os.path.exists(whisper_path):
        raise ResourceError(
            f"whisper.cpp not found at {whisper_path}. Please install it first."
        )
    
    if not os.path.exists(model_path):
        # Check if models directory exists
        models_dir = os.path.join(whisper_path, "models")
        if not os.path.exists(models_dir):
            try:
                os.makedirs(models_dir)
                logger.info(f"Created models directory at {models_dir}", emoji_key="info")
            except Exception as e:
                raise ResourceError(f"Failed to create models directory: {e}") from e
        
        # Try to automatically download the model using httpx
        logger.info(f"Whisper model '{model}' not found at {model_path}", emoji_key="info")
        logger.info(f"Attempting to download model '{model}' now...", emoji_key="download")
        
        # Download the model directly using httpx - first check if model exists
        if model == "large-v3":
            # Double check if the model actually exists
            test_url = f"https://huggingface.co/ggerganov/whisper.cpp/resolve/main/ggml-{model}.bin"
            async with httpx.AsyncClient(timeout=10.0) as client:
                try:
                    head_response = await client.head(test_url)
                    if head_response.status_code != 200:
                        # Model might not be directly available - show a clear error
                        if head_response.status_code == 404:
                            raise ResourceError(
                                f"Whisper model 'large-v3' not found in the HuggingFace repository at {test_url}\n"
                                f"Please download it manually with one of these commands:\n"
                                f"1. ~/whisper.cpp/models/download-ggml-model.sh large-v3\n"
                                f"2. Or try a different model like 'large-v3-turbo' which is known to be available"
                            )
                except Exception as e:
                    logger.warning(f"Error checking model availability: {e}", emoji_key="warning")
                    # Continue with download attempt anyway
        
        # Attempt to download the model
        download_success = await download_whisper_model(model, model_path)
        if not download_success:
            # Modified error message with alternative suggestions
            if model == "large-v3":
                raise ResourceError(
                    f"Failed to download Whisper model '{model}'.\n"
                    f"You can:\n"
                    f"1. Try using a different model like 'large-v3-turbo' which is more reliable\n"
                    f"2. Or download manually with: ~/whisper.cpp/models/download-ggml-model.sh {model}"
                )
            else:
                raise ResourceError(
                    f"Failed to download Whisper model '{model}'.\n"
                    f"Please download it manually with: ~/whisper.cpp/models/download-ggml-model.sh {model}"
                )
        
        # Verify that the model was downloaded
        if not os.path.exists(model_path):
            raise ResourceError(
                f"Model download completed but model file not found at {model_path}. "
                f"Please check the download and try again."
            )
        
        # Verify model file size
        actual_size = os.path.getsize(model_path)
        expected_min_size = WHISPER_MODEL_SIZES.get(model, 100000000)  # Default to 100MB minimum if unknown
        
        if actual_size < expected_min_size:
            actual_size_mb = actual_size / (1024 * 1024)
            expected_mb = expected_min_size / (1024 * 1024)
            raise ResourceError(
                f"Downloaded model file at {model_path} is too small ({actual_size_mb:.1f} MB). "
                f"Expected at least {expected_mb:.1f} MB. File may be corrupted or incomplete. "
                f"Please download it manually with: ~/whisper.cpp/models/download-ggml-model.sh {model}"
            )
        
        logger.info(f"Successfully downloaded Whisper model '{model}'", emoji_key="success")
    else:
        # Verify existing model file size
        actual_size = os.path.getsize(model_path)
        expected_min_size = WHISPER_MODEL_SIZES.get(model, 100000000)  # Default to 100MB minimum if unknown
        
        file_size_mb = actual_size / (1024 * 1024)
        if actual_size < expected_min_size:
            expected_mb = expected_min_size / (1024 * 1024)
            logger.warning(
                f"Existing model at {model_path} is suspiciously small ({file_size_mb:.1f} MB). "
                f"Expected at least {expected_mb:.1f} MB. Model may be corrupted.", 
                emoji_key="warning"
            )
        else:
            logger.info(f"Found existing model at {model_path} ({file_size_mb:.1f} MB)", emoji_key="dependency")
    
    # Check if whisper binary is available in PATH using shlex for command safety
    try:
        result = subprocess.run(["which", "whisper-cli"], capture_output=True, text=True)
        if result.returncode == 0:
            whisper_path_found = result.stdout.strip()
            logger.debug(f"Found whisper-cli in PATH: {whisper_path_found}", emoji_key="dependency")
        else:
            # Check in the expected location
            whisper_path = os.path.expanduser("~/whisper.cpp")
            whisper_bin = os.path.join(whisper_path, "build", "bin", "whisper-cli")
            if not os.path.exists(whisper_bin):
                raise ResourceError(
                    f"whisper-cli binary not found at {whisper_bin}. "
                    f"Please build whisper.cpp first with: "
                    f"cd ~/whisper.cpp && cmake -B build && cmake --build build -j --config Release"
                )
            logger.debug(f"Found whisper-cli at {whisper_bin}", emoji_key="dependency")
    except FileNotFoundError as e:
        raise ResourceError("Command 'which' not found. Cannot check for whisper-cli installation.") from e
    
    logger.debug(f"Found whisper model: {model}", emoji_key="dependency")
    return True


async def get_detailed_audio_info(file_path: str) -> Dict[str, Any]:
    """Gets detailed information about an audio file using ffprobe."""
    cmd = [
        "ffprobe",
        "-v", "error",
        "-show_entries", "format=duration,bit_rate,size",
        "-select_streams", "a:0",
        "-show_entries", "stream=channels,sample_rate,codec_name,bits_per_sample",
        "-of", "json",
        file_path
    ]
    
    try:
        process = await asyncio.create_subprocess_exec(
            *cmd,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE
        )
        stdout, stderr = await process.communicate()
        
        if process.returncode != 0:
            logger.warning(f"Failed to get audio info: {stderr.decode('utf-8', errors='ignore')}", emoji_key="warning")
            return {
                "duration": 0,
                "channels": 0,
                "sample_rate": 0,
                "format": os.path.splitext(file_path)[1][1:],
                "codec": None,
                "bit_depth": None,
                "bitrate": None,
                "size_bytes": os.path.getsize(file_path) if os.path.exists(file_path) else 0
            }
        
        info = json.loads(stdout)
        format_info = info.get("format", {})
        stream_info = info.get("streams", [{}])[0] if info.get("streams") else {}
        
        # Extract information
        duration = float(format_info.get("duration", 0))
        channels = int(stream_info.get("channels", 0))
        sample_rate = int(stream_info.get("sample_rate", 0))
        codec = stream_info.get("codec_name")
        bit_depth = int(stream_info.get("bits_per_sample", 0)) or None
        bitrate = int(format_info.get("bit_rate", 0)) or None
        size_bytes = int(format_info.get("size", 0)) or os.path.getsize(file_path)
        audio_format = os.path.splitext(file_path)[1][1:]
        
        return {
            "duration": duration,
            "channels": channels,
            "sample_rate": sample_rate,
            "format": audio_format,
            "codec": codec,
            "bit_depth": bit_depth,
            "bitrate": bitrate,
            "size_bytes": size_bytes
        }
    except Exception as e:
        logger.warning(f"Error getting audio info: {e}", emoji_key="warning", exc_info=True)
        try:
            size_bytes = os.path.getsize(file_path) if os.path.exists(file_path) else 0
        except Exception:
            size_bytes = 0
            
        return {
            "duration": 0,
            "channels": 0,
            "sample_rate": 0,
            "format": os.path.splitext(file_path)[1][1:],
            "codec": None,
            "bit_depth": None,
            "bitrate": None,
            "size_bytes": size_bytes
        }


async def enhance_audio(context: ProcessingContext, audio_info: Dict[str, Any]) -> Optional[str]:
    """Enhances audio quality using ffmpeg preprocessing."""
    # Create output path in temp directory
    output_path = os.path.join(context.temp_dir, f"{context.base_filename}_enhanced.wav")
    
    # Use optimized audio enhancement settings
    # Build the complete command with the standardized enhancement settings
    cmd = [
        "ffmpeg",
        "-i", context.file_path,
        "-threads", str(os.cpu_count() or 1),
        "-af", "volume=1.5, highpass=f=200, lowpass=f=3000, afftdn=nr=10:nf=-20, "
              "compand=attacks=0:points=-80/-80|-45/-15|-27/-9|0/-7|20/-7:gain=5, "
              "dynaudnorm=f=150:g=15:p=1:m=1:s=0, "
              "pan=stereo|c0=c0|c1=c0",
        "-ar", "16000",
        "-ac", "2",
        "-c:a", "pcm_s16le",
        "-y",  # Overwrite output if exists
        output_path
    ]
    
    logger.debug(f"Running ffmpeg command: {' '.join(cmd)}", emoji_key="command")
    
    try:
        process = await asyncio.create_subprocess_exec(
            *cmd,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE
        )
        stdout, stderr = await process.communicate()
        
        if process.returncode != 0:
            error_msg = stderr.decode('utf-8', errors='ignore')
            logger.error(f"FFmpeg error: {error_msg}", emoji_key="error")
            return None
        
        # Verify the output file was created and has a reasonable size
        if os.path.exists(output_path) and os.path.getsize(output_path) > 1000:
            file_size_mb = os.path.getsize(output_path) / (1024 * 1024)
            logger.info(f"Enhanced audio saved to {output_path} ({file_size_mb:.2f} MB)", emoji_key="audio")
        else:
            logger.warning("Enhanced audio file is suspiciously small or doesn't exist", emoji_key="warning")
            return None
        
        # If we're keeping enhanced audio, copy it to a persistent location
        if context.options.save_enhanced_audio:
            original_dir = os.path.dirname(context.file_path)
            persistent_path = os.path.join(original_dir, f"{context.base_filename}_enhanced.wav")
            shutil.copy2(output_path, persistent_path)
            logger.info(f"Saved enhanced audio to {persistent_path}", emoji_key="save")
        
        logger.info("Audio enhancement completed", emoji_key="audio")
        return output_path
    except Exception as e:
        logger.error(f"Error enhancing audio: {e}", emoji_key="error", exc_info=True)
        return None


async def transcribe_with_whisper(context: ProcessingContext) -> Dict[str, Any]:
    """Transcribes audio using Whisper.cpp with advanced options."""
    # Create output base name in temp directory
    output_base = os.path.join(context.temp_dir, context.base_filename)
    output_json = f"{output_base}.json"
    output_txt = f"{output_base}.txt"
    
    # Get whisper parameters
    params = context.options.whisper_params
    
    # Build command with configurable parameters
    whisper_bin = os.path.expanduser("~/whisper.cpp/build/bin/whisper-cli")
    model_path = os.path.expanduser(f"~/whisper.cpp/models/ggml-{params.model}.bin")
    
    # Validate required files exist
    if not os.path.exists(whisper_bin):
        logger.error(f"Whisper binary not found at {whisper_bin}", emoji_key="error")
        raise ToolError(f"Whisper binary not found at {whisper_bin}")
    
    if not os.path.exists(model_path):
        logger.error(f"Whisper model not found at {model_path}", emoji_key="error")
        raise ToolError(f"Whisper model not found at {model_path}")
    
    # Verify model file size
    actual_size = os.path.getsize(model_path)
    expected_min_size = WHISPER_MODEL_SIZES.get(params.model, 100000000)  # Default to 100MB minimum if unknown
    if actual_size < expected_min_size:
        actual_size_mb = actual_size / (1024 * 1024)
        expected_mb = expected_min_size / (1024 * 1024)
        logger.warning(
            f"Model file at {model_path} is smaller than expected ({actual_size_mb:.1f} MB, expected {expected_mb:.1f} MB)",
            emoji_key="warning"
        )
    
    # Use file_path as fallback if enhanced_audio_path is None
    audio_path = context.enhanced_audio_path or context.file_path
    if not os.path.exists(audio_path):
        logger.error(f"Audio file not found at {audio_path}", emoji_key="error")
        raise ToolError(f"Audio file not found at {audio_path}")
    
    cmd = [
        whisper_bin,
        "-m", model_path,
        "-f", audio_path,
        "-of", output_base,
        "-oj"  # Always output JSON for post-processing
    ]
    
    # Add boolean flags
    if params.word_timestamps:
        cmd.append("-pc")
    
    if params.translate:
        cmd.append("-tr")
    
    # Always output text for readability
    cmd.append("-otxt")
    
    # Add numeric parameters
    cmd.extend(["-t", str(os.cpu_count() if params.processors <= 0 else params.processors)])
    
    if params.beam_size:
        cmd.extend(["-bs", str(params.beam_size)])
    
    # Add language parameter if specified
    if params.language:
        cmd.extend(["-l", params.language])
    
    # Add max context parameter if specified
    if params.max_context > 0:
        cmd.extend(["-mc", str(params.max_context)])
    
    # Additional optimizations
    cmd.append("-fa")  # Full sentence timestamps (improved segmentation)
    cmd.append("-pp")  # Enable post-processing
    
    # Add custom vocab if specified (create a vocab file)
    if params.custom_vocab:
        vocab_path = os.path.join(context.temp_dir, "custom_vocab.txt")
        try:
            async with aiofiles.open(vocab_path, 'w') as f:
                await f.write("\n".join(params.custom_vocab))
            cmd.extend(["-kv", vocab_path])
        except Exception as e:
            logger.warning(f"Failed to create custom vocab file: {e}", emoji_key="warning")
    
    # Add diarization if requested
    if params.diarize:
        cmd.append("-dm")
    
    cmd_str = ' '.join(cmd)
    logger.debug(f"Running whisper command: {cmd_str}", emoji_key="command")
    
    try:
        process = await asyncio.create_subprocess_exec(
            *cmd,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE
        )
        stdout, stderr = await process.communicate()
        
        stderr_output = stderr.decode('utf-8', errors='ignore') if stderr else ""
        stdout_output = stdout.decode('utf-8', errors='ignore') if stdout else ""
        
        # Log outputs for debugging
        if stdout_output:
            logger.debug(f"Whisper stdout (first 500 chars): {stdout_output[:500]}", emoji_key="info")
        
        if stderr_output:
            if process.returncode != 0:
                logger.error(f"Whisper stderr: {stderr_output}", emoji_key="error")
            elif "error" in stderr_output.lower() or "warning" in stderr_output.lower():
                logger.warning(f"Whisper warnings/errors: {stderr_output}", emoji_key="warning")
            else:
                logger.debug(f"Whisper stderr: {stderr_output}", emoji_key="info")
        
        if process.returncode != 0:
            error_msg = stderr_output or "Unknown error"
            logger.error(f"Whisper transcription error (exit code {process.returncode}): {error_msg}", emoji_key="error")
            raise ToolError(f"Whisper transcription failed with exit code {process.returncode}: {error_msg}")
            
        # Check if output files exist
        if not os.path.exists(output_json) and not os.path.exists(output_txt):
            logger.error("Whisper completed successfully but no output files were created", emoji_key="error")
            raise ToolError("Whisper completed successfully but no output files were created")
            
        # Read results from the JSON file
        if os.path.exists(output_json):
            json_file_size = os.path.getsize(output_json)
            logger.debug(f"Reading JSON output file: {output_json} ({json_file_size} bytes)", emoji_key="info")
            
            if json_file_size < 10:  # Suspiciously small
                logger.warning(f"JSON output file is suspiciously small: {json_file_size} bytes", emoji_key="warning")
            
            try:
                async with aiofiles.open(output_json, 'r') as f:
                    content = await f.read()
                    
                try:
                    result = json.loads(content)
                    
                    # Fix missing fields in result if needed
                    if "segments" not in result:
                        logger.warning("No segments found in Whisper JSON output", emoji_key="warning")
                        result["segments"] = []
                    
                    if "text" not in result or not result.get("text"):
                        logger.warning("No transcript text found in Whisper JSON output", emoji_key="warning")
                        # Try to construct text from segments
                        if result.get("segments"):
                            reconstructed_text = " ".join([seg.get("text", "") for seg in result["segments"]])
                            if reconstructed_text:
                                logger.info("Reconstructed transcript text from segments", emoji_key="info")
                                result["text"] = reconstructed_text
                            else:
                                result["text"] = ""
                        else:
                            result["text"] = ""
                    
                    # Extract metadata
                    metadata = {
                        "language": context.language_code or result.get("language"),
                        "duration": result.get("duration", 0)
                    }
                    result["metadata"] = metadata
                    
                except json.JSONDecodeError as e:
                    logger.error(f"Failed to parse Whisper JSON output: {e}", emoji_key="error")
                    logger.error(f"JSON content: {content[:1000]}...", emoji_key="error")
                    raise ToolError(f"Failed to parse Whisper output JSON: {e}") from e
            except Exception as e:
                logger.error(f"Failed to read Whisper JSON output file: {e}", emoji_key="error")
                raise ToolError(f"Failed to read Whisper output: {e}") from e
        else:
            logger.warning(f"Whisper JSON output not found at expected path: {output_json}", emoji_key="warning")
            # Fallback to text file
            if os.path.exists(output_txt):
                txt_file_size = os.path.getsize(output_txt)
                logger.info(f"Falling back to text output file: {output_txt} ({txt_file_size} bytes)", emoji_key="info")
                
                if txt_file_size < 10:  # Suspiciously small
                    logger.warning(f"Text output file is suspiciously small: {txt_file_size} bytes", emoji_key="warning")
                
                try:
                    async with aiofiles.open(output_txt, 'r') as f:
                        text = await f.read()
                    
                    # Create minimal result structure
                    result = {
                        "text": text,
                        "segments": [{"text": text, "start": 0, "end": 0}],
                        "metadata": {
                            "language": context.language_code,
                            "duration": 0
                        }
                    }
                    
                    if not text:
                        logger.warning("Text output file is empty", emoji_key="warning")
                except Exception as e:
                    logger.error(f"Failed to read text output file: {e}", emoji_key="error")
                    raise ToolError(f"Failed to read Whisper text output: {e}") from e
            else:
                logger.error(f"No output files found from Whisper at {output_base}.*", emoji_key="error")
                raise ToolError("No output files found from Whisper transcription")
        
        # Check if we actually got a transcript
        if not result.get("text"):
            logger.warning("Whisper returned an empty transcript", emoji_key="warning")
        
        # Clean up results (remove empty/duplicate segments)
        cleaned_segments = clean_segments(result.get("segments", []))
        result["segments"] = cleaned_segments
        
        # Clean up text (remove dots-only lines, etc.)
        result["text"] = clean_raw_transcript(result.get("text", ""))
        
        return result
    
    except ToolError:
        raise
    except Exception as e:
        logger.error(f"Error in Whisper transcription: {e}", emoji_key="error", exc_info=True)
        raise ToolError(f"Whisper transcription failed: {str(e)}") from e


def clean_segments(segments: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """Clean and normalize segment data."""
    cleaned_segments = []
    seen_texts = set()
    
    for segment in segments:
        # Skip segments with empty or meaningless text
        text = segment.get("text", "").strip()
        if not text or re.match(r'^[\s.]*$', text):
            continue
        
        # Skip exact duplicates unless they have meaningful timing differences
        # (within 0.5s of a previous segment)
        is_duplicate = False
        if text in seen_texts:
            for prev_segment in cleaned_segments:
                if prev_segment.get("text") == text:
                    start_diff = abs(prev_segment.get("start", 0) - segment.get("start", 0))
                    end_diff = abs(prev_segment.get("end", 0) - segment.get("end", 0))
                    if start_diff < 0.5 and end_diff < 0.5:
                        is_duplicate = True
                        break
        
        if is_duplicate:
            continue
        
        # Add to seen texts
        seen_texts.add(text)
        
        # Standardize segment structure
        clean_segment = {
            "start": float(segment.get("start", 0)),
            "end": float(segment.get("end", 0)),
            "text": text
        }
        
        # Add optional fields if available
        for field in ["speaker", "words", "confidence"]:
            if field in segment:
                clean_segment[field] = segment[field]
        
        cleaned_segments.append(clean_segment)
    
    # Sort by start time
    cleaned_segments.sort(key=lambda x: x["start"])
    
    return cleaned_segments


def clean_raw_transcript(text: str) -> str:
    """Cleans raw transcript."""
    if not text:
        return ""
    
    # Split into lines and process
    lines = text.split("\n")
    cleaned_lines = []
    seen_lines = set()
    
    for line in lines:
        line = line.strip()
        # Skip empty lines
        if not line:
            continue
        # Skip lines with just dots or other meaningless patterns
        if re.match(r'^[\s.]*$', line) or line == '[BLANK_AUDIO]':
            continue
        
        # Standardize multiple spaces
        line = re.sub(r'\s+', ' ', line)
        
        # Keep duplicates if they're long (likely not duplicates but legitimate repetition)
        if line in seen_lines and len(line) <= 50:
            continue
        
        seen_lines.add(line)
        cleaned_lines.append(line)
    
    # Join but ensure there's proper spacing
    return "\n".join(cleaned_lines)


# --- Transcript Enhancement Functions ---

async def enhance_transcript(
    context: ProcessingContext, 
    transcript: str,
    metadata: Dict[str, Any]
) -> Dict[str, Any]:
    """Enhance a transcript with formatting, readability and semantic structuring."""
    if not transcript or transcript.strip() == "":
        return {
            "transcript": "",
            "tokens": {"input": 0, "output": 0, "total": 0},
            "cost": 0.0,
            "topics": [],
            "title": None
        }
    
    # Extract key parameters
    params = context.options.enhancement_params
    
    # Split transcript into manageable chunks
    chunks = await chunk_text(transcript, params.max_chunk_size)
    
    if not chunks:
        return {
            "transcript": transcript,
            "tokens": {"input": 0, "output": 0, "total": 0},
            "cost": 0.0,
            "topics": [],
            "title": None
        }
    
    # First analyze context to get a summary of the content
    context_data = await detect_subject_matter(
        chunks[0], 
        params.provider, 
        params.model,
        metadata
    )
    
    # Track topics if available
    topics = context_data.get("topics", [])
    title = context_data.get("title")
    context_info = context_data.get("context", "")
    
    logger.info(f"Content analysis complete: {len(topics)} topics identified", emoji_key="analyze")
    
    # Process chunks concurrently if parallel processing is enabled
    if context.options.parallel_processing and len(chunks) > 1:
        enhanced_chunks = await process_chunks_parallel(
            context, 
            chunks, 
            context_info, 
            params
        )
    else:
        enhanced_chunks = await process_chunks_sequential(
            context, 
            chunks, 
            context_info, 
            params
        )
    
    # Calculate total metrics
    total_tokens = {"input": 0, "output": 0, "total": 0}
    total_cost = 0.0
    
    for chunk_data in enhanced_chunks:
        chunk_tokens = chunk_data.get("tokens", {})
        total_tokens["input"] += chunk_tokens.get("input", 0)
        total_tokens["output"] += chunk_tokens.get("output", 0)
        total_tokens["total"] += chunk_tokens.get("total", 0)
        total_cost += chunk_data.get("cost", 0.0)
    
    # Join the enhanced chunks
    enhanced_transcript = "\n\n".join(chunk_data["text"] for chunk_data in enhanced_chunks)
    
    # If sections are enabled, try to add section headings
    if params.sections and topics:
        enhanced_transcript = await add_section_headings(
            enhanced_transcript, 
            topics, 
            params.provider, 
            params.model
        )
    
    return {
        "transcript": enhanced_transcript,
        "tokens": total_tokens,
        "cost": total_cost,
        "topics": topics,
        "title": title
    }


async def chunk_text(text: str, max_chunk_size: int = 6500) -> List[str]:
    """Split text into chunks with intelligent boundary detection."""
    if len(text) <= max_chunk_size:
        return [text]
    
    # Define patterns for natural breaks, prioritized
    patterns = [
        r'\n\s*\n\s*\n',  # Triple line break (highest priority)
        r'\n\s*\n',       # Double line break
        r'(?<=[.!?])\s+(?=[A-Z])',  # Sentence boundary with capital letter following
        r'(?<=[.!?])\s',  # Any sentence boundary
        r'(?<=[,:;])\s'   # Phrase boundary (lowest priority)
    ]
    
    chunks = []
    remaining_text = text
    
    while remaining_text:
        if len(remaining_text) <= max_chunk_size:
            chunks.append(remaining_text)
            break
        
        # Start with an initial chunk at max size
        chunk_candidate = remaining_text[:max_chunk_size]
        split_position = None
        
        # Try each pattern in order of priority
        for pattern in patterns:
            # Look for the last occurrence of the pattern
            matches = list(re.finditer(pattern, chunk_candidate))
            if matches:
                # Use the last match as the split point
                split_position = matches[-1].end()
                break
        
        # Fallback if no natural breaks found
        if split_position is None or split_position < max_chunk_size // 2:
            # Look for the last space after a minimum chunk size
            min_size = max(max_chunk_size // 2, 1000)
            last_space = chunk_candidate.rfind(' ', min_size)
            if last_space > min_size:
                split_position = last_space
            else:
                # Forced split at max_chunk_size
                split_position = max_chunk_size
        
        # Create chunk and update remaining text
        chunks.append(remaining_text[:split_position].strip())
        remaining_text = remaining_text[split_position:].strip()
    
    # Validate chunks
    validated_chunks = []
    for chunk in chunks:
        if chunk and len(chunk) >= 100:  # Minimum viable chunk size
            validated_chunks.append(chunk)
        elif chunk:
            # If chunk is too small, combine with previous or next chunk
            if validated_chunks:
                validated_chunks[-1] += "\n\n" + chunk
            else:
                # This is the first chunk and it's too small - rare case
                validated_chunks.append(chunk)
    
    return validated_chunks


async def detect_subject_matter(
    text: str, 
    provider: str, 
    model: Optional[str] = None,
    metadata: Optional[Dict[str, Any]] = None
) -> Dict[str, Any]:
    """Analyze transcript to determine subject matter, topics, and title."""
    prompt = """Analyze this transcript excerpt for the following:

1. CONTEXT: The primary domain or topic being discussed (e.g., technology, business, healthcare, etc.)
2. SPEAKERS: The likely type and number of speakers (e.g., interview, panel, lecture, etc.)
3. TOPICS: List 2-5 specific topics covered, in order of importance
4. TITLE: A short, descriptive title for this content (under 10 words)

Return your analysis in JSON format ONLY:
{
  "context": "Brief description of the domain and conversation type",
  "topics": ["Topic 1", "Topic 2", "Topic 3"],
  "title": "Concise descriptive title"
}

Transcript excerpt:
{text}"""

    # Include metadata if available
    if metadata and metadata.get("language"):
        language = metadata.get("language")
        prompt += f"\n\nMetadata: The transcript language is {language}."
    
    try:
        result = await chat_completion(
            messages=[{"role": "user", "content": prompt.format(text=text)}],
            provider=provider,
            model=model,
            temperature=0
        )
        
        if result.get("success") and "message" in result:
            content = result["message"].get("content", "")
            
            # Try to parse JSON from content
            try:
                # Extract JSON if it's embedded in text
                json_match = re.search(r'({[\s\S]*})', content)
                if json_match:
                    analysis = json.loads(json_match.group(1))
                else:
                    analysis = json.loads(content)
                
                return {
                    "context": analysis.get("context", ""),
                    "topics": analysis.get("topics", []),
                    "title": analysis.get("title")
                }
            except json.JSONDecodeError:
                # Fallback: extract fields manually
                context_match = re.search(r'context["\s:]+([^"}\n]+)', content, re.IGNORECASE)
                title_match = re.search(r'title["\s:]+([^"}\n]+)', content, re.IGNORECASE)
                topics_match = re.search(r'topics["\s:]+\[(.*?)\]', content, re.IGNORECASE | re.DOTALL)
                
                context = context_match.group(1).strip() if context_match else ""
                title = title_match.group(1).strip() if title_match else None
                
                topics = []
                if topics_match:
                    topics_text = topics_match.group(1)
                    topics = [t.strip().strip('"\'') for t in re.findall(r'"([^"]+)"', topics_text)]
                    if not topics:
                        topics = [t.strip() for t in topics_text.split(',') if t.strip()]
                
                return {
                    "context": context,
                    "topics": topics,
                    "title": title
                }
    except Exception as e:
        logger.warning(f"Subject matter detection failed: {e}", emoji_key="warning")
    
    return {
        "context": "",
        "topics": [],
        "title": None
    }


async def process_chunks_parallel(
    context: ProcessingContext,
    chunks: List[str],
    context_info: str,
    params: TranscriptEnhancementParams
) -> List[Dict[str, Any]]:
    """Process transcript chunks in parallel for better performance."""
    # Limit max workers to CPU count or specified max
    max_workers = min(context.options.max_workers, os.cpu_count() or 4, len(chunks))
    
    # Create a semaphore to limit concurrency
    sem = asyncio.Semaphore(max_workers)
    
    # Create a thread pool for parallel processing
    chunk_results = []
    
    async def process_chunk(i, chunk):
        """Process an individual chunk."""
        async with sem:  # Use semaphore to limit concurrency
            logger.info(f"Enhancing chunk {i+1}/{len(chunks)}", emoji_key="enhance")
            try:
                result = await enhance_chunk(chunk, context_info, params, i, len(chunks))
                return result
            except Exception as e:
                logger.error(f"Error enhancing chunk {i+1}: {e}", emoji_key="error", exc_info=True)
                return {"text": chunk, "tokens": {"input": 0, "output": 0, "total": 0}, "cost": 0.0}
    
    # Create tasks for parallel execution
    tasks = [process_chunk(i, chunk) for i, chunk in enumerate(chunks)]
    chunk_results = await asyncio.gather(*tasks)
    
    # Make sure results are in original order (they should be)
    return chunk_results


async def process_chunks_sequential(
    context: ProcessingContext,
    chunks: List[str],
    context_info: str,
    params: TranscriptEnhancementParams
) -> List[Dict[str, Any]]:
    """Process transcript chunks sequentially to preserve context flow."""
    enhanced_chunks = []
    accumulated_context = context_info
    
    for i, chunk in enumerate(chunks):
        logger.info(f"Enhancing chunk {i+1}/{len(chunks)}", emoji_key="enhance")
        
        # Update context with information from previous chunks
        if i > 0 and enhanced_chunks:
            # Add brief summary of what was covered in previous chunk
            previous_text = enhanced_chunks[-1]["text"]
            if len(previous_text) > 500:
                accumulated_context += f"\nPrevious chunk ended with: {previous_text[-500:]}"
            else:
                accumulated_context += f"\nPrevious chunk: {previous_text}"
        
        try:
            result = await enhance_chunk(chunk, accumulated_context, params, i, len(chunks))
            enhanced_chunks.append(result)
        except Exception as e:
            logger.error(f"Error enhancing chunk {i+1}: {e}", emoji_key="error", exc_info=True)
            # Use original text on error
            enhanced_chunks.append({"text": chunk, "tokens": {"input": 0, "output": 0, "total": 0}, "cost": 0.0})
    
    return enhanced_chunks


async def enhance_chunk(
    chunk: str, 
    context_info: str, 
    params: TranscriptEnhancementParams,
    chunk_index: int,
    total_chunks: int
) -> Dict[str, Any]:
    """Enhance a single transcript chunk with LLM."""
    # Build the prompt based on enhancement parameters
    style_instructions = _get_style_instructions(params.style)
    fix_instructions = []
    
    if params.add_paragraphs:
        fix_instructions.append("- Add paragraph breaks at natural topic transitions")
    
    if params.fix_spelling:
        fix_instructions.append("- Fix obvious spelling errors while preserving domain-specific terms")
    
    if params.fix_grammar:
        fix_instructions.append("- Fix basic grammatical errors without changing style or meaning")
    
    if params.format_numbers:
        fix_instructions.append("- Format numbers consistently (e.g., '25' instead of 'twenty-five')")
    
    if params.identify_speakers:
        fix_instructions.append("- Try to identify different speakers and label them as Speaker 1, Speaker 2, etc.")
        fix_instructions.append("- Format speaker changes as 'Speaker N: text' on a new line")
    
    fix_section = "\n".join(fix_instructions) if fix_instructions else "None"
    
    # Add custom instructions if provided
    custom_section = f"\nADDITIONAL INSTRUCTIONS:\n{params.custom_instructions}" if params.custom_instructions else ""
    
    # Mention chunk position in context
    position_info = f"This is chunk {chunk_index+1} of {total_chunks}." if total_chunks > 1 else ""
    
    prompt = f"""You are cleaning up a raw transcript from a recorded conversation. {position_info}

CONTENT CONTEXT: {context_info}

ENHANCEMENT STYLE: {style_instructions}

CLEANUP INSTRUCTIONS:
1. Remove filler sounds: "um", "uh", "er", "ah", "hmm"
2. Remove stutters and word repetitions: "the- the", "I- I"
3. Remove meaningless filler phrases when used as pure filler: "you know", "like", "sort of"
4. Fix clear transcription errors and garbled text
5. Add proper punctuation for readability
{fix_section}

STRICT PRESERVATION RULES:
1. DO NOT modify, rephrase, or restructure ANY of the speaker's content
2. DO NOT add ANY new content or explanations
3. DO NOT make the language more formal or technical
4. DO NOT summarize or condense anything
5. PRESERVE ALL technical terms, numbers, and specific details exactly as spoken
6. PRESERVE the speaker's unique speaking style and personality
{custom_section}

Here's the transcript chunk to clean:
{chunk}

Return ONLY the cleaned transcript text with no explanations, comments, or metadata."""

    try:
        result = await chat_completion(
            messages=[{"role": "user", "content": prompt}],
            provider=params.provider,
            model=params.model,
            temperature=0.1,  # Low temperature for consistent results
            max_tokens=min(len(chunk) * 2, 8192)  # Reasonable token limit based on input size
        )
        
        if result.get("success") and "message" in result:
            enhanced_text = result["message"].get("content", "").strip()
            
            # Validation: if enhanced text is much shorter, it might have been summarized
            if len(enhanced_text) < len(chunk) * 0.6:
                logger.warning(
                    f"Enhanced text suspiciously short ({len(enhanced_text)} vs {len(chunk)} chars), "
                    f"may have been summarized. Using original.",
                    emoji_key="warning"
                )
                enhanced_text = chunk
            
            return {
                "text": enhanced_text,
                "tokens": result.get("tokens", {"input": 0, "output": 0, "total": 0}),
                "cost": result.get("cost", 0.0)
            }
        
        return {"text": chunk, "tokens": {"input": 0, "output": 0, "total": 0}, "cost": 0.0}
    except Exception as e:
        logger.error(f"Chunk enhancement error: {e}", emoji_key="error")
        return {"text": chunk, "tokens": {"input": 0, "output": 0, "total": 0}, "cost": 0.0}


def _get_style_instructions(style: EnhancementStyle) -> str:
    """Get instructions for the specified enhancement style."""
    styles = {
        EnhancementStyle.RAW: (
            "Minimal cleaning only. Preserve all speech patterns and informality. "
            "Focus on removing only transcription errors and unintelligible elements."
        ),
        EnhancementStyle.READABLE: (
            "Basic readability improvements. Light cleanup while preserving natural speech patterns. "
            "Remove only clear disfluencies and maintain conversational flow."
        ),
        EnhancementStyle.POLISHED: (
            "Well-formatted with proper punctuation and clean sentences. "
            "Remove speech disfluencies but preserve the speaker's voice and style. "
            "Create a professional but authentic reading experience."
        ),
        EnhancementStyle.VERBATIM: (
            "Preserve all speech patterns, hesitations, and repetitions. "
            "Format for readability but maintain every verbal quirk and pause. "
            "Indicate hesitations with ellipses [...] and preserve every repeated word or phrase."
        ),
        EnhancementStyle.STRUCTURED: (
            "Add semantic structure with clear paragraphs around topics. "
            "Clean speech for maximum readability while preserving content accuracy. "
            "Organize into logical sections while keeping all original information."
        )
    }
    
    return styles.get(style, styles[EnhancementStyle.READABLE])


async def add_section_headings(
    transcript: str, 
    topics: List[str], 
    provider: str, 
    model: Optional[str] = None
) -> str:
    """Add section headings to the transcript based on topic changes."""
    if not transcript or not topics:
        return transcript
    
    prompt = """Add clear section headings to this transcript based on topic changes. 

TOPICS COVERED (in approximate order):
{topics}

RULES:
1. Insert section headings as "## [Topic]" on their own line
2. Place headings ONLY where there is a clear topic change
3. Use at most {max_sections} headings total
4. NEVER add content or edit the existing text
5. NEVER remove any original content
6. Base headings on the given topics list, but you can adjust wording for clarity
7. Don't duplicate headings for the same topic
8. Keep headings short and descriptive (2-6 words each)

TRANSCRIPT:
{text}

Return the full transcript with section headings added."""

    # Adjust max sections based on transcript length
    token_estimate = len(transcript) // 4
    max_sections = min(len(topics) + 1, token_estimate // 1000 + 1)
    topics_text = "\n".join([f"- {topic}" for topic in topics])
    
    try:
        result = await chat_completion(
            messages=[{
                "role": "user", 
                "content": prompt.format(topics=topics_text, text=transcript, max_sections=max_sections)
            }],
            provider=provider,
            model=model,
            temperature=0.1,
            max_tokens=min(len(transcript) * 2, 8192)
        )
        
        if result.get("success") and "message" in result:
            return result["message"].get("content", "").strip()
        
        return transcript
    except Exception as e:
        logger.warning(f"Failed to add section headings: {e}", emoji_key="warning")
        return transcript


# --- Output File Generation ---

async def generate_output_files(
    context: ProcessingContext,
    raw_transcript: str,
    enhanced_transcript: str,
    segments: List[Dict[str, Any]],
    metadata: Dict[str, Any]
) -> Dict[str, Any]:
    """Generate output files in requested formats."""
    artifact_paths = {
        "output_files": {}
    }
    
    # Save enhanced audio path if requested
    if context.options.save_enhanced_audio and context.enhanced_audio_path:
        original_dir = os.path.dirname(context.file_path)
        persistent_path = os.path.join(original_dir, f"{context.base_filename}_enhanced.wav")
        
        # May have already been saved during enhancement
        if not os.path.exists(persistent_path) and os.path.exists(context.enhanced_audio_path):
            shutil.copy2(context.enhanced_audio_path, persistent_path)
            
        artifact_paths["enhanced_audio"] = persistent_path
    
    # Generate requested output formats
    output_formats = context.options.output_formats
    output_dir = os.path.dirname(context.file_path)
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    
    # Convert temp files to Path objects for easier path manipulation
    output_dir = Path(os.path.dirname(context.file_path))
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    
    # Generate JSON output
    if OutputFormat.JSON in output_formats:
        json_path = output_dir / f"{context.base_filename}_transcript_{timestamp}.json"
        
        # Create JSON data structure
        json_data = {
            "metadata": {
                "filename": context.original_filename,
                "processed_at": timestamp,
                **metadata
            },
            "raw_transcript": raw_transcript,
            "enhanced_transcript": enhanced_transcript,
            "segments": segments
        }
        
        # Write JSON file
        async with aiofiles.open(str(json_path), 'w') as f:
            await f.write(json.dumps(json_data, indent=2))
            
        artifact_paths["output_files"]["json"] = json_path
    
    # Generate TEXT output
    if OutputFormat.TEXT in output_formats:
        text_path = output_dir / f"{context.base_filename}_transcript_{timestamp}.txt"
        
        # Create plain text file
        async with aiofiles.open(str(text_path), 'w') as f:
            # Add metadata header if available
            if metadata:
                if "title" in metadata and metadata["title"]:
                    await f.write(f"Title: {metadata['title']}\n")
                if "language" in metadata and metadata["language"]:
                    await f.write(f"Language: {metadata['language']}\n")
                if "topics" in metadata and metadata["topics"]:
                    topics_str = ", ".join(metadata["topics"])
                    await f.write(f"Topics: {topics_str}\n")
                await f.write("\n")
            
            # Write enhanced transcript if available, otherwise use raw transcript
            await f.write(enhanced_transcript or raw_transcript)
        
        artifact_paths["output_files"]["text"] = text_path
    
    # Generate SRT output
    if OutputFormat.SRT in output_formats:
        srt_path = output_dir / f"{context.base_filename}_transcript_{timestamp}.srt"
        
        # Convert segments to SRT format
        srt_content = generate_srt(segments)
        
        # Write SRT file
        async with aiofiles.open(str(srt_path), 'w') as f:
            await f.write(srt_content)
            
        artifact_paths["output_files"]["srt"] = srt_path
    
    # Generate VTT output
    if OutputFormat.VTT in output_formats:
        vtt_path = output_dir / f"{context.base_filename}_transcript_{timestamp}.vtt"
        
        # Convert segments to VTT format
        vtt_content = generate_vtt(segments)
        
        # Write VTT file
        async with aiofiles.open(str(vtt_path), 'w') as f:
            await f.write(vtt_content)
            
        artifact_paths["output_files"]["vtt"] = vtt_path
    
    # Generate Markdown output
    if OutputFormat.MARKDOWN in output_formats:
        md_path = output_dir / f"{context.base_filename}_transcript_{timestamp}.md"
        
        # Create markdown content
        md_content = generate_markdown(enhanced_transcript, metadata)
        
        # Write markdown file
        async with aiofiles.open(str(md_path), 'w') as f:
            await f.write(md_content)
            
        artifact_paths["output_files"]["markdown"] = md_path
    
    # Generate DOCX output (if supported)
    if OutputFormat.DOCX in output_formats:
        try:
            docx_path = output_dir / f"{context.base_filename}_transcript_{timestamp}.docx"
            
            # Generate DOCX in a thread pool to avoid blocking
            with concurrent.futures.ThreadPoolExecutor() as executor:
                await asyncio.get_event_loop().run_in_executor(
                    executor, 
                    generate_docx, 
                    docx_path, 
                    enhanced_transcript, 
                    metadata
                )
                
            artifact_paths["output_files"]["docx"] = docx_path
        except (ImportError, Exception) as e:
            logger.warning(f"Failed to generate DOCX output: {e}", emoji_key="warning")
    
    return artifact_paths


def generate_srt(segments: List[Dict[str, Any]]) -> str:
    """Generate SRT format from segments."""
    srt_lines = []
    
    for i, segment in enumerate(segments):
        # Convert times to SRT format (HH:MM:SS,mmm)
        start_time = segment.get("start", 0)
        end_time = segment.get("end", 0)
        
        start_str = format_srt_time(start_time)
        end_str = format_srt_time(end_time)
        
        # Format text
        text = segment.get("text", "").replace("\n", " ")
        
        # Add speaker if available
        if "speaker" in segment and segment["speaker"]:
            text = f"[{segment['speaker']}] {text}"
        
        # Add to SRT
        srt_lines.append(f"{i+1}")
        srt_lines.append(f"{start_str} --> {end_str}")
        srt_lines.append(f"{text}")
        srt_lines.append("")  # Empty line between entries
    
    return "\n".join(srt_lines)


def format_srt_time(seconds: float) -> str:
    """Format seconds as SRT time: HH:MM:SS,mmm."""
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    sec_int = int(seconds % 60)
    ms = int((seconds % 1) * 1000)
    return f"{hours:02d}:{minutes:02d}:{sec_int:02d},{ms:03d}"


def generate_vtt(segments: List[Dict[str, Any]]) -> str:
    """Generate WebVTT format from segments."""
    vtt_lines = ["WEBVTT", ""]
    
    for segment in segments:
        # Convert times to VTT format (HH:MM:SS.mmm)
        start_time = segment.get("start", 0)
        end_time = segment.get("end", 0)
        
        start_str = format_vtt_time(start_time)
        end_str = format_vtt_time(end_time)
        
        # Format text
        text = segment.get("text", "").replace("\n", " ")
        
        # Add speaker if available
        if "speaker" in segment and segment["speaker"]:
            text = f"<v {segment['speaker']}>{text}</v>"
        
        # Add to VTT
        vtt_lines.append(f"{start_str} --> {end_str}")
        vtt_lines.append(f"{text}")
        vtt_lines.append("")  # Empty line between entries
    
    return "\n".join(vtt_lines)


def format_vtt_time(seconds: float) -> str:
    """Format seconds as WebVTT time: HH:MM:SS.mmm."""
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    sec_fractional = seconds % 60
    return f"{hours:02d}:{minutes:02d}:{sec_fractional:06.3f}"


def generate_markdown(transcript: str, metadata: Dict[str, Any]) -> str:
    """Generate Markdown format for the transcript."""
    md_lines = []
    
    # Add title
    if "title" in metadata and metadata["title"]:
        md_lines.append(f"# {metadata['title']}")
        md_lines.append("")
    else:
        md_lines.append("# Transcript")
        md_lines.append("")
    
    # Add metadata section
    md_lines.append("## Metadata")
    md_lines.append("")
    
    if "language" in metadata and metadata["language"]:
        md_lines.append(f"- **Language:** {metadata['language']}")
    
    if "duration" in metadata and metadata["duration"]:
        duration_min = int(metadata["duration"] // 60)
        duration_sec = int(metadata["duration"] % 60)
        md_lines.append(f"- **Duration:** {duration_min} min {duration_sec} sec")
    
    if "topics" in metadata and metadata["topics"]:
        topics_str = ", ".join(metadata["topics"])
        md_lines.append(f"- **Topics:** {topics_str}")
    
    md_lines.append("")
    md_lines.append("## Transcript")
    md_lines.append("")
    
    # Add transcript with proper line breaks preserved
    for line in transcript.split("\n"):
        md_lines.append(line)
    
    return "\n".join(md_lines)


def generate_docx(docx_path: str, transcript: str, metadata: Dict[str, Any]) -> None:
    """Generate DOCX format for the transcript."""
    # Must be run in a ThreadPoolExecutor since python-docx is not async
    doc = Document()
    
    # Add title
    if "title" in metadata and metadata["title"]:
        title = doc.add_heading(metadata["title"], 0)
    else:
        title = doc.add_heading("Transcript", 0)  # noqa: F841
    
    # Add metadata section
    doc.add_heading("Metadata", 1)
    
    if "language" in metadata and metadata["language"]:
        doc.add_paragraph(f"Language: {metadata['language']}")
    
    if "duration" in metadata and metadata["duration"]:
        duration_min = int(metadata["duration"] // 60)
        duration_sec = int(metadata["duration"] % 60)
        doc.add_paragraph(f"Duration: {duration_min} min {duration_sec} sec")
    
    if "topics" in metadata and metadata["topics"]:
        topics_str = ", ".join(metadata["topics"])
        doc.add_paragraph(f"Topics: {topics_str}")
    
    # Add transcript
    doc.add_heading("Transcript", 1)
    
    # Split into paragraphs and add
    for paragraph in transcript.split("\n\n"):
        if paragraph.strip():
            p = doc.add_paragraph()
            
            # Check if paragraph starts with a heading marker
            if paragraph.startswith("##"):
                parts = paragraph.split(" ", 1)
                if len(parts) > 1:
                    doc.add_heading(parts[1], 2)
                    continue
            
            # Regular paragraph
            p.add_run(paragraph)
    
    # Save the document
    doc.save(docx_path)


async def chat_with_transcript(
    transcript: str,
    query: str,
    provider: str = Provider.ANTHROPIC.value,
    model: Optional[str] = None,
    context: Optional[str] = None
) -> Dict[str, Any]:
    """Chat with a transcript to extract information or answer questions about its content.
    
    Args:
        transcript: The transcript text to analyze
        query: The question or instruction to process regarding the transcript
        provider: LLM provider to use (default: Anthropic)
        model: Specific model to use (default: provider's default model)
        context: Optional additional context about the audio/transcript
        
    Returns:
        A dictionary containing the response and related metadata
    """
    if not transcript or not isinstance(transcript, str):
        raise ToolInputError("Transcript must be a non-empty string.")
    
    if not query or not isinstance(query, str):
        raise ToolInputError("Query must be a non-empty string.")
    
    # Calculate token count for logging
    try:
        transcript_tokens = count_tokens(transcript, model)
        query_tokens = count_tokens(query, model)
        logger.info(
            f"Transcript: {transcript_tokens} tokens, Query: {query_tokens} tokens",
            emoji_key=TaskType.CHAT.value
        )
    except Exception as e:
        logger.warning(f"Failed to count tokens: {e}", emoji_key="warning")
    
    # Build the prompt
    system_prompt = """You are an expert at analyzing transcripts and extracting information.
Provide concise, accurate answers based solely on the provided transcript.
If the answer is not in the transcript, say so clearly."""

    if context:
        system_prompt += f"\n\nAdditional context about this transcript: {context}"
    
    # Get provider instance to ensure it exists and is available
    try:
        provider_instance = await get_provider(provider)
        if model is None:
            # Check if the provider has a default model or use claude-3-7-sonnet as fallback
            default_models = await provider_instance.list_models()
            if default_models and len(default_models) > 0:
                model = default_models[0].get("id")
            else:
                model = "claude-3-7-sonnet-20250219" if provider == Provider.ANTHROPIC.value else None
        
        logger.info(f"Using model: {provider}/{model}", emoji_key="model")
    except Exception as e:
        raise ProviderError(
            f"Failed to initialize provider '{provider}': {str(e)}",
            provider=provider,
            cause=e
        ) from e
    
    # Use relative file paths for any file references
    rel_transcript_path = None
    if os.path.exists(transcript):
        rel_transcript_path = Path(transcript).relative_to(Path.cwd())  # noqa: F841
    
    # Create the message with the transcript and query
    user_message = f"""Here is a transcript to analyze:

---TRANSCRIPT BEGIN---
{transcript}
---TRANSCRIPT END---

{query}"""
    
    # Send to LLM
    result = await chat_completion(
        messages=[{"role": "user", "content": user_message}],
        provider=provider,
        model=model,
        system_prompt=system_prompt,
        temperature=0.1
    )
    
    return result


@with_cache(ttl=24 * 60 * 60)  # Cache results for 24 hours
@with_tool_metrics
@with_retry(max_retries=1, retry_delay=1.0)
@with_error_handling
async def extract_audio_transcript_key_points(
    file_path_or_transcript: str,
    is_file: bool = True,
    provider: str = Provider.ANTHROPIC.value,
    model: Optional[str] = None,
    max_points: int = 10
) -> Dict[str, Any]:
    """Extracts the most important key points from an audio transcript.
    
    This tool can process either an audio file (which it will transcribe first)
    or directly analyze an existing transcript to identify the most important 
    information, main topics, and key takeaways.
    
    Args:
        file_path_or_transcript: Path to audio file or transcript text content
        is_file: Whether the input is a file path (True) or transcript text (False)
        provider: LLM provider to use for analysis
        model: Specific model to use (provider default if None)
        max_points: Maximum number of key points to extract
        
    Returns:
        A dictionary containing:
        {
            "key_points": ["Point 1", "Point 2", ...],
            "summary": "Brief summary of the content",
            "topics": ["Topic 1", "Topic 2", ...],
            "speakers": ["Speaker 1", "Speaker 2", ...] (if multiple speakers detected),
            "tokens": { statistics about token usage },
            "cost": estimated cost of the operation,
            "processing_time": total processing time in seconds
        }
    """
    start_time = time.time()
    
    # Get transcript from file or use provided text
    transcript = ""
    if is_file:
        try:
            # Validate file path
            file_path = os.path.abspath(os.path.expanduser(file_path_or_transcript))
            if not os.path.exists(file_path):
                raise ToolInputError(f"File not found: {file_path}")
            
            # Get file info for logging
            file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
            file_name = Path(file_path).name
            
            logger.info(
                f"Extracting key points from audio file: {file_name} ({file_size_mb:.2f} MB)",
                emoji_key="audio"
            )
            
            # Transcribe audio
            transcription_result = await transcribe_audio(file_path, {
                "enhance_audio": True,
                "enhance_transcript": True,
                "output_formats": ["json"]
            })
            
            transcript = transcription_result.get("enhanced_transcript", "")
            if not transcript:
                transcript = transcription_result.get("raw_transcript", "")
                
            if not transcript:
                raise ToolError("Failed to generate transcript from audio")
                
        except Exception as e:
            if isinstance(e, (ToolError, ToolInputError)):
                raise
            raise ToolError(f"Failed to process audio file: {str(e)}") from e
    else:
        # Input is already a transcript
        transcript = file_path_or_transcript
        if not transcript or not isinstance(transcript, str):
            raise ToolInputError("Transcript text must be a non-empty string")
    
    # Calculate token count for the transcript
    try:
        token_count = count_tokens(transcript, model)
        logger.info(f"Transcript token count: {token_count}", emoji_key="tokens")
    except Exception as e:
        logger.warning(f"Failed to count tokens: {e}")
    
    # Create prompt for key points extraction
    prompt = f"""Extract the most important key points from this transcript.

Identify:
1. The {max_points} most important key points or takeaways
2. Main topics discussed
3. Any speakers or main entities mentioned (if identifiable)
4. A brief summary (2-3 sentences max)

Format your response as JSON with these fields:
{{
  "key_points": ["Point 1", "Point 2", ...],
  "topics": ["Topic 1", "Topic 2", ...],
  "speakers": ["Speaker 1", "Speaker 2", ...],
  "summary": "Brief summary here"
}}

TRANSCRIPT:
{transcript}
"""

    # Get provider instance
    try:
        provider_instance = await get_provider(provider)  # noqa: F841
    except Exception as e:
        raise ProviderError(
            f"Failed to initialize provider '{provider}': {str(e)}",
            provider=provider,
            cause=e
        ) from e

    # Generate completion
    try:
        completion_result = await generate_completion(
            prompt=prompt,
            provider=provider,
            model=model,
            temperature=0.1,
            max_tokens=1000
        )
        
        # Parse JSON response
        response_text = completion_result.get("text", "")
        
        # Find JSON in the response
        json_match = re.search(r'({[\s\S]*})', response_text)
        if json_match:
            try:
                extracted_data = json.loads(json_match.group(1))
            except json.JSONDecodeError:
                # Fallback to regex extraction
                extracted_data = _extract_key_points_with_regex(response_text)
        else:
            # Fallback to regex extraction
            extracted_data = _extract_key_points_with_regex(response_text)
        
        processing_time = time.time() - start_time
        
        # Add token and cost info
        result = {
            "key_points": extracted_data.get("key_points", []),
            "topics": extracted_data.get("topics", []),
            "speakers": extracted_data.get("speakers", []),
            "summary": extracted_data.get("summary", ""),
            "tokens": completion_result.get("tokens", {"input": 0, "output": 0, "total": 0}),
            "cost": completion_result.get("cost", 0.0),
            "processing_time": processing_time
        }
        
        return result
        
    except Exception as e:
        error_model = model or f"{provider}/default"
        raise ProviderError(
            f"Key points extraction failed for model '{error_model}': {str(e)}",
            provider=provider,
            model=error_model,
            cause=e
        ) from e


def _extract_key_points_with_regex(text: str) -> Dict[str, Any]:
    """Extract key points data using regex when JSON parsing fails."""
    result = {
        "key_points": [],
        "topics": [],
        "speakers": [],
        "summary": ""
    }
    
    # Extract key points
    key_points_pattern = r'key_points"?\s*:?\s*\[\s*"([^"]+)"(?:\s*,\s*"([^"]+)")*\s*\]'
    key_points_match = re.search(key_points_pattern, text, re.IGNORECASE | re.DOTALL)
    if key_points_match:
        point_list = re.findall(r'"([^"]+)"', key_points_match.group(0))
        # Filter out empty strings
        result["key_points"] = [p for p in point_list if p.strip()]
    else:
        # Try alternative pattern for non-JSON format
        point_list = re.findall(r'(?:^|\n)(?:•|\*|-|[0-9]+\.)\s*([^\n]+?)(?:\n|$)', text)
        # Filter out empty strings
        result["key_points"] = [p.strip() for p in point_list if p.strip()][:10]  # Limit to 10 points
    
    # Extract topics
    topics_pattern = r'topics"?\s*:?\s*\[\s*"([^"]+)"(?:\s*,\s*"([^"]+)")*\s*\]'
    topics_match = re.search(topics_pattern, text, re.IGNORECASE | re.DOTALL)
    if topics_match:
        topic_list = re.findall(r'"([^"]+)"', topics_match.group(0))
        # Filter out empty strings
        result["topics"] = [t for t in topic_list if t.strip()]
    
    # Extract speakers
    speakers_pattern = r'speakers"?\s*:?\s*\[\s*"([^"]+)"(?:\s*,\s*"([^"]+)")*\s*\]'
    speakers_match = re.search(speakers_pattern, text, re.IGNORECASE | re.DOTALL)
    if speakers_match:
        speaker_list = re.findall(r'"([^"]+)"', speakers_match.group(0))
        # Filter out empty strings
        result["speakers"] = [s for s in speaker_list if s.strip()]
    
    # Extract summary
    summary_pattern = r'summary"?\s*:?\s*"([^"]+)"'
    summary_match = re.search(summary_pattern, text, re.IGNORECASE)
    if summary_match and summary_match.group(1).strip():
        result["summary"] = summary_match.group(1).strip()
    
    return result