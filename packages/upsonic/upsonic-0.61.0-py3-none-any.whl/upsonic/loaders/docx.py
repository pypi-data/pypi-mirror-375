from __future__ import annotations
from typing import List, Any, Dict, Optional
import os
import time
from datetime import datetime

from .base import DocumentLoader
from .config import DOCXLoaderConfig
from ..schemas.data_models import Document

try:
    import docx
    from docx.table import Table as DocxTable
except ImportError:
    raise ImportError(
        "python-docx is not installed. It is required for the DOCXLoader. "
        "Please run: 'pip install upsonic[docx]'"
    )


class DOCXLoader(DocumentLoader):
    """
    A master-class, structure-aware ingestion engine for Microsoft Word (`.docx`) files.

    This loader intelligently traverses the document's structure, going beyond
    simple text extraction. It is built on three pillars:
    1.  **Full Document Traversal:** Iterates through all block-level items,
        identifying both paragraphs and tables.
    2.  **Intelligent Table-to-Text Conversion:** Parses `docx` tables and
        serializes them into a clean, LLM-friendly textual format.
    3.  **Deep Metadata Archaeology:** Extracts both filesystem metadata and the
        document's core properties (e.g., author, title).
    """

    def __init__(self, config: Optional[DOCXLoaderConfig] = None):
        """
        Initialize the DOCXLoader with configuration.
        
        Args:
            config: DOCXLoaderConfig object with all configuration options
        """
        if config is None:
            config = DOCXLoaderConfig()
        
        
        super().__init__(config)
        self.config = config

    def load(self, source: str) -> List[Document]:
        """
        Loads a .docx file, parsing its structure into a single, rich Document object.
        """
        return self._load_with_error_handling(source)

    def _load_with_error_handling(self, source: str) -> List[Document]:
        """Override to call internal method instead of public load method."""
        start_time = time.time()
        
        try:
            if not self._validate_source(source):
                raise ValueError(f"Invalid source: {source}")
            
            if self.config and self.config.max_file_size:
                if os.path.exists(source):
                    file_size = os.path.getsize(source)
                    if file_size > self.config.max_file_size:
                        raise ValueError(f"File size {file_size} exceeds limit {self.config.max_file_size}")
            
            documents = self._load_document_internal(source)
            
            documents = self._post_process_documents(documents, source)
            
            processing_time = time.time() - start_time
            self._update_stats(len(documents), processing_time, success=True)
            
            return documents
            
        except Exception as e:
            processing_time = time.time() - start_time
            self._update_stats(0, processing_time, success=False)
            
            if self.config:
                if self.config.error_handling == "ignore":
                    return []
                elif self.config.error_handling == "warn":
                    print(f"Warning: Failed to load {source}: {e}")
                    return []
                else:
                    raise
            else:
                raise

    def _load_document_internal(self, source: str) -> List[Document]:
        """
        Internal method to load a .docx file, parsing its structure into a single, rich Document object.
        """
        try:
            file_path = os.path.abspath(source)
            if not os.path.isfile(file_path):
                raise FileNotFoundError(f"Source path '{source}' is not a valid file.")

            stats = os.stat(file_path)
            base_metadata: Dict[str, Any] = {
                "source": source,
                "file_name": os.path.basename(file_path),
                "file_path": file_path,
                "file_size": stats.st_size,
                "creation_time": datetime.fromtimestamp(stats.st_ctime).isoformat(),
                "last_modified_time": datetime.fromtimestamp(stats.st_mtime).isoformat(),
            }

            doc = docx.Document(file_path)

            cp = doc.core_properties
            core_meta = {
                "author": cp.author, "category": cp.category,
                "comments": cp.comments, "title": cp.title,
                "subject": cp.subject, "keywords": cp.keywords,
                "last_modified_by": cp.last_modified_by,
            }
            base_metadata.update({k: v for k, v in core_meta.items() if v})

            content_parts: List[str] = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content_parts.append(paragraph.text)
            
            for table in doc.tables:
                table_text = self._parse_table_to_text(table)
                if table_text:
                    content_parts.append(table_text)
            
            full_content = "\n\n".join(part for part in content_parts if part and part.strip())

            return [Document(content=full_content, metadata=base_metadata)]

        except FileNotFoundError as e:
            print(f"Error: [DOCXLoader] File not found at path: {e}")
            return []
        except Exception as e:
            print(f"Error: [DOCXLoader] An unexpected error occurred while loading '{source}': {e}")
            return []

    def _parse_table_to_text(self, table: DocxTable) -> str:
        """Converts a docx.table.Table object into a human-readable text block."""
        try:
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            
            text_rows = []
            for row in table.rows[1:]:
                row_cells = [cell.text.strip() for cell in row.cells]
                if len(row_cells) == len(headers):
                    row_text = ", ".join([f"{headers[i]}: {cell}" for i, cell in enumerate(row_cells)])
                    text_rows.append(f"- {row_text}")
            
            if not text_rows:
                return ""

            return f"[Structured Table Data]:\n" + "\n".join(text_rows)
        except Exception:
            return "[Unable to parse table content]"

    @classmethod
    def get_supported_extensions(cls) -> List[str]:
        """Get list of file extensions supported by this loader."""
        return ['.docx']