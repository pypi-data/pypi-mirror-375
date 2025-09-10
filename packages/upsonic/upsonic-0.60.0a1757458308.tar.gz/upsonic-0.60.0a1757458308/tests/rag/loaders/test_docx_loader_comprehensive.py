import os
import sys
import tempfile
import time
from pathlib import Path
import unittest

from upsonic.loaders.docx import DOCXLoader
from upsonic.loaders.config import DOCXLoaderConfig
from upsonic.schemas.data_models import Document

try:
    import docx
    from docx import Document as DocxDocument
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    raise ImportError("python-docx is not installed. Please install it with: pip install python-docx")

class TestDOCXLoaderComprehensive(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.temp_dir = tempfile.mkdtemp()
        cls.test_files = {}
        cls.create_basic_docx(cls)
        cls.create_table_docx(cls)
        cls.create_complex_docx(cls)
        cls.create_empty_docx(cls)
        cls.create_metadata_docx(cls)
        cls.create_large_docx(cls)
        cls.create_special_formatting_docx(cls)

    @classmethod
    def tearDownClass(cls):
        if cls.temp_dir and os.path.exists(cls.temp_dir):
            import shutil
            shutil.rmtree(cls.temp_dir)

    @staticmethod
    def create_basic_docx(self):
        file_path = os.path.join(self.temp_dir, "basic_test.docx")
        doc = DocxDocument()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph('This is a basic test document for the DOCX loader.')
        doc.add_paragraph('It contains multiple paragraphs with different content.')
        doc.add_paragraph('The loader should extract all text content properly.')
        doc.add_heading('Section 1: Introduction', level=1)
        doc.add_paragraph('This is the introduction section with some important information.')
        doc.add_heading('Section 2: Details', level=2)
        doc.add_paragraph('Here are the detailed information and specifications.')
        doc.save(file_path)
        self.test_files['basic'] = file_path

    @staticmethod
    def create_table_docx(self):
        file_path = os.path.join(self.temp_dir, "table_test.docx")
        doc = DocxDocument()
        doc.add_heading('Table Test Document', 0)
        doc.add_paragraph('This document contains various tables for testing.')
        table1 = doc.add_table(rows=4, cols=3)
        table1.style = 'Table Grid'
        hdr_cells = table1.rows[0].cells
        hdr_cells[0].text = 'Name'
        hdr_cells[1].text = 'Age'
        hdr_cells[2].text = 'City'
        data = [
            ['John Doe', '30', 'New York'],
            ['Jane Smith', '25', 'San Francisco'],
            ['Bob Johnson', '35', 'Chicago']
        ]
        for i, row_data in enumerate(data, 1):
            row_cells = table1.rows[i].cells
            for j, cell_data in enumerate(row_data):
                row_cells[j].text = cell_data
        doc.add_paragraph('Here is another table with different data:')
        table2 = doc.add_table(rows=3, cols=4)
        table2.style = 'Table Grid'
        hdr_cells2 = table2.rows[0].cells
        hdr_cells2[0].text = 'Product'
        hdr_cells2[1].text = 'Price'
        hdr_cells2[2].text = 'Stock'
        hdr_cells2[3].text = 'Category'
        data2 = [
            ['Laptop', '$999', '50', 'Electronics'],
            ['Book', '$15', '200', 'Education']
        ]
        for i, row_data in enumerate(data2, 1):
            row_cells = table2.rows[i].cells
            for j, cell_data in enumerate(row_data):
                row_cells[j].text = cell_data
        doc.save(file_path)
        self.test_files['table'] = file_path

    @staticmethod
    def create_complex_docx(self):
        file_path = os.path.join(self.temp_dir, "complex_test.docx")
        doc = DocxDocument()
        doc.add_heading('Complex Document Structure', 0)
        doc.add_paragraph('This is a complex document that tests various DOCX features.')
        doc.add_heading('Executive Summary', level=1)
        doc.add_paragraph('This section contains the executive summary of our findings.')
        doc.add_paragraph('Key points include:')
        doc.add_paragraph('• First important point', style='List Bullet')
        doc.add_paragraph('• Second important point', style='List Bullet')
        doc.add_paragraph('• Third important point', style='List Bullet')
        doc.add_heading('Data Analysis', level=2)
        doc.add_paragraph('The following table shows our analysis results:')
        table = doc.add_table(rows=5, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Metric'
        hdr_cells[1].text = 'Value'
        hdr_cells[2].text = 'Status'
        analysis_data = [
            ['Revenue', '$1,000,000', 'Good'],
            ['Costs', '$750,000', 'Acceptable'],
            ['Profit', '$250,000', 'Excellent'],
            ['Growth Rate', '15%', 'Positive']
        ]
        for i, row_data in enumerate(analysis_data, 1):
            row_cells = table.rows[i].cells
            for j, cell_data in enumerate(row_data):
                row_cells[j].text = cell_data
        doc.add_heading('Conclusion', level=1)
        doc.add_paragraph('Based on our analysis, we recommend the following actions.')
        doc.add_paragraph('This concludes our comprehensive report.')
        doc.save(file_path)
        self.test_files['complex'] = file_path

    @staticmethod
    def create_empty_docx(self):
        file_path = os.path.join(self.temp_dir, "empty_test.docx")
        doc = DocxDocument()
        doc.save(file_path)
        self.test_files['empty'] = file_path

    @staticmethod
    def create_metadata_docx(self):
        file_path = os.path.join(self.temp_dir, "metadata_test.docx")
        doc = DocxDocument()
        doc.core_properties.author = "Test Author"
        doc.core_properties.title = "Test Document with Metadata"
        doc.core_properties.subject = "Testing DOCX Loader"
        doc.core_properties.keywords = "test, docx, loader, metadata"
        doc.core_properties.comments = "This is a test document for metadata extraction"
        doc.core_properties.category = "Test Documents"
        doc.core_properties.last_modified_by = "Test User"
        doc.add_heading('Metadata Test Document', 0)
        doc.add_paragraph('This document contains rich metadata for testing.')
        doc.add_paragraph('The loader should extract all document properties.')
        doc.save(file_path)
        self.test_files['metadata'] = file_path

    @staticmethod
    def create_large_docx(self):
        file_path = os.path.join(self.temp_dir, "large_test.docx")
        doc = DocxDocument()
        doc.add_heading('Large Document for Performance Testing', 0)
        for i in range(100):
            doc.add_heading(f'Section {i+1}', level=1)
            doc.add_paragraph(f'This is paragraph {i+1} of the large document.')
            doc.add_paragraph(f'It contains detailed information about topic {i+1}.')
            doc.add_paragraph(f'Additional details and specifications for section {i+1}.')
        for i in range(10):
            doc.add_heading(f'Data Table {i+1}', level=2)
            table = doc.add_table(rows=6, cols=4)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'ID'
            hdr_cells[1].text = 'Name'
            hdr_cells[2].text = 'Value'
            hdr_cells[3].text = 'Status'
            for j in range(1, 6):
                row_cells = table.rows[j].cells
                row_cells[0].text = str(j)
                row_cells[1].text = f'Item {j}'
                row_cells[2].text = str(100 + j)
                row_cells[3].text = 'Active'
        doc.save(file_path)
        self.test_files['large'] = file_path

    @staticmethod
    def create_special_formatting_docx(self):
        file_path = os.path.join(self.temp_dir, "special_formatting_test.docx")
        doc = DocxDocument()
        doc.add_heading('Special Formatting Test', 0)
        doc.add_paragraph('This document tests special formatting and edge cases.')
        doc.add_paragraph('')
        doc.add_paragraph('   ')
        doc.add_paragraph('Normal paragraph with content.')
        doc.add_paragraph('Special characters: é, ñ, ü, ç, €, £, ¥, ©, ®, ™')
        doc.add_paragraph('Numbers: 123, 45.67, 1,000,000, 1.23e-4')
        doc.add_paragraph('Symbols: @#$%^&*()_+-=[]{}|;:,.<>?')
        table = doc.add_table(rows=3, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Column 1'
        hdr_cells[1].text = 'Column 2'
        hdr_cells[2].text = 'Column 3'
        row1_cells = table.rows[1].cells
        row1_cells[0].text = 'Data 1'
        row1_cells[1].text = ''
        row1_cells[2].text = 'Data 3'
        row2_cells = table.rows[2].cells
        row2_cells[0].text = ''
        row2_cells[1].text = 'Data 2'
        row2_cells[2].text = ''
        doc.save(file_path)
        self.test_files['special_formatting'] = file_path

    def test_basic_loading(self):
        loader = DOCXLoader()
        documents = loader.load(self.test_files['basic'])
        self.assertEqual(len(documents), 1)
        doc = documents[0]
        self.assertTrue(doc.content)
        self.assertIn('Test Document', doc.content)
        self.assertIn('basic test document', doc.content.lower())
        self.assertIn('Section 1: Introduction', doc.content)
        self.assertIn('Section 2: Details', doc.content)
        self.assertEqual(doc.metadata['file_name'], 'basic_test.docx')
        self.assertIn('file_size', doc.metadata)
        self.assertIn('creation_time', doc.metadata)

    def test_table_parsing(self):
        loader = DOCXLoader()
        documents = loader.load(self.test_files['table'])
        self.assertEqual(len(documents), 1)
        doc = documents[0]
        content = doc.content
        self.assertIn('[Structured Table Data]', content)
        self.assertIn('Name: John Doe', content)
        self.assertIn('Age: 30', content)
        self.assertIn('Product: Laptop', content)
        self.assertIn('Price: $999', content)
        table_count = content.count('[Structured Table Data]')
        self.assertEqual(table_count, 2)

    def test_complex_structure(self):
        loader = DOCXLoader()
        documents = loader.load(self.test_files['complex'])
        self.assertEqual(len(documents), 1)
        doc = documents[0]
        content = doc.content
        self.assertIn('Complex Document Structure', content)
        self.assertIn('Executive Summary', content)
        self.assertIn('Data Analysis', content)
        self.assertIn('First important point', content)
        self.assertIn('Revenue', content)
        self.assertIn('Conclusion', content)
        self.assertIn('\n\n', content)

    def test_empty_document(self):
        loader = DOCXLoader()
        documents = loader.load(self.test_files['empty'])
        if len(documents) == 0:
            pass
        else:
            self.assertEqual(len(documents), 1)
            doc = documents[0]
            self.assertIsInstance(doc.content, str)
            self.assertEqual(doc.metadata['file_name'], 'empty_test.docx')

    def test_metadata_extraction(self):
        loader = DOCXLoader()
        documents = loader.load(self.test_files['metadata'])
        self.assertEqual(len(documents), 1)
        doc = documents[0]
        metadata = doc.metadata
        self.assertEqual(metadata.get('author'), 'Test Author')
        self.assertEqual(metadata.get('title'), 'Test Document with Metadata')
        self.assertEqual(metadata.get('subject'), 'Testing DOCX Loader')
        self.assertEqual(metadata.get('keywords'), 'test, docx, loader, metadata')
        self.assertEqual(metadata.get('comments'), 'This is a test document for metadata extraction')
        self.assertEqual(metadata.get('category'), 'Test Documents')
        self.assertEqual(metadata.get('last_modified_by'), 'Test User')

    def test_special_formatting(self):
        loader = DOCXLoader()
        documents = loader.load(self.test_files['special_formatting'])
        self.assertEqual(len(documents), 1)
        doc = documents[0]
        content = doc.content
        self.assertIn('é, ñ, ü, ç', content)
        self.assertIn('€, £, ¥', content)
        self.assertIn('©, ®, ™', content)
        self.assertIn('123, 45.67', content)
        self.assertIn('@#$%^&*()', content)
        self.assertIn('[Structured Table Data]', content)
        self.assertIn('Data 1', content)
        self.assertIn('Data 2', content)

    def test_performance(self):
        loader = DOCXLoader()
        start_time = time.time()
        documents = loader.load(self.test_files['large'])
        end_time = time.time()
        processing_time = end_time - start_time
        self.assertEqual(len(documents), 1)
        self.assertLess(processing_time, 10.0)
        stats = loader.get_stats()
        self.assertEqual(stats['total_files_processed'], 1)
        self.assertEqual(stats['total_documents_created'], 1)
        doc = documents[0]
        content = doc.content
        self.assertIn('Large Document for Performance Testing', content)
        self.assertIn('Section 1', content)
        self.assertIn('Section 100', content)
        self.assertIn('Data Table 1', content)
        self.assertIn('Data Table 10', content)

    def test_file_not_found(self):
        config_ignore = DOCXLoaderConfig(error_handling="ignore")
        loader_ignore = DOCXLoader(config_ignore)
        docs = loader_ignore.load("non_existent_file.docx")
        self.assertEqual(len(docs), 0)
        config_warn = DOCXLoaderConfig(error_handling="warn")
        loader_warn = DOCXLoader(config_warn)
        docs_warn = loader_warn.load("non_existent_file.docx")
        self.assertEqual(len(docs_warn), 0)
        config_raise = DOCXLoaderConfig(error_handling="raise")
        loader_raise = DOCXLoader(config_raise)
        try:
            docs_raise = loader_raise.load("non_existent_file.docx")
            self.assertEqual(len(docs_raise), 0)
        except Exception as e:
            self.assertTrue("Invalid source" in str(e) or "not a valid file" in str(e))

    def test_configuration_options(self):
        custom_metadata = {"source_type": "test", "version": "1.0"}
        config_custom = DOCXLoaderConfig(custom_metadata=custom_metadata)
        loader_custom = DOCXLoader(config_custom)
        docs_custom = loader_custom.load(self.test_files['basic'])
        self.assertEqual(len(docs_custom), 1)
        doc = docs_custom[0]
        self.assertEqual(doc.metadata['source_type'], 'test')
        self.assertEqual(doc.metadata['version'], '1.0')
        config_size = DOCXLoaderConfig(max_file_size=1000)
        loader_size = DOCXLoader(config_size)
        docs_size = loader_size.load(self.test_files['large'])
        self.assertEqual(len(docs_size), 0)

    def test_batch_loading(self):
        loader = DOCXLoader()
        sources = [self.test_files['basic'], self.test_files['table'], self.test_files['metadata']]
        results = loader.load_batch(sources)
        self.assertEqual(len(results), 3)
        self.assertTrue(all(result.success for result in results))
        self.assertTrue(all(len(result.documents) == 1 for result in results))

    def test_directory_loading(self):
        loader = DOCXLoader()
        results = loader.load_directory(self.temp_dir, file_patterns=["*.docx"])
        self.assertGreaterEqual(len(results), 7)
        successful_results = [result for result in results if result.success]
        total_docs = sum(len(result.documents) for result in successful_results)
        self.assertGreaterEqual(len(successful_results), 6)
        self.assertGreaterEqual(total_docs, 6)

if __name__ == "__main__":
    unittest.main()
